#!/usr/bin/env python3
"""Small browser UI for the local Ollama tools."""

from __future__ import annotations

import asyncio
import io
import json
import os
import re
import secrets
import socket
import shutil
import subprocess
import sys
import tempfile
import threading
import urllib.error
import urllib.request
import wave
import ipaddress
from dataclasses import asdict
from datetime import datetime
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

try:
    import edge_tts as _edge_tts
    EDGE_TTS_AVAILABLE = True
except ImportError:
    EDGE_TTS_AVAILABLE = False

# ── Kokoro self-hosted TTS ────────────────────────────────────────────────────
# Set KOKORO_URL to your LAN server, e.g. http://tts-host:8880
KOKORO_URL = os.environ.get("KOKORO_URL", "")  # resolved after imports below

KOKORO_FEMALE_VOICES = [
    {"name": "preset:lilith_dark", "label": "Lilith Dark — low, slow preset (offline if Piper/Kokoro is configured)"},
    {"name": "af_sky",     "label": "Sky — warm & expressive (Piper/Kokoro offline)"},
    {"name": "af_nova",    "label": "Nova — clear & bright (Piper/Kokoro offline)"},
    {"name": "af_bella",   "label": "Bella — soft & natural (Piper/Kokoro offline)"},
    {"name": "af_jessica", "label": "Jessica — friendly (Piper/Kokoro offline)"},
    {"name": "af_sarah",   "label": "Sarah — calm & professional (Piper/Kokoro offline)"},
    {"name": "af_nicole",  "label": "Nicole — conversational (Piper/Kokoro offline)"},
    {"name": "af_hfc_female", "label": "HFC Female — darker US voice (Piper offline)"},
    {"name": "af_lessac",  "label": "Lessac — clear US voice (Piper offline)"},
    {"name": "af_amy",     "label": "Amy — natural US English (Piper offline)"},
    {"name": "bf_emma",    "label": "Emma — British English (Piper/Kokoro offline)"},
    {"name": "bf_isabella","label": "Isabella — British English (Piper/Kokoro offline)"},
]

# Edge TTS voices — internet required
EDGE_TTS_FEMALE_VOICES = [
    {"name": "en-US-AriaNeural",    "label": "Aria — US, natural & expressive (Edge)"},
    {"name": "en-US-JennyNeural",   "label": "Jenny — US, friendly & warm (Edge)"},
    {"name": "en-US-MichelleNeural","label": "Michelle — US, clear & professional (Edge)"},
    {"name": "en-US-MonicaNeural",  "label": "Monica — US, conversational (Edge)"},
    {"name": "en-GB-SoniaNeural",   "label": "Sonia — British English (Edge)"},
    {"name": "en-AU-NatashaNeural", "label": "Natasha — Australian English (Edge)"},
    {"name": "en-IE-EmilyNeural",   "label": "Emily — Irish English (Edge)"},
    {"name": "en-CA-ClaraNeural",   "label": "Clara — Canadian English (Edge)"},
    {"name": "en-IN-NeerjaNeural",  "label": "Neerja — Indian English (Edge)"},
]

_KOKORO_VOICE_NAMES = {v["name"] for v in KOKORO_FEMALE_VOICES}
_EDGE_VOICE_NAMES   = {v["name"] for v in EDGE_TTS_FEMALE_VOICES}
_PIPER_DYNAMIC_VOICE_NAMES: set[str] = set()
_PIPER_LAST_VOICE_NAMES: set[str] = set()

# Speed: slider sends -75..+100 as percent string; convert for each backend
def _rate_to_kokoro_speed(rate_str: str) -> float:
    """Convert '+25%' style string to Kokoro speed float (1.0 = normal)."""
    try:
        pct = int(rate_str.replace("%", "").replace("+", ""))
        return max(0.5, min(2.0, 1.0 + pct / 100))
    except ValueError:
        return 1.0


def kokoro_tts_speak(text: str, voice: str, rate: str = "+0%") -> tuple[bytes, str]:
    """Call the self-hosted TTS endpoint (Kokoro or Piper).

    Returns (audio_bytes, content_type).
    """
    import urllib.request
    speed = _rate_to_kokoro_speed(rate)
    body = json.dumps({"model": "kokoro", "input": text, "voice": voice, "speed": speed}).encode()
    req = urllib.request.Request(
        f"{KOKORO_URL}/v1/audio/speech",
        data=body,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            content_type = resp.headers.get("Content-Type", "audio/mpeg")
            return resp.read(), content_type
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore").strip()
        raise RuntimeError(f"Piper/Kokoro TTS failed for voice '{voice}': {detail or exc.reason}") from exc


def kokoro_voice_names() -> list[str]:
    if not KOKORO_URL:
        return []
    try:
        with urllib.request.urlopen(f"{KOKORO_URL}/v1/voices", timeout=3) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception:
        return []
    voices = data.get("voices") if isinstance(data, dict) else []
    if not isinstance(voices, list):
        return []
    return sorted({str(item) for item in voices if str(item).strip()})


def _clamp_percent(rate_str: str, delta: int = 0) -> str:
    try:
        pct = int(rate_str.replace("%", "").replace("+", ""))
    except ValueError:
        pct = 0
    pct = max(-75, min(100, pct + delta))
    return f"{pct:+d}%"


def _resolve_tts_request(voice: str, rate: str) -> tuple[str, str, str, bool]:
    """Return backend voice, rate, pitch, and whether Kokoro should be skipped."""
    if voice == "preset:lilith_dark":
        if KOKORO_URL:
            for candidate in ("af_hfc_female", "hfc_female", "en_US-hfc_female-medium", "af_lessac", "lessac", "en_US-lessac-medium"):
                if candidate in _PIPER_LAST_VOICE_NAMES or candidate in _PIPER_DYNAMIC_VOICE_NAMES:
                    return candidate, _clamp_percent(rate, -25), "+0Hz", False
            available = sorted(_PIPER_LAST_VOICE_NAMES or _PIPER_DYNAMIC_VOICE_NAMES)
            if available:
                return available[0], _clamp_percent(rate, -25), "+0Hz", False
            return "af_hfc_female", _clamp_percent(rate, -25), "+0Hz", False
        return "en-US-MichelleNeural", _clamp_percent(rate, -15), "-18Hz", True
    if voice in _EDGE_VOICE_NAMES:
        return voice, rate, "+0Hz", True
    return voice, rate, "+0Hz", False


def _clamp_int(value: object, low: int, high: int, default: int = 0) -> int:
    try:
        ivalue = int(value)
    except (TypeError, ValueError):
        return default
    return max(low, min(high, ivalue))


def _audio_sample_rate(audio: bytes) -> int:
    try:
        with wave.open(io.BytesIO(audio), "rb") as wav:
            rate = int(wav.getframerate())
            if 8_000 <= rate <= 192_000:
                return rate
    except (EOFError, wave.Error):
        pass
    ffprobe = shutil.which("ffprobe")
    if ffprobe:
        proc = subprocess.run(
            [
                ffprobe,
                "-v", "error",
                "-select_streams", "a:0",
                "-show_entries", "stream=sample_rate",
                "-of", "default=noprint_wrappers=1:nokey=1",
                "-i", "pipe:0",
            ],
            input=audio,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            check=False,
        )
        try:
            rate = int(proc.stdout.decode("utf-8", errors="ignore").strip().splitlines()[0])
            if 8_000 <= rate <= 192_000:
                return rate
        except (IndexError, ValueError):
            pass
    return 48_000


def _mixer_filter(*, pitch: int = 0, tone: str = "natural", volume: int = 0, sample_rate: int = 48_000) -> str:
    filters: list[str] = []
    if pitch:
        factor = 2 ** (pitch / 12)
        tempo = max(0.5, min(2.0, 1 / factor))
        sample_rate = max(8_000, min(192_000, sample_rate))
        filters.append(f"asetrate={sample_rate}*{factor:.6f},aresample={sample_rate},atempo={tempo:.6f}")
    if tone == "dark":
        filters.extend(["bass=g=7:f=120:w=0.6", "treble=g=-4:f=3200:w=0.5"])
    elif tone == "bright":
        filters.extend(["bass=g=-3:f=160:w=0.7", "treble=g=5:f=3600:w=0.6"])
    elif tone == "radio":
        filters.extend(["highpass=f=260", "lowpass=f=3600", "acompressor=threshold=-18dB:ratio=3:attack=8:release=80"])
    elif tone == "robotic":
        filters.extend(["aecho=0.6:0.45:18:0.35", "chorus=0.5:0.9:50:0.4:0.25:2"])
    if volume:
        filters.append(f"volume={1 + (volume / 100):.2f}")
    return ",".join(filters)


def mix_audio(audio: bytes, content_type: str, *, pitch: int = 0, tone: str = "natural", volume: int = 0) -> tuple[bytes, str]:
    filters = _mixer_filter(pitch=pitch, tone=tone, volume=volume, sample_rate=_audio_sample_rate(audio))
    if not filters:
        return audio, content_type
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        return audio, content_type
    output_format = "mp3" if "mpeg" in content_type or "mp3" in content_type else "wav"
    output_type = "audio/mpeg" if output_format == "mp3" else "audio/wav"
    proc = subprocess.run(
        [ffmpeg, "-hide_banner", "-loglevel", "error", "-i", "pipe:0", "-af", filters, "-f", output_format, "pipe:1"],
        input=audio,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        check=False,
    )
    if proc.returncode != 0 or not proc.stdout:
        raise RuntimeError(proc.stderr.decode("utf-8", errors="ignore").strip() or "ffmpeg audio mixer failed")
    return proc.stdout, output_type


def edge_tts_speak(text: str, voice: str, rate: str = "+0%", pitch: str = "+0Hz") -> bytes:
    """Synthesise text with edge-tts and return raw MP3 bytes."""
    async def _run() -> bytes:
        chunks: list[bytes] = []
        try:
            communicate = _edge_tts.Communicate(text, voice, rate=rate, pitch=pitch)
        except TypeError:
            communicate = _edge_tts.Communicate(text, voice, rate=rate)
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                chunks.append(chunk["data"])
        return b"".join(chunks)
    return asyncio.run(_run())


def tts_speak(
    text: str,
    voice: str,
    rate: str = "+0%",
    *,
    pitch_steps: int = 0,
    tone: str = "natural",
    volume: int = 0,
) -> tuple[bytes, str]:
    """Speak text using self-hosted server first, fall back to edge-tts.

    Returns (audio_bytes, content_type).
    """
    backend_voice, backend_rate, pitch, skip_kokoro = _resolve_tts_request(voice, rate)
    if pitch_steps and (skip_kokoro or not KOKORO_URL):
        pitch = f"{_clamp_int(pitch_steps, -12, 12, 0) * 6:+d}Hz"
    last_error = ""
    if KOKORO_URL and not skip_kokoro:
        try:
            audio, content_type = kokoro_tts_speak(text, backend_voice, backend_rate)
            return mix_audio(audio, content_type, pitch=pitch_steps, tone=tone, volume=volume)
        except Exception as exc:
            last_error = str(exc)
    if EDGE_TTS_AVAILABLE:
        audio = edge_tts_speak(text, backend_voice, backend_rate, pitch)
        return mix_audio(audio, "audio/mpeg", pitch=0, tone=tone, volume=volume)
    detail = f" Last local TTS error: {last_error}" if last_error else ""
    raise RuntimeError("No working TTS backend available. Install edge-tts, configure Piper/Kokoro, or download Piper voices." + detail)


def clean_text_for_tts(text: str) -> str:
    """Convert markdown-heavy assistant output into speech-friendly text."""
    text = re.sub(r"```[\s\S]*?```", " Code block omitted. ", text or "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        line = re.sub(r"[✅❌⚠️✔✖✗✓]", "", line)
        if not line:
            lines.append("")
            continue
        if re.fullmatch(r"\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?", line):
            continue
        if "|" in line and line.count("|") >= 2:
            cells = [cell.strip() for cell in line.strip("|").split("|") if cell.strip()]
            if cells:
                lines.append("; ".join(cells) + ".")
            continue
        line = re.sub(r"^#{1,6}\s*", "", line)
        line = re.sub(r"^\s*[-*+]\s+", "", line)
        line = re.sub(r"^\s*\d+[.)]\s+", "", line)
        line = re.sub(r"^\s*>\s?", "", line)
        line = re.sub(r"[*_~]{1,3}", "", line)
        line = re.sub(r"\s{2,}", " ", line)
        if line:
            lines.append(line)
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned).strip()
    return cleaned

import base64

ROOT = Path(__file__).resolve().parents[1]
VENV_PYTHON = ROOT / ".venv" / "bin" / "python"
VENV_ROOT = ROOT / ".venv"
if (
    VENV_PYTHON.exists()
    and Path(sys.prefix).resolve() != VENV_ROOT.resolve()
    and os.environ.get("OLLAMA_WEB_NO_VENV", "").lower() not in {"1", "true", "yes"}
):
    os.execve(str(VENV_PYTHON), [str(VENV_PYTHON), *sys.argv], os.environ)

if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from ollama_tools.langchain_orchestrator import (  # noqa: E402
    LangChainUnavailableError,
    invoke_langchain_agent_with_trace,
    is_langchain_available,
    stream_langchain_agent_events,
)
from ollama_tools.config import (  # noqa: E402
    get_ollama_base_url,
    get_web_host, get_web_port, get_piper_url,
    get_web_cert_file, get_web_key_file,
    get_gpu_urls, get_gpu_labels,
    get_stt_enabled,
    get_whisper_model, get_whisper_device, get_whisper_device_index,
    get_auth_enabled, get_auth_skip_allowed, get_cookie_same_site, get_cookie_secure,
    get_dangerous_mode,
    get_data_dir, get_setup_completed, is_deploy_mode,
    get_ollama_base_url, get_piper_url, get_searxng_url, get_brave_api_key,
    get_ssh_config_path, get_tool_enabled, get_work_dir, get_allow_work_dir_writes,
    get_thunderbird_enabled, get_thunderbird_token, get_thunderbird_max_messages,
    get_thunderbird_max_chars_per_message,
    reload_config,
    set_hook_override, set_skill_override, set_tool_override,
)
from ollama_tools.auth import (  # noqa: E402
    authenticate, create_admin_user, create_session,
    delete_all_users, get_first_username, has_users,
    rotate_session_secret, verify_session,
)
from ollama_tools.first_run import WARNING_TEXT, normalize_ollama_url, reset_runtime_state, write_first_run_config  # noqa: E402
from ollama_tools.agent_profiles import AGENT_PROFILES, agent_profile_prompt, normalize_agent_profile  # noqa: E402
from ollama_tools.context_loader import context_prompt, select_context  # noqa: E402
from ollama_tools.tool_registry import all_metadata  # noqa: E402
from ollama_tools.ollama_api import ChatStats, OllamaError, ThinkingChunk, list_models, stream_chat  # noqa: E402
from scripts.ollama_tui import (  # noqa: E402
    CHAT_MODES,
    CHAT_PERSONALITIES,
    SYSTEM_PROMPT,
    whitespace_columns_to_markdown,
    run_tool_command,
)


def _auto_tls_enabled() -> bool:
    return os.environ.get("OLLAMA_WEB_AUTO_TLS", "").strip().lower() in {"1", "true", "yes", "on"}


def _local_tls_names() -> tuple[list[str], list[str]]:
    dns_names = {"localhost", socket.gethostname()}
    ip_addrs = {"127.0.0.1"}
    try:
        dns_names.add(socket.getfqdn())
    except Exception:
        pass
    for item in os.environ.get("OLLAMA_WEB_TLS_HOSTS", "").split(","):
        item = item.strip()
        if item:
            try:
                ip_addrs.add(str(ipaddress.ip_address(item)))
            except ValueError:
                dns_names.add(item)

    try:
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as sock:
            sock.connect(("8.8.8.8", 80))
            ip_addrs.add(sock.getsockname()[0])
    except Exception:
        pass
    try:
        for info in socket.getaddrinfo(socket.gethostname(), None, family=socket.AF_INET):
            ip_addrs.add(info[4][0])
    except Exception:
        pass
    return sorted(name for name in dns_names if name), sorted(ip_addrs)


def _ensure_auto_tls_cert(cert_file: str, key_file: str) -> None:
    cert_path = Path(cert_file)
    key_path = Path(key_file)
    dns_names, ip_addrs = _local_tls_names()
    san_parts = [f"DNS:{name}" for name in dns_names]
    san_parts.extend(f"IP:{ip}" for ip in ip_addrs)
    san_text = ",".join(san_parts)
    marker_path = cert_path.with_suffix(cert_path.suffix + ".hosts")
    if cert_path.exists() and key_path.exists() and marker_path.exists():
        if marker_path.read_text(encoding="utf-8", errors="ignore").strip() == san_text:
            return
    openssl = shutil.which("openssl")
    if not openssl:
        raise RuntimeError("OLLAMA_WEB_AUTO_TLS is enabled, but openssl is not installed in the container.")
    cert_path.parent.mkdir(parents=True, exist_ok=True)
    key_path.parent.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            openssl,
            "req",
            "-x509",
            "-newkey",
            "rsa:2048",
            "-nodes",
            "-sha256",
            "-days",
            "825",
            "-keyout",
            str(key_path),
            "-out",
            str(cert_path),
            "-subj",
            "/CN=Ollama Command Deck",
            "-addext",
            "subjectAltName=" + san_text,
        ],
        check=True,
        text=True,
        capture_output=True,
    )
    try:
        key_path.chmod(0o600)
        cert_path.chmod(0o644)
    except OSError:
        pass
    marker_path.write_text(san_text + "\n", encoding="utf-8")
    print(f"[web] Generated self-signed HTTPS cert at {cert_path}", file=sys.stderr)


HOST = get_web_host()
PORT = get_web_port()
CERT_FILE = get_web_cert_file()
KEY_FILE = get_web_key_file()
if _auto_tls_enabled() and CERT_FILE and KEY_FILE:
    _ensure_auto_tls_cert(CERT_FILE, KEY_FILE)
TLS_ENABLED = bool(CERT_FILE and KEY_FILE)

# ── Vision / multi-modal helpers ──────────────────────────────────────────────
_VISION_PATTERNS = ("llava", "bakllava", "moondream", "minicpm-v", "qwen2-vl",
                    "gemma3", "mistral-small3", "phi3-vision", "phi4-vision",
                    "llama3.2-vision", "llama-3.2-vision")

_CODE_PATTERNS = ("codellama", "deepseek-coder", "qwen2.5-coder", "starcoder",
                  "codegemma", "magicoder", "phind-codellama", "codestral",
                  "granite-code", "coder")

_THINK_PATTERNS = ("qwen3", "deepseek-r1", "phi4-reasoning", "qwq", "marco-o1",
                   "skywork-o1", "r1")

_DOC_PATTERNS = ("llama", "mistral", "gemma", "qwen", "phi", "dolphin",
                 "neural-chat", "openchat", "orca", "vicuna", "wizard")


def model_capabilities(name: str) -> list[str]:
    """Return a list of capability tags for a model name."""
    low = name.lower()
    caps: list[str] = ["chat"]
    if _is_vision_model(name):
        caps.append("vision")
    if any(p in low for p in _CODE_PATTERNS):
        caps.append("code")
    if any(p in low for p in _THINK_PATTERNS):
        caps.append("think")
    # Agent mode works with any general text model (not pure vision-only)
    if "chat" in caps:
        caps.append("agent")
    # Document reading works on any model that can handle long context
    caps.append("document")
    return caps


def _is_vision_model(name: str) -> bool:
    low = name.lower()
    return any(p in low for p in _VISION_PATTERNS)


def pick_vision_model(base_url: str | None = None) -> str | None:
    """Return the name of the first available vision model, or None."""
    try:
        for model in list_models(base_url=base_url):
            if _is_vision_model(model.name):
                return model.name
    except Exception:
        pass
    return None


def pdf_to_images_b64(data_b64: str, max_pages: int = 4) -> list[str]:
    """Render PDF pages to base64 PNG images using pymupdf. Returns [] on failure."""
    try:
        import fitz  # pymupdf
        raw = base64.b64decode(data_b64)
        doc = fitz.open(stream=raw, filetype="pdf")
        out: list[str] = []
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2× zoom for legibility
            out.append(base64.b64encode(pix.tobytes("png")).decode())
        return out
    except Exception:
        return []


def extract_document_text(data_b64: str, filename: str) -> str:
    """Convert an uploaded document to plain text for context injection."""
    import io
    raw = base64.b64decode(data_b64)
    fname_lower = filename.lower()

    if fname_lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n\n".join(pages).strip() or "[PDF contained no extractable text]"
        except ImportError:
            return "[pypdf not installed — run: pip install pypdf]"
        except Exception as exc:
            return f"[Could not extract PDF text: {exc}]"

    if fname_lower.endswith((".docx",)):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            parts: list[str] = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)
            # Also pull text from tables
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            return "\n".join(parts).strip() or "[Word document contained no extractable text]"
        except ImportError:
            return "[python-docx not installed — run: pip install python-docx]"
        except Exception as exc:
            return f"[Could not extract Word document text: {exc}]"

    if fname_lower.endswith(".doc"):
        return "[Legacy .doc format not supported — please save as .docx and re-upload]"

    if fname_lower.endswith(".csv"):
        # Decode then pretty-format as a text table
        for enc in ("utf-8", "latin-1"):
            try:
                text = raw.decode(enc).strip()
                lines = text.splitlines()
                # Keep as-is but cap at 500 rows to avoid overflowing context
                if len(lines) > 500:
                    text = "\n".join(lines[:500]) + f"\n... ({len(lines)-500} more rows truncated)"
                return text
            except UnicodeDecodeError:
                continue
        return "[Could not decode CSV file]"

    # Plain text fallback (txt, md, json, etc.)
    for enc in ("utf-8", "latin-1"):
        try:
            return raw.decode(enc).strip()
        except UnicodeDecodeError:
            continue
    return "[Could not decode file as text]"
PORT_SEARCH_LIMIT = int(os.environ.get("OLLAMA_WEB_PORT_SEARCH_LIMIT", "20"))
# Resolve KOKORO_URL now that config is loaded (env var set early overrides config)
KOKORO_URL = (os.environ.get("KOKORO_URL", "") or get_piper_url()).rstrip("/")

def _gpu_base_url(gpu_id: str) -> str | None:
    """Return the Ollama URL for a GPU ID, or None to use the default."""
    # Do not cache this at module import. First-run setup can rewrite
    # config.toml and call reload_config() without restarting the process.
    return get_gpu_urls().get(str(gpu_id)) or None


def _gpu_rows() -> list[dict[str, str]]:
    """Return current GPU/Ollama endpoint rows for the browser selector."""
    gpu_urls = get_gpu_urls()
    gpu_labels = get_gpu_labels()
    rows = []
    for gid, url in sorted(gpu_urls.items()):
        from urllib.parse import urlparse as _uparse
        parsed = _uparse(url)
        default_label = f"GPU {gid} — port {parsed.port or 11434}"
        label = gpu_labels.get(gid, default_label)
        rows.append({"id": gid, "url": url, "label": label})
    return rows

# ── Whisper STT ───────────────────────────────────────────────────────────────
_whisper_model_inst = None
_whisper_lock = threading.Lock()
WHISPER_AVAILABLE: bool | None = None  # None = not yet checked

def _load_whisper():
    """Lazy-load the Whisper model. Returns the model or None if unavailable."""
    global _whisper_model_inst, WHISPER_AVAILABLE
    if _whisper_model_inst is not None:
        return _whisper_model_inst
    with _whisper_lock:
        if _whisper_model_inst is not None:
            return _whisper_model_inst
        try:
            from faster_whisper import WhisperModel
            model_size = os.environ.get("WHISPER_MODEL", get_whisper_model())
            device = os.environ.get("WHISPER_DEVICE", get_whisper_device())
            device_index = int(os.environ.get("WHISPER_DEVICE_INDEX", str(get_whisper_device_index())))
            print(f"[stt] Loading Whisper {model_size} on {device}:{device_index}…", file=sys.stderr)
            _whisper_model_inst = WhisperModel(
                model_size, device=device, device_index=device_index, compute_type="int8"
            )
            WHISPER_AVAILABLE = True
            print("[stt] Whisper ready.", file=sys.stderr)
            return _whisper_model_inst
        except Exception as exc:
            WHISPER_AVAILABLE = False
            print(f"[stt] Whisper not available: {exc}", file=sys.stderr)
            return None

def is_whisper_available() -> bool:
    """Quick check without loading the model."""
    global WHISPER_AVAILABLE
    if not get_stt_enabled():
        WHISPER_AVAILABLE = False
        return False
    if WHISPER_AVAILABLE is not None:
        return WHISPER_AVAILABLE
    try:
        import faster_whisper  # noqa: F401
        WHISPER_AVAILABLE = True
    except ImportError:
        WHISPER_AVAILABLE = False
    return WHISPER_AVAILABLE

def transcribe_audio(audio_bytes: bytes, content_type: str = "audio/webm") -> str:
    """Transcribe audio bytes via Whisper. Returns transcribed text."""
    ext = ".webm"
    if "ogg" in content_type:
        ext = ".ogg"
    elif "wav" in content_type:
        ext = ".wav"
    elif "mp4" in content_type or "m4a" in content_type:
        ext = ".mp4"
    model = _load_whisper()
    if model is None:
        raise RuntimeError("Whisper model not loaded. Install faster-whisper and check config.")
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
        f.write(audio_bytes)
        tmp_path = f.name
    try:
        segments, _info = model.transcribe(tmp_path, beam_size=5)
        return " ".join(seg.text for seg in segments).strip()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ── Sessions (chat history persistence) ───────────────────────────────────────
_SESSIONS_DIR = get_data_dir() / "web_sessions"
_SESSION_ID_RE = re.compile(r"^[A-Za-z0-9_-]{1,64}$")


def _ensure_sessions_dir() -> Path:
    _SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
    try:
        os.chmod(_SESSIONS_DIR, 0o700)
    except OSError:
        pass
    return _SESSIONS_DIR


def _session_path(session_id: str) -> Path | None:
    if not _SESSION_ID_RE.match(session_id):
        return None
    return _SESSIONS_DIR / f"{session_id}.json"


def _new_session_id() -> str:
    import uuid
    return uuid.uuid4().hex[:16]


def _session_summary(path: Path) -> dict[str, Any] | None:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    return {
        "id": data.get("id") or path.stem,
        "title": data.get("title") or "(untitled)",
        "model": data.get("model") or "",
        "created_at": data.get("created_at") or "",
        "updated_at": data.get("updated_at") or data.get("created_at") or "",
        "message_count": len(data.get("messages") or []),
    }


def list_sessions() -> list[dict[str, Any]]:
    _ensure_sessions_dir()
    items: list[dict[str, Any]] = []
    for p in _SESSIONS_DIR.glob("*.json"):
        s = _session_summary(p)
        if s:
            items.append(s)
    items.sort(key=lambda s: s.get("updated_at") or "", reverse=True)
    return items


def load_session(session_id: str) -> dict[str, Any] | None:
    p = _session_path(session_id)
    if not p or not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None


def save_session(payload: dict[str, Any]) -> dict[str, Any]:
    _ensure_sessions_dir()
    sid = str(payload.get("id") or "").strip() or _new_session_id()
    p = _session_path(sid)
    if not p:
        raise ValueError("invalid session id")
    now = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
    existing = load_session(sid) if p.exists() else {}
    messages = payload.get("messages") or existing.get("messages") or []
    title = str(payload.get("title") or existing.get("title") or "").strip()
    if not title:
        # Derive title from the first user message
        first_user = next((m.get("content", "") for m in messages if m.get("role") == "user"), "")
        title = (first_user.strip().splitlines()[0] if first_user else "(empty)")[:80] or "(empty)"
    record = {
        "id": sid,
        "title": title,
        "model": str(payload.get("model") or existing.get("model") or ""),
        "created_at": existing.get("created_at") or now,
        "updated_at": now,
        "messages": messages,
        "agentMode": bool(payload.get("agentMode", existing.get("agentMode", True))),
        "chatMode": str(payload.get("chatMode") or existing.get("chatMode") or "default"),
        "personality": str(payload.get("personality") or existing.get("personality") or "default"),
        "agentProfile": normalize_agent_profile(str(payload.get("agentProfile") or existing.get("agentProfile") or "general")),
        "autoProfile": bool(payload.get("autoProfile", existing.get("autoProfile", False))),
        "assistantName": _assistant_name(str(payload.get("assistantName") or existing.get("assistantName") or "Lilith")),
    }
    p.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
    try:
        os.chmod(p, 0o600)
    except OSError:
        pass
    return _session_summary(p) or {"id": sid}


def delete_session(session_id: str) -> bool:
    p = _session_path(session_id)
    if not p or not p.exists():
        return False
    try:
        p.unlink()
        return True
    except OSError:
        return False


# ── Canvas files ──────────────────────────────────────────────────────────────
_CANVAS_DIR = get_data_dir() / "canvas_files"
_CANVAS_NAME_RE = re.compile(r"^[A-Za-z0-9 _.\-]{1,80}$")


def _ensure_canvas_dir() -> Path:
    _CANVAS_DIR.mkdir(parents=True, exist_ok=True)
    try:
        os.chmod(_CANVAS_DIR, 0o700)
    except OSError:
        pass
    return _CANVAS_DIR


def _canvas_path(name: str) -> Path | None:
    name = name.strip()
    if not _CANVAS_NAME_RE.match(name):
        return None
    if not name.lower().endswith(".md"):
        name = name + ".md"
    p = _CANVAS_DIR / name
    # Reject path traversal
    try:
        p.resolve().relative_to(_CANVAS_DIR.resolve())
    except ValueError:
        return None
    return p


def list_canvas_files() -> list[dict[str, Any]]:
    _ensure_canvas_dir()
    items: list[dict[str, Any]] = []
    for p in _CANVAS_DIR.glob("*.md"):
        try:
            st = p.stat()
        except OSError:
            continue
        items.append({
            "name": p.stem,
            "size": st.st_size,
            "updated_at": datetime.utcfromtimestamp(st.st_mtime).strftime("%Y-%m-%dT%H:%M:%SZ"),
        })
    items.sort(key=lambda i: i["updated_at"], reverse=True)
    return items


def load_canvas_file(name: str) -> str | None:
    p = _canvas_path(name)
    if not p or not p.exists():
        return None
    try:
        return p.read_text(encoding="utf-8")
    except OSError:
        return None


def save_canvas_file(name: str, content: str) -> dict[str, Any] | None:
    _ensure_canvas_dir()
    p = _canvas_path(name)
    if not p:
        return None
    p.write_text(content, encoding="utf-8")
    try:
        os.chmod(p, 0o600)
    except OSError:
        pass
    st = p.stat()
    return {
        "name": p.stem,
        "size": st.st_size,
        "updated_at": datetime.utcfromtimestamp(st.st_mtime).strftime("%Y-%m-%dT%H:%M:%SZ"),
    }


def delete_canvas_file(name: str) -> bool:
    p = _canvas_path(name)
    if not p or not p.exists():
        return False
    try:
        p.unlink()
        return True
    except OSError:
        return False


def fetch_model_context_info(model: str, base_url: str | None) -> dict[str, int | None]:
    """Query Ollama /api/show for advertised and runtime context windows."""
    result: dict[str, int | None] = {"advertised_context_length": None, "runtime_context_length": None}
    if not model:
        return result
    try:
        url = f"{base_url or get_ollama_base_url()}/api/show"
        request = urllib.request.Request(
            url,
            data=json.dumps({"model": model}).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        with urllib.request.urlopen(request, timeout=10) as response:
            data = json.loads(response.read().decode("utf-8"))
    except (urllib.error.URLError, ValueError, json.JSONDecodeError):
        return result
    info = data.get("model_info") or {}
    for key, val in info.items():
        if isinstance(val, int) and key.endswith(".context_length"):
            result["advertised_context_length"] = val
            break
    pl = data.get("parameters") or ""
    if isinstance(pl, str):
        m = re.search(r"num_ctx\s+(\d+)", pl)
        if m:
            result["runtime_context_length"] = int(m.group(1))
    return result


def fetch_model_context_length(model: str, base_url: str | None) -> int | None:
    """Backward-compatible helper returning the advertised context length."""
    info = fetch_model_context_info(model, base_url)
    return info.get("advertised_context_length") or info.get("runtime_context_length")


def _assistant_name(value: str | None) -> str:
    name = re.sub(r"[^A-Za-z0-9 _.-]", "", str(value or "")).strip()
    return (name[:40] or "Lilith")


def build_system_prompt(
    chat_mode: str,
    personality: str,
    agent_profile: str = "general",
    user_text: str = "",
    assistant_name: str = "Lilith",
) -> str:
    mode = chat_mode if chat_mode in CHAT_MODES else "default"
    persona = personality if personality in CHAT_PERSONALITIES else "default"
    name = _assistant_name(assistant_name)
    prompt = re.sub(r"Your name is Lilith\.", f"Your name is {name}.", SYSTEM_PROMPT, count=1)
    if name != "Lilith":
        prompt += f"\n\nThe user has renamed the assistant to {name}. Use that name for self-reference."
    prompt += "\n\n" + agent_profile_prompt(agent_profile)
    if normalize_agent_profile(agent_profile) == "builder":
        prompt += (
            f"\n\nBuilder workspace: the configured work directory is `{get_work_dir()}`. "
            "Inspect and edit only inside this directory. If the user asks for a path outside this directory, "
            "stop and report the configuration mismatch; do not fall back to inspecting `/app`, sibling repos, "
            "or parent workspace directories. The Web runtime creates and appends `BUILDER_RUN.md` automatically; "
            "do not create it yourself, and do not report it missing just because no shell write command created it."
        )
    prompt += context_prompt(ROOT, agent_profile, user_text)
    if CHAT_MODES[mode]:
        prompt += "\n\n" + CHAT_MODES[mode]
    if CHAT_PERSONALITIES[persona]:
        prompt += "\n\n" + CHAT_PERSONALITIES[persona]
    return prompt


def stats_to_dict(stats: ChatStats | None) -> dict[str, Any] | None:
    if stats is None:
        return None
    return {
        "prompt_tokens": stats.prompt_tokens,
        "response_tokens": stats.response_tokens,
        "total_duration_ms": stats.total_duration_ms,
        "text": str(stats),
    }


def _markdown_fence(text: str, lang: str = "") -> str:
    body = str(text or "").replace("```", "'''").rstrip()
    return f"```{lang}\n{body}\n```"


def _builder_log_append(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(text)


def _builder_log_start(model: str, prompt: str, paths: list[str]) -> tuple[Path | None, str | None]:
    if not get_allow_work_dir_writes():
        return None, "Builder run log disabled because work-directory writes are disabled."
    work_dir = get_work_dir()
    path = work_dir / "BUILDER_RUN.md"
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    try:
        if not path.exists():
            _builder_log_append(path, "# Builder Run Log\n\n")
        _builder_log_append(
            path,
            (
                f"## {stamp}\n\n"
                f"- Model: `{model}`\n"
                f"- Work directory: `{work_dir}`\n"
                f"- Context: {', '.join(paths) if paths else 'none'}\n"
                "- Status: started\n\n"
                "### Prompt\n\n"
                f"{_markdown_fence(prompt)}\n\n"
            ),
        )
    except OSError as exc:
        return None, f"Builder run log unavailable at {path}: {exc}"
    return path, None


def _builder_log_event(path: Path | None, heading: str, text: str, lang: str = "") -> None:
    if path is None:
        return
    try:
        _builder_log_append(path, f"### {heading}\n\n{_markdown_fence(text, lang)}\n\n")
    except OSError:
        return


def _builder_log_final_note(answer_text: str, path: Path | None) -> str:
    if path is None:
        return answer_text
    note = (
        f"\n\nRuntime log: `{path}` was created/opened by the Web runtime. "
        "Ignore any model statement above that says `BUILDER_RUN.md` was missing or not shown in shell evidence."
    )
    if str(path) in answer_text and "Runtime log:" in answer_text:
        return answer_text
    return answer_text.rstrip() + note


def collect_chat(model: str, messages: list[dict[str, str]], verbose: bool, keep_alive: str | None = None) -> tuple[str, ChatStats | None]:
    parts: list[str] = []
    stats: ChatStats | None = None
    for chunk in stream_chat(model, messages, collect_stats=verbose, keep_alive=keep_alive):
        if isinstance(chunk, ChatStats):
            stats = chunk
        elif isinstance(chunk, ThinkingChunk):
            continue  # skip reasoning trace in non-streaming collection
        else:
            parts.append(str(chunk))
    return "".join(parts), stats


def _default_thunderbird_model() -> str:
    preferred = {"qwen3.5:latest", "qwen3:latest", "qwen3.5", "qwen3"}
    try:
        models = list(list_models(base_url=get_ollama_base_url()))
    except Exception:
        return "qwen3.5:latest"
    for item in models:
        if item.name in preferred:
            return item.name
    return models[0].name if models else "qwen3.5:latest"


def _message_text_from_payload(item: dict[str, Any], max_chars: int) -> str:
    body = str(item.get("body") or item.get("text") or item.get("excerpt") or "")
    if len(body) > max_chars:
        body = body[:max_chars].rstrip() + "\n... [message truncated]"
    subject = str(item.get("subject") or "(no subject)")
    author = str(item.get("author") or item.get("from") or "")
    date = str(item.get("date") or "")
    folder = str(item.get("folder") or "")
    return (
        f"Subject: {subject}\n"
        f"From: {author}\n"
        f"Date: {date}\n"
        f"Folder: {folder}\n"
        f"Body:\n{body}"
    ).strip()


def handle_thunderbird_analyze(payload: dict[str, Any]) -> dict[str, Any]:
    if not get_thunderbird_enabled():
        return {"ok": False, "error": "Thunderbird bridge is disabled."}
    expected = get_thunderbird_token()
    provided = str(payload.get("token") or "")
    if not expected or not provided or not secrets.compare_digest(expected, provided):
        return {"ok": False, "error": "Invalid Thunderbird bridge token."}
    raw_messages = payload.get("messages")
    if not isinstance(raw_messages, list) or not raw_messages:
        return {"ok": False, "error": "messages must be a non-empty list."}

    max_messages = get_thunderbird_max_messages()
    max_chars = get_thunderbird_max_chars_per_message()
    messages = [item for item in raw_messages[:max_messages] if isinstance(item, dict)]
    if not messages:
        return {"ok": False, "error": "No valid message objects supplied."}

    question = str(payload.get("question") or payload.get("query") or "Find the relevant email details.").strip()
    model = str(payload.get("model") or "").strip() or _default_thunderbird_model()
    email_context = "\n\n---\n\n".join(_message_text_from_payload(item, max_chars) for item in messages)
    chat_messages = [
        {
            "role": "system",
            "content": (
                "You are a read-only email search assistant for Thunderbird. Use only the email snippets "
                "provided in this request. Do not claim you can send, delete, move, mark, or modify email. "
                "If the snippets do not contain the answer, say what is missing and suggest a narrower search."
            ),
        },
        {
            "role": "user",
            "content": f"Question:\n{question}\n\nEmail snippets:\n{email_context}",
        },
    ]
    answer, stats = collect_chat(model, chat_messages, verbose=True, keep_alive="10m")
    return {
        "ok": True,
        "model": model,
        "text": whitespace_columns_to_markdown(answer),
        "stats": stats_to_dict(stats),
        "messages_used": len(messages),
    }


def chat_messages_from_payload(payload: dict[str, Any], text: str) -> list[dict[str, str]]:
    chat_mode = str(payload.get("chatMode") or "default")
    personality = str(payload.get("personality") or "default")
    agent_profile = normalize_agent_profile(str(payload.get("agentProfile") or "general"))
    assistant_name = _assistant_name(str(payload.get("assistantName") or "Lilith"))
    canvas_requested = bool(payload.get("canvasRequested"))
    history = payload.get("messages") if isinstance(payload.get("messages"), list) else []
    messages: list[dict[str, str]] = [{"role": "system", "content": build_system_prompt(chat_mode, personality, agent_profile, text, assistant_name)}]
    if canvas_requested:
        canvas_instruction = (
            "The user explicitly requested a canvas document. Return the complete content as a "
            "standalone Markdown document suitable for direct insertion into the canvas. "
            "Do not include chat transcript labels, tool logs, or meta commentary such as "
            "'here is the content for the canvas'."
        )
        if agent_profile == "deep_research":
            canvas_instruction += (
                " For deep research, the document must be a detailed report with a title, executive summary, "
                "key findings, detailed analysis, risks/gaps, recommendations, and sources/notes. Cite source "
                "URLs with Markdown links when tool evidence provides them. Do not invent citations."
            )
        messages.append({"role": "system", "content": canvas_instruction})
    for item in history[-30:]:
        if isinstance(item, dict) and item.get("role") in {"user", "assistant", "system"}:
            messages.append({"role": str(item["role"]), "content": str(item.get("content", ""))})
    messages.append({"role": "user", "content": text})
    return messages


def stream_chat_events(payload: dict[str, Any]):
    model = str(payload.get("model") or "")
    text = str(payload.get("text") or "").strip()
    verbose = bool(payload.get("verbose"))
    agent_mode = bool(payload.get("agentMode"))
    canvas_requested = bool(payload.get("canvasRequested"))
    ollama_base_url = _gpu_base_url(str(payload.get("gpu") or "0"))
    keep_alive = "30m" if bool(payload.get("keepAlive")) else None
    attachment: dict[str, Any] | None = payload.get("attachment") if isinstance(payload.get("attachment"), dict) else None
    requested_agent_profile = normalize_agent_profile(str(payload.get("agentProfile") or "general"))
    auto_profile = bool(payload.get("autoProfile"))
    assistant_name = _assistant_name(str(payload.get("assistantName") or "Lilith"))

    # Resolve attachment — images auto-switch to a vision model
    images: list[str] = []
    if attachment:
        att_type = str(attachment.get("type", ""))
        if att_type == "image":
            vision_model = pick_vision_model(base_url=ollama_base_url)
            if not vision_model:
                yield {"type": "error", "error": "No vision model found. Pull one with: ollama pull llava"}
                return
            if vision_model != model:
                yield {"type": "cmd", "command": f"Auto-switching to vision model: {vision_model}"}
                model = vision_model
            images = [str(attachment.get("data", ""))]
            if not text:
                text = "Describe this image."
            # Vision models don't support tool-calling — force chat mode
            if agent_mode:
                agent_mode = False
                yield {"type": "cmd", "command": "Switched to Chat mode (vision model active)"}
        elif att_type == "document":
            att_data = str(attachment.get("data", ""))
            filename = str(attachment.get("filename", "file"))
            doc_text = extract_document_text(att_data, filename)
            scanned = len(doc_text.strip()) < 40 and filename.lower().endswith(".pdf")

            if scanned:
                page_images = pdf_to_images_b64(att_data)
                if page_images:
                    vision_model = pick_vision_model(base_url=ollama_base_url)
                    if not vision_model:
                        yield {"type": "error", "error": "Scanned PDF detected but no vision model available. Run: ollama pull llava"}
                        return
                    if vision_model != model:
                        yield {"type": "cmd", "command": f"Scanned PDF — switching to vision model: {vision_model}"}
                        model = vision_model
                    images = page_images
                    if not text:
                        text = (
                            "Please transcribe ALL text visible in this document verbatim, "
                            "page by page, preserving the layout as much as possible. "
                            "Include every field, label, value, and signature line you can read."
                        )
                    agent_mode = False  # vision models don't support tool-calling
                    yield {"type": "cmd", "command": f"Rendering {len(page_images)} page(s) → {model}"}
                    yield {"type": "cmd", "command": "📋 Transcription will be sent to canvas automatically"}
                else:
                    yield {"type": "cmd", "command": "Scanned PDF detected but page rendering failed"}
            else:
                preamble = (
                    f"[DOCUMENT UPLOADED BY USER: \"{filename}\"]\n"
                    f"[Full extracted text is provided below — do NOT use tools to search for this file.]\n\n"
                    f"{doc_text}"
                )
                text = f"{preamble}\n\n{text}" if text else preamble
                if agent_mode and _is_vision_model(model):
                    agent_mode = False
                    yield {"type": "cmd", "command": f"Reading document with {model} (Chat mode — vision model)"}
                else:
                    mode_label = "Agent mode" if agent_mode else "Chat mode"
                    yield {"type": "cmd", "command": f"Reading document with {model} ({mode_label})"}

    if not text and not images:
        yield {"type": "error", "error": "Message is empty."}
        return

    control_prefixes = ("/model", "/chat", "/agent on", "/agent off", "/verbose on", "/verbose off", "/chat_mode ", "/chat_personality ", "/agent_profile ", "/quit")
    if not attachment and (text == "/clear" or text == "/help" or any(text == item or text.startswith(item) for item in control_prefixes)):
        result = handle_chat(payload)
        yield {"type": "final", **result}
        return

    if not model:
        yield {"type": "error", "error": "No model selected."}
        return

    if verbose:
        yield {"type": "cmd", "command": f"Model: {model}"}

    canvas_content = str(payload.get("canvasContent") or "").strip()

    # Canvas operation — just answer from context, no tools needed
    _canvas_op = bool(canvas_content) and bool(
        re.search(r'\b(canvas|put|send|copy|move|append|add).{0,30}\bcanvas\b'
                  r'|\bcanvas\b.{0,30}(text|content|document|pdf|result)'
                  r'|\b(to|into)\s+canvas\b', text, re.I)
    )
    if _canvas_op and agent_mode:
        agent_mode = False

    context_selection = select_context(ROOT, requested_agent_profile, text, auto_route=auto_profile)
    if context_selection.routed:
        payload = dict(payload)
        payload["agentProfile"] = context_selection.active_profile
        yield {
            "type": "context",
            "profile": context_selection.active_profile,
            "requested_profile": context_selection.requested_profile,
            "routed": True,
            "reason": context_selection.route_reason,
            "paths": context_selection.paths,
            "text": (
                f"Profile routed: {context_selection.requested_profile} -> "
                f"{context_selection.active_profile} ({context_selection.route_reason})\n"
                "Context: " + (", ".join(context_selection.paths) or "none")
            ),
        }
    else:
        yield {
            "type": "context",
            "profile": context_selection.active_profile,
            "requested_profile": context_selection.requested_profile,
            "routed": False,
            "reason": "",
            "paths": context_selection.paths,
            "text": (
                f"Profile: {context_selection.active_profile}\n"
                "Context: " + (", ".join(context_selection.paths) or "none")
            ),
        }

    messages = chat_messages_from_payload(payload, text)
    if canvas_content:
        if canvas_requested:
            canvas_note = (
                "[The user's document canvas contains the current full document below. "
                "Use it as the source document, preserve useful existing content, integrate any new user edits, "
                "and return the complete updated Markdown document for the canvas. Make the document polished for "
                "a human reader. Prefer markdown headings, tables, numbered workflows, and bullets. Avoid fragile "
                "ASCII box diagrams unless they are simple and perfectly aligned; use markdown tables or mermaid "
                "fenced diagrams for complex architecture/process flows. Do NOT use tools.]\n\n"
            )
        else:
            canvas_note = (
                "[The user's document canvas contains the following text. "
                "Answer questions about it directly — do NOT use any tools.]\n\n"
            )
        canvas_ctx = {"role": "system", "content": canvas_note + canvas_content}
        messages = messages[:-1] + [canvas_ctx, messages[-1]]
    builder_log_path: Path | None = None
    try:
        if not images and (text.startswith("/agent ") or (agent_mode and not text.startswith("/"))):
            agent_text = text[len("/agent "):].strip() if text.startswith("/agent ") else text
            agent_messages = messages[:-1] + [{"role": "user", "content": agent_text}]
            final_event: dict[str, Any] = {}
            if context_selection.active_profile == "builder":
                tool_rounds = 20
            elif context_selection.active_profile == "deep_research":
                tool_rounds = 12
            else:
                tool_rounds = 4
            if context_selection.active_profile == "builder":
                builder_log_path, builder_log_error = _builder_log_start(model, agent_text, context_selection.paths)
                if builder_log_path:
                    yield {"type": "cmd", "command": f"Builder run log: {builder_log_path}"}
                    agent_messages.insert(
                        -1,
                        {
                            "role": "system",
                            "content": (
                                f"The Web runtime has already created or opened the Builder run log at {builder_log_path}. "
                                "Treat this as runtime evidence. Do not create that file yourself and do not mark it "
                                "missing in the final answer unless the runtime reported an error."
                            ),
                        },
                    )
                elif builder_log_error:
                    yield {"type": "cmd", "command": builder_log_error}
                    agent_messages.insert(
                        -1,
                        {"role": "system", "content": f"Builder run log error: {builder_log_error}"},
                    )
            enforce_work_dir = context_selection.active_profile == "builder"
            for event in stream_langchain_agent_events(
                model,
                agent_messages,
                max_tool_rounds=tool_rounds,
                keep_alive=keep_alive,
                enforce_work_dir=enforce_work_dir,
            ):
                if event["type"] == "cmd":
                    command = str(event.get("command", ""))
                    _builder_log_event(builder_log_path, "Command", command, "bash")
                    yield {"type": "cmd", "command": command}
                elif event["type"] == "tool":
                    role = str(event.get("role") or "Tool result")
                    tool_text = str(event.get("text") or "")
                    _builder_log_event(builder_log_path, role, tool_text)
                    yield {"type": "tool", "role": role, "text": tool_text}
                elif event["type"] == "final":
                    final_event = event
            stats = final_event.get("stats")
            answer_text = whitespace_columns_to_markdown(str(final_event.get("text", "")))
            answer_text = _builder_log_final_note(answer_text, builder_log_path)
            _builder_log_event(builder_log_path, "Final response", answer_text)
            if builder_log_path:
                _builder_log_event(builder_log_path, "Status", "complete")
            yield {
                "type": "final",
                "text": answer_text,
                "role": assistant_name,
                "agentProfile": context_selection.active_profile,
                "stats": stats_to_dict(stats) if stats else None,
            }
            return

        tool_result = run_tool_command(text) if not images else None
        stream_messages = messages
        if tool_result:
            tool_role, tool_text, tool_cmd = tool_result
            if tool_cmd:
                yield {"type": "cmd", "command": tool_cmd}
            yield {"type": "tool", "role": tool_role, "text": tool_text, "command": tool_cmd}
            stream_messages = messages + [{"role": "system", "content": "Tool result:\n" + tool_text}]

        stats: ChatStats | None = None
        answer_text_parts: list[str] = []
        # Always collect stats (cheap) so the client can update its context-window bar.
        low_latency_chat = not agent_mode and str(payload.get("chatMode") or "") == "live"
        for chunk in stream_chat(
            model,
            stream_messages,
            collect_stats=True,
            images=images or None,
            base_url=ollama_base_url,
            keep_alive=keep_alive,
            think=False if low_latency_chat else None,
        ):
            if isinstance(chunk, ChatStats):
                stats = chunk
            elif isinstance(chunk, ThinkingChunk):
                if verbose and not low_latency_chat:
                    yield {"type": "thinking", "text": chunk.text}
            else:
                chunk_text = str(chunk)
                answer_text_parts.append(chunk_text)
                yield {"type": "chunk", "text": chunk_text}
        if stats:
            yield {"type": "stats", "stats": stats_to_dict(stats)}
        answer_text = "".join(answer_text_parts)
        if answer_text:
            yield {
                "type": "final",
                "text": whitespace_columns_to_markdown(answer_text),
                "role": assistant_name,
                "agentProfile": context_selection.active_profile,
            }
        yield {"type": "done"}
    except (OllamaError, LangChainUnavailableError, ValueError, RuntimeError) as exc:
        _builder_log_event(builder_log_path, "Error", str(exc))
        yield {"type": "error", "error": str(exc)}


def handle_chat(payload: dict[str, Any]) -> dict[str, Any]:
    model = str(payload.get("model") or "")
    text = str(payload.get("text") or "").strip()
    if not text:
        return {"ok": False, "error": "Message is empty."}

    agent_mode = bool(payload.get("agentMode"))
    verbose = bool(payload.get("verbose"))
    chat_mode = str(payload.get("chatMode") or "default")
    personality = str(payload.get("personality") or "default")
    agent_profile = normalize_agent_profile(str(payload.get("agentProfile") or "general"))
    assistant_name = _assistant_name(str(payload.get("assistantName") or "Lilith"))
    keep_alive = "30m" if bool(payload.get("keepAlive")) else None
    history = payload.get("messages") if isinstance(payload.get("messages"), list) else []

    if text == "/model":
        models = [item.name for item in list_models()]
        return {
            "ok": True,
            "role": "Tool",
            "text": "Available models:\n" + "\n".join(f"- {item}" for item in models) + "\n\nUse /model <name> to switch.",
            "models": models,
        }
    if text.startswith("/model "):
        requested = text[len("/model "):].strip()
        models = [item.name for item in list_models()]
        matches = [item for item in models if item == requested]
        if not matches:
            matches = [item for item in models if item.startswith(requested)]
        if len(matches) == 1:
            return {"ok": True, "model": matches[0], "role": "Tool", "text": f"Model switched to {matches[0]}."}
        if matches:
            return {
                "ok": True,
                "role": "Tool",
                "text": "Multiple models matched:\n" + "\n".join(f"- {item}" for item in matches),
                "models": models,
            }
        return {
            "ok": True,
            "role": "Tool",
            "text": f"No model matched '{requested}'. Use /model to list available models.",
            "models": models,
        }

    if text == "/chat":
        return {"ok": True, "mode": "chat", "chatMode": "live", "role": "Tool", "text": "Live chat mode enabled. Agent mode disabled."}
    if text == "/agent on":
        if not is_langchain_available():
            return {"ok": True, "mode": "chat", "role": "Tool", "text": "LangChain is not available. Install langchain and langchain-ollama in .venv first."}
        return {"ok": True, "mode": "agent", "role": "Tool", "text": "Agent mode enabled."}
    if text == "/agent off":
        return {"ok": True, "mode": "chat", "role": "Tool", "text": "Agent mode disabled."}
    if text == "/verbose on":
        return {"ok": True, "verbose": True, "role": "Tool", "text": "Verbose mode enabled."}
    if text == "/verbose off":
        return {"ok": True, "verbose": False, "role": "Tool", "text": "Verbose mode disabled."}
    if text.startswith("/chat_mode "):
        requested = text[len("/chat_mode "):].strip()
        if requested not in CHAT_MODES:
            return {"ok": True, "role": "Tool", "text": "usage: /chat_mode <mode>  (valid: " + ", ".join(CHAT_MODES) + ")"}
        return {"ok": True, "chatMode": requested, "role": "Tool", "text": f"Chat mode set to '{requested}'."}
    if text.startswith("/chat_personality "):
        requested = text[len("/chat_personality "):].strip()
        if requested not in CHAT_PERSONALITIES:
            return {"ok": True, "role": "Tool", "text": "usage: /chat_personality <name>  (valid: " + ", ".join(CHAT_PERSONALITIES) + ")"}
        return {"ok": True, "personality": requested, "role": "Tool", "text": f"Personality set to '{requested}'."}
    if text.startswith("/agent_profile "):
        requested = normalize_agent_profile(text[len("/agent_profile "):].strip())
        return {"ok": True, "agentProfile": requested, "role": "Tool", "text": f"Agent profile set to '{requested}'."}
    if text == "/quit":
        return {"ok": True, "role": "Tool", "text": "Close this browser tab to exit the web UI."}
    if not model:
        return {"ok": False, "error": "No model selected."}

    messages: list[dict[str, str]] = [{"role": "system", "content": build_system_prompt(chat_mode, personality, agent_profile, text, assistant_name)}]
    for item in history[-30:]:
        if isinstance(item, dict) and item.get("role") in {"user", "assistant", "system"}:
            messages.append({"role": str(item["role"]), "content": str(item.get("content", ""))})
    messages.append({"role": "user", "content": text})

    try:
        if text.startswith("/agent ") or (agent_mode and not text.startswith("/")):
            if not is_langchain_available():
                answer, stats = collect_chat(model, messages, verbose, keep_alive=keep_alive)
                notice = (
                    "LangChain is not available, so Agent mode was disabled and this was answered in normal chat.\n"
                    "Install langchain and langchain-ollama in the Python environment running this app."
                )
                return {
                    "ok": True,
                    "mode": "chat",
                    "role": assistant_name,
                    "tool": {"role": "Tool", "text": notice, "command": None},
                    "text": whitespace_columns_to_markdown(answer),
                    "stats": stats_to_dict(stats),
                }
            agent_text = text[len("/agent "):].strip() if text.startswith("/agent ") else text
            agent_messages = messages[:-1] + [{"role": "user", "content": agent_text}]
            answer, stats, commands = invoke_langchain_agent_with_trace(
                model,
                agent_messages,
                max_tool_rounds=20 if agent_profile == "builder" else 4,
                keep_alive=keep_alive,
                enforce_work_dir=agent_profile == "builder",
            )
            return {
                "ok": True,
                "role": "Agent",
                "text": whitespace_columns_to_markdown(answer),
                "stats": stats_to_dict(stats),
                "commands": commands,
            }

        tool_result = run_tool_command(text)
        commands: list[str] = []
        if tool_result:
            tool_role, tool_text, tool_cmd = tool_result
            if tool_cmd:
                commands.append(tool_cmd)
            follow_up = messages + [{"role": "system", "content": "Tool result:\n" + tool_text}]
            answer, stats = collect_chat(model, follow_up, verbose, keep_alive=keep_alive)
            return {
                "ok": True,
                "role": assistant_name,
                "tool": {"role": tool_role, "text": tool_text, "command": tool_cmd},
                "commands": commands,
                "text": whitespace_columns_to_markdown(answer),
                "stats": stats_to_dict(stats),
            }

        answer, stats = collect_chat(model, messages, verbose, keep_alive=keep_alive)
        return {"ok": True, "role": assistant_name, "text": whitespace_columns_to_markdown(answer), "stats": stats_to_dict(stats)}
    except (OllamaError, LangChainUnavailableError, ValueError, RuntimeError) as exc:
        return {"ok": False, "error": str(exc)}


INDEX_HTML = r"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Ollama Web TUI</title>
  <style>
    :root {
      color-scheme: dark;
      --bg: #0f1216;
      --panel: #171b21;
      --panel-2: #222832;
      --panel-3: #11151b;
      --border: #303844;
      --border-soft: rgba(255,255,255,0.08);
      --text: #edf1f7;
      --muted: #9aa6b4;
      --accent: #60b7ff;
      --ok: #83d18f;
      --warn: #f0c36a;
      --error: #ff8d8d;
      --shadow: 0 18px 50px rgba(0,0,0,0.28);
      --code-bg: #0d1117;
      --code-header-bg: #161b22;
      --canvas-bg: #0e1117;
      --canvas-preview-bg: #fafaf7;
      --canvas-preview-text: #1a1a1a;
      --canvas-preview-border: #d0d0d0;
      --canvas-preview-muted-bg: #f0efe9;
      --canvas-h1: #0b66c3;
      --canvas-h2: #7a3db8;
      --canvas-h3: #0b7a53;
      --canvas-table-head: #e6f0fb;
      --canvas-accent-soft: rgba(11, 102, 195, 0.12);
    }
    :root[data-theme="light"] {
      color-scheme: light;
      --bg: #f5f7fa;
      --panel: #ffffff;
      --panel-2: #eef2f7;
      --panel-3: #f9fbfd;
      --border: #c8d0da;
      --border-soft: rgba(24,33,47,0.08);
      --text: #18212f;
      --muted: #5d6978;
      --accent: #0b66c3;
      --shadow: 0 18px 50px rgba(40,54,75,0.14);
      --ok: #147a32;
      --warn: #8a5a00;
      --error: #b42318;
      --code-bg: #f6f8fa;
      --code-header-bg: #e9eef5;
      --canvas-bg: #ffffff;
      --canvas-preview-bg: #ffffff;
      --canvas-preview-text: #161616;
      --canvas-preview-border: #cfd6de;
      --canvas-preview-muted-bg: #f4f6f8;
      --canvas-h1: #0b66c3;
      --canvas-h2: #7a3db8;
      --canvas-h3: #0b7a53;
      --canvas-table-head: #e6f0fb;
      --canvas-accent-soft: rgba(11, 102, 195, 0.12);
    }
    :root[data-theme="dim"] {
      color-scheme: dark;
      --bg: #171a1f;
      --panel: #20242b;
      --panel-2: #2a3038;
      --panel-3: #191d24;
      --border: #47515f;
      --border-soft: rgba(255,255,255,0.1);
      --text: #edf1f5;
      --muted: #b8c0ca;
      --accent: #8fcef8;
      --ok: #8fd49b;
      --warn: #e0c476;
      --error: #ff9b9b;
      --code-bg: #141820;
      --code-header-bg: #202632;
      --canvas-bg: #151a22;
      --canvas-preview-bg: #f3f4f0;
      --canvas-preview-text: #1d2228;
      --canvas-preview-border: #c7ced6;
      --canvas-preview-muted-bg: #eceee8;
      --canvas-h1: #2474c6;
      --canvas-h2: #7a4ab4;
      --canvas-h3: #177a5a;
      --canvas-table-head: #e4edf7;
      --canvas-accent-soft: rgba(36, 116, 198, 0.12);
    }
    :root[data-theme="warm"] {
      color-scheme: light;
      --bg: #f6f3ee;
      --panel: #fffaf2;
      --panel-2: #f0e8dc;
      --panel-3: #fffdf8;
      --border: #d3c7b8;
      --border-soft: rgba(33,27,21,0.08);
      --text: #211b15;
      --muted: #6f6255;
      --accent: #9b4d13;
      --ok: #3f7d3f;
      --warn: #936b00;
      --error: #b33a2b;
      --code-bg: #f2ede5;
      --code-header-bg: #e7ddcf;
      --canvas-bg: #fffaf2;
      --canvas-preview-bg: #fffdf7;
      --canvas-preview-text: #211b15;
      --canvas-preview-border: #d8ccb9;
      --canvas-preview-muted-bg: #f3eadf;
      --canvas-h1: #9b4d13;
      --canvas-h2: #6d4c9b;
      --canvas-h3: #3f7d3f;
      --canvas-table-head: #f1e2d0;
      --canvas-accent-soft: rgba(155, 77, 19, 0.13);
    }
    :root[data-theme="contrast"] {
      color-scheme: dark;
      --bg: #000000;
      --panel: #050505;
      --panel-2: #111111;
      --panel-3: #000000;
      --border: #777777;
      --border-soft: #777777;
      --text: #ffffff;
      --muted: #d0d0d0;
      --accent: #00d5ff;
      --ok: #00ff66;
      --warn: #ffd400;
      --error: #ff4d4d;
      --code-bg: #000000;
      --code-header-bg: #111111;
      --canvas-bg: #000000;
      --canvas-preview-bg: #ffffff;
      --canvas-preview-text: #000000;
      --canvas-preview-border: #555555;
      --canvas-preview-muted-bg: #eeeeee;
      --canvas-h1: #005fcc;
      --canvas-h2: #6b00cc;
      --canvas-h3: #006b3a;
      --canvas-table-head: #dbeafe;
      --canvas-accent-soft: rgba(0, 95, 204, 0.16);
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      background: var(--bg);
      color: var(--text);
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
      font-size: var(--ui-font-size, 14px);
      line-height: 1.45;
      height: 100vh;
      overflow: hidden;
    }
    .app {
      height: 100vh;
      display: grid;
      grid-template-rows: auto 1fr auto;
      min-width: 320px;
    }
    header {
      background: var(--panel);
      border-bottom: 1px solid var(--border);
      padding: 10px 12px;
      box-shadow: 0 1px 0 var(--border-soft);
    }
    .topbar {
      display: grid;
      grid-template-columns: minmax(190px, auto) minmax(0, 1fr);
      gap: 12px;
      align-items: start;
      margin-bottom: 10px;
    }
    .brand {
      display: flex;
      align-items: center;
      gap: 10px;
      min-width: 0;
      cursor: pointer;
      width: fit-content;
      border: 0;
      padding: 0;
      background: transparent;
      color: inherit;
      text-align: left;
    }
    .brand:hover .brand-mark,
    .brand:focus-visible .brand-mark {
      border-color: var(--accent);
      color: var(--accent);
    }
    .brand-wrap {
      position: relative;
    }
    .brand-mark {
      width: 34px;
      height: 34px;
      display: grid;
      place-items: center;
      border: 1px solid var(--border);
      border-radius: 8px;
      background: var(--panel-2);
      color: var(--accent);
      font-weight: 800;
      letter-spacing: 0;
      flex: 0 0 auto;
    }
    .brand-title {
      font-size: 16px;
      font-weight: 750;
      line-height: 1.1;
      letter-spacing: 0;
    }
    .brand-subtitle {
      color: var(--muted);
      font-size: 12px;
      margin-top: 2px;
    }
    .name-popover {
      display: none;
      position: absolute;
      top: calc(100% + 8px);
      left: 0;
      z-index: 900;
      width: min(320px, 88vw);
      padding: 10px;
      border: 1px solid var(--border);
      border-radius: 10px;
      background: var(--panel);
      box-shadow: var(--shadow);
    }
    .name-popover.open { display: block; }
    .name-popover label {
      display: block;
      color: var(--muted);
      font-size: 12px;
      margin-bottom: 6px;
    }
    .name-popover-row {
      display: grid;
      grid-template-columns: 1fr auto;
      gap: 8px;
    }
    .brand + .runtime-banner { margin-top: 8px; }
    .status {
      display: flex;
      flex-wrap: wrap;
      gap: 6px;
      align-items: center;
      justify-content: flex-end;
    }
    .status > span {
      display: inline-flex;
      align-items: center;
      gap: 4px;
      min-height: 28px;
      padding: 4px 8px;
      border: 1px solid var(--border);
      border-radius: 999px;
      background: var(--panel-3);
      color: var(--muted);
      white-space: nowrap;
      font-size: 12px;
    }
    .status strong { color: var(--accent); font-weight: 700; }
    #ollamaStatus { transition: color 0.4s; }
    #ollamaStatus.online  { color: var(--ok); }
    #ollamaStatus.offline { color: var(--error); }
    #ollamaStatus.checking { color: var(--muted); }
    .runtime-banner {
      display: flex; flex-wrap: wrap; gap: 6px; align-items: center;
      padding: 0;
      background: transparent;
      border: 0;
      border-radius: 0;
      font-size: 12px;
    }
    .rt-pill {
      padding: 4px 8px; border-radius: 999px; background: var(--panel-3);
      color: var(--muted); border: 1px solid var(--border);
    }
    .rt-pill.rt-danger {
      background: #2d0f0f; color: #ff7b72; border-color: #da3633; font-weight: 700;
    }
    .rt-link { text-decoration: none; }
    .rt-link:hover { color: var(--accent); }
    .runtime-banner .hide { display: none; }
    .controls {
      display: grid;
      grid-template-columns: minmax(220px, 2fr) repeat(4, minmax(130px, 1fr)) minmax(160px, 1fr) auto auto;
      gap: 8px;
      align-items: center;
      padding: 8px;
      border: 1px solid var(--border);
      border-radius: 10px;
      background: var(--panel-3);
    }
    #gpuSelect { display: none; }  /* shown via JS only when multiple GPUs configured */
    .cap-tag {
      font-size: 10px; color: var(--muted); opacity: 0.8;
    }
    .tts-controls {
      display: flex;
      gap: 8px;
      align-items: center;
      margin-top: 8px;
      flex-wrap: wrap;
      padding: 7px 8px;
      border: 1px solid var(--border-soft);
      border-radius: 10px;
      background: rgba(255,255,255,0.018);
    }
    .tts-controls select { flex: 1; min-width: 160px; }
    .tts-controls input[type=range] { flex: 0 0 110px; accent-color: var(--accent); cursor: pointer; }
    .tts-speed-label { color: var(--muted); white-space: nowrap; }
    .check-control {
      display: inline-flex;
      align-items: center;
      justify-content: center;
      gap: 6px;
      min-height: 34px;
      padding: 7px 9px;
      border: 1px solid var(--border);
      border-radius: 8px;
      background: var(--panel-2);
      color: var(--text);
      white-space: nowrap;
      cursor: pointer;
      user-select: none;
    }
    .check-control input { width: auto; margin: 0; accent-color: var(--accent); cursor: pointer; }
    .tts-btn {
      background: none;
      border: none;
      cursor: pointer;
      font-size: 15px;
      padding: 1px 4px;
      border-radius: 4px;
      color: var(--muted);
      min-height: unset;
      line-height: 1;
    }
    .tts-btn:hover { color: var(--accent); }
    .tts-btn.speaking { color: var(--ok); }
    select, button, input {
      background: var(--panel-2);
      color: var(--text);
      border: 1px solid var(--border);
      border-radius: 8px;
      padding: 7px 9px;
      font: inherit;
      min-height: 34px;
    }
    button { cursor: pointer; }
    select:focus, button:focus-visible, input:focus, textarea:focus {
      outline: 2px solid rgba(96,183,255,0.35);
      outline-offset: 1px;
    }
    button:hover, select:hover { border-color: color-mix(in srgb, var(--accent) 55%, var(--border)); }
    button.active { border-color: var(--accent); color: var(--accent); background: color-mix(in srgb, var(--accent) 14%, var(--panel-2)); }
    .workspace {
      min-height: 0;
      display: grid;
      grid-template-columns: minmax(0, 1fr) minmax(320px, 26vw);
      overflow: hidden;
    }
    .app:not(.verbose-on) .workspace { grid-template-columns: minmax(0, 1fr); }
    .app:not(.verbose-on) #verbosePane { display: none; }
    main {
      overflow: auto;
      padding: 18px clamp(12px, 2vw, 28px);
      scroll-behavior: smooth;
      height: 100%;
    }
    .msg {
      max-width: 1120px;
      border: 1px solid var(--border);
      border-left: 4px solid var(--border);
      border-radius: 10px;
      padding: 10px 12px;
      margin: 0 0 12px;
      background: var(--panel);
      box-shadow: 0 1px 0 var(--border-soft);
      white-space: pre-wrap;
      word-break: break-word;
    }
    .msg.user { border-left-color: var(--accent); margin-left: auto; background: color-mix(in srgb, var(--accent) 8%, var(--panel)); }
    .msg.ollama, .msg.agent { border-left-color: var(--warn); }
    .msg.tool { border-left-color: var(--ok); color: var(--muted); background: var(--panel-3); }
    .msg.error { border-left-color: var(--error); color: var(--error); }
    #verbosePane {
      min-width: 0;
      display: grid;
      grid-template-rows: auto 1fr;
      border-left: 1px solid var(--border);
      background: var(--panel-3);
      overflow: hidden;
    }
    .verbose-title {
      display: flex;
      align-items: center;
      justify-content: space-between;
      gap: 8px;
      padding: 6px 8px;
      border-bottom: 1px solid var(--border);
      color: var(--accent);
      font-weight: 700;
      font-size: 12px;
      background: var(--panel);
    }
    #verboseClearBtn {
      min-height: unset;
      padding: 2px 6px;
      font-size: 11px;
      color: var(--muted);
      background: none;
    }
    #verboseLog {
      overflow: auto;
      padding: 10px;
      font: 12px/1.4 ui-monospace, SFMono-Regular, Menlo, Consolas, "Liberation Mono", monospace;
      line-height: 1.35;
      color: var(--muted);
    }
    .vmsg {
      border: 1px solid var(--border-soft);
      border-left: 3px solid var(--border);
      border-radius: 8px;
      margin: 0 0 8px;
      padding: 7px 8px;
      background: var(--panel);
      white-space: pre-wrap;
      overflow-wrap: anywhere;
    }
    .vmsg.cmd { border-color: var(--accent); color: var(--text); }
    .vmsg.step { border-color: var(--muted); }
    .vmsg.stats { border-color: var(--ok); color: var(--text); }
    .vmsg.context { border-color: var(--warn); color: var(--text); background: color-mix(in srgb, var(--warn) 8%, var(--panel)); }
    .vmsg.tool { border-color: var(--ok); }
    /* Rendered markdown inside messages */
    .msg-body { white-space: normal; }
    .msg-body p { margin: 0 0 8px; }
    .msg-body p:last-child { margin-bottom: 0; }
    .msg-body h1,.msg-body h2,.msg-body h3 { margin: 10px 0 4px; color: var(--accent); font-size: 1em; }
    .msg-body ul,.msg-body ol { margin: 4px 0 8px 18px; padding: 0; }
    .msg-body li { margin-bottom: 2px; }
    .msg-body blockquote { border-left: 3px solid var(--muted); margin: 4px 0; padding: 2px 8px; color: var(--muted); }
    .msg-body code { background: var(--panel-2); border: 1px solid var(--border); border-radius: 5px; padding: 1px 5px; font: 0.92em/1.3 ui-monospace, SFMono-Regular, Menlo, Consolas, "Liberation Mono", monospace; }
    .msg-body pre { background: var(--code-bg); border: 1px solid var(--border); border-radius: 8px; margin: 8px 0; overflow-x: auto; position: relative; }
    .msg-body pre code { background: none; border: none; padding: 12px 14px; display: block; white-space: pre; font-size: 0.9em; }
    .msg-body strong { color: var(--text); }
    .msg-body .table-wrap { margin: 8px 0; max-width: 100%; overflow-x: auto; }
    .msg-body table { border-collapse: separate; border-spacing: 0; width: auto; min-width: min(420px, 100%); margin: 0; font-size: 0.92em; border: 1px solid var(--border); border-radius: 8px; overflow: hidden; }
    .msg-body th,.msg-body td { border: 1px solid var(--border); padding: 6px 8px; text-align: left; vertical-align: top; }
    .msg-body th { background: var(--panel-2); color: var(--accent); font-weight: 700; }
    .msg-body tr:nth-child(even) td { background: rgba(255,255,255,0.025); }
    .code-header { display: flex; justify-content: space-between; align-items: center; background: var(--code-header-bg); border-bottom: 1px solid var(--border); padding: 4px 10px; border-radius: 6px 6px 0 0; font-size: 11px; color: var(--muted); }
    .copy-btn { background: none; border: none; color: var(--muted); cursor: pointer; font-size: 11px; min-height: unset; padding: 2px 6px; }
    .copy-btn:hover { color: var(--accent); }
    /* Canvas popup window */
    #canvasWindow {
      display: none;
      position: fixed;
      inset: 12px;
      width: auto;
      height: auto;
      min-height: 360px; min-width: 380px;
      transform: none;
      background: var(--panel);
      border: 1px solid var(--border);
      border-radius: 8px;
      box-shadow: 0 12px 48px rgba(0,0,0,0.65);
      z-index: 500;
      flex-direction: column;
      overflow: hidden;
      resize: both;
    }
    #canvasWindow.open { display: flex; }
    #canvasWindow.maximized {
      inset: 0;
      width: auto;
      height: auto;
      min-width: 0;
      min-height: 0;
      border-radius: 0;
      resize: none;
    }
    .canvas-titlebar {
      display: flex; align-items: center; gap: 6px;
      padding: 6px 10px;
      background: var(--panel-2);
      border-bottom: 1px solid var(--border);
      cursor: move; user-select: none; flex-shrink: 0;
    }
    .canvas-titlebar span { flex: 1; color: var(--accent); font-weight: 700; font-size: 13px; }
    .canvas-titlebar button { min-height: unset; padding: 3px 8px; font-size: 11px; }
    .canvas-toolbar {
      display: flex;
      gap: 6px; align-items: center; flex-wrap: wrap;
      padding: 6px 8px;
      background: var(--panel);
      border-bottom: 1px solid var(--border);
      flex-shrink: 0;
    }
    .canvas-toolbar button { min-height: unset; padding: 4px 9px; font-size: 12px; white-space: nowrap; }
    .canvas-toolbar select { min-height: unset; padding: 4px 8px; font-size: 12px; flex: 0 1 240px; min-width: 150px; }
    .canvas-toolbar .export-group { display: flex; gap: 4px; justify-content: flex-end; flex-wrap: wrap; margin-left: auto; }
    #canvasStatus { color: var(--muted); font-size: 11px; flex: 1 1 150px; min-width: 120px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
    #canvasFileDeleteBtn { background: none; border-color: var(--border); color: var(--muted); }
    #canvasFileDeleteBtn:hover { color: var(--error); border-color: var(--error); }
    #canvasBody {
      flex: 1 1 auto;
      min-height: 0;
      display: flex;
      overflow: hidden;
      background: var(--canvas-bg);
    }
    #canvasEditor {
      flex: 1;
      width: 100%;
      height: 100%;
      min-height: 0;
      max-height: none;
      box-sizing: border-box;
      resize: none;
      background: var(--canvas-bg);
      color: var(--text);
      border: none;
      padding: 20px 24px;
      font: 14px/1.6 "SFMono-Regular", Consolas, "Liberation Mono", monospace;
      outline: none;
      tab-size: 4;
      overflow: auto;
    }
    #canvasPreview {
      flex: 1;
      width: 100%;
      height: 100%;
      min-height: 0;
      box-sizing: border-box;
      overflow: auto;
      background: var(--canvas-preview-bg);
      color: var(--canvas-preview-text);
      padding: 34px clamp(24px, 5vw, 72px);
      font: 14px/1.7 Georgia, "Times New Roman", serif;
    }
    #canvasPreview > *:first-child { margin-top: 0; }
    #canvasPreview h1, #canvasPreview h2, #canvasPreview h3,
    #canvasPreview h4, #canvasPreview h5, #canvasPreview h6 {
      font-family: "Helvetica Neue", Arial, sans-serif;
      line-height: 1.25;
      margin: 1.4em 0 0.5em;
      font-weight: 600;
    }
    #canvasPreview h1 {
      font-size: 28px;
      color: var(--canvas-h1);
      border-bottom: 3px solid var(--canvas-h1);
      padding-bottom: 8px;
      margin-top: 0.2em;
    }
    #canvasPreview h2 {
      font-size: 21px;
      color: var(--canvas-h2);
      border-left: 5px solid var(--canvas-h2);
      background: var(--canvas-accent-soft);
      padding: 6px 10px;
      border-radius: 4px;
    }
    #canvasPreview h3 { font-size: 17px; color: var(--canvas-h3); }
    #canvasPreview h4 { font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px; color: var(--canvas-h2); }
    #canvasPreview p { margin: 0.6em 0; }
    #canvasPreview ul, #canvasPreview ol { margin: 0.5em 0 0.8em; padding-left: 28px; }
    #canvasPreview li { margin: 0.2em 0; }
    #canvasPreview li::marker { color: var(--canvas-h2); font-weight: 700; }
    #canvasPreview a { color: #1a5fb4; text-decoration: none; border-bottom: 1px dotted #1a5fb4; }
    #canvasPreview a:hover { border-bottom-style: solid; }
    #canvasPreview blockquote {
      margin: 1em 0; padding: 6px 14px;
      border-left: 3px solid var(--canvas-preview-border); background: var(--canvas-preview-muted-bg);
      color: var(--canvas-preview-text); font-style: italic;
    }
    #canvasPreview hr { border: none; border-top: 1px solid var(--canvas-preview-border); margin: 1.6em 0; }
    #canvasPreview code {
      font-family: "SFMono-Regular", Consolas, "Liberation Mono", monospace;
      font-size: 0.92em; background: var(--canvas-preview-muted-bg); padding: 1px 5px; border-radius: 3px;
    }
    #canvasPreview pre {
      background: var(--canvas-preview-muted-bg); border: 1px solid var(--canvas-preview-border); border-left: 5px solid var(--canvas-h3); border-radius: 5px;
      padding: 12px 14px; overflow-x: auto; line-height: 1.5;
    }
    #canvasPreview pre code { background: none; padding: 0; font-size: 0.92em; }
    #canvasPreview pre.diagram {
      background: #0d1117;
      color: #dbeafe;
      border-color: #27496d;
      border-left-color: var(--canvas-h1);
      font-family: "SFMono-Regular", Consolas, "Liberation Mono", monospace;
      font-size: 12px;
      line-height: 1.35;
      white-space: pre;
    }
    #canvasPreview pre.diagram code { color: inherit; font-size: inherit; }
    #canvasPreview table {
      border-collapse: collapse; margin: 1em 0; width: 100%;
      font-size: 13px;
    }
    #canvasPreview th, #canvasPreview td {
      border: 1px solid var(--canvas-preview-border); padding: 6px 10px; vertical-align: top;
    }
    #canvasPreview thead th { background: var(--canvas-table-head); color: var(--canvas-h1); font-weight: 700; }
    #canvasPreview tbody tr:nth-child(even) { background: var(--canvas-preview-muted-bg); }
    #canvasPreview tbody tr:hover { background: var(--canvas-accent-soft); }
    #canvasPreview .code-header { display: none; }
    #canvasChat {
      flex: 0 0 auto;
      max-height: 24vh;
      overflow: auto;
      padding: 8px 10px;
      background: var(--panel);
      border-top: 1px solid var(--border);
    }
    #canvasChat:empty { display: none; }
    .canvas-chat-msg {
      border-left: 3px solid var(--border);
      margin: 0 0 8px;
      padding: 5px 8px;
      white-space: pre-wrap;
      word-break: break-word;
    }
    .canvas-chat-msg.user { border-color: var(--accent); }
    .canvas-chat-msg.ai { border-color: var(--warn); }
    .canvas-chat-msg.tool { border-color: var(--ok); color: var(--muted); }
    .canvas-chat-role { display: block; color: var(--muted); margin-bottom: 2px; font-size: 12px; }
    #canvasChatBar {
      display: grid;
      grid-template-columns: 1fr auto;
      gap: 8px;
      padding: 8px 10px;
      background: var(--panel);
      border-top: 1px solid var(--border);
      flex-shrink: 0;
    }
    #canvasChatInput {
      resize: vertical;
      min-height: 38px;
      max-height: 120px;
      background: var(--panel-2);
      color: var(--text);
      border: 1px solid var(--border);
      border-radius: 6px;
      padding: 8px 10px;
      font: inherit;
    }
    .canvas-toolbar #canvasViewBtn.active { border-color: var(--accent); color: var(--accent); }
    #canvasToggleBtn { border-color: var(--border); color: var(--muted); }
    #canvasToggleBtn.active { border-color: var(--accent); color: var(--accent); }
    #canvasToggleBtn.updated { border-color: #5b9ef7; color: #5b9ef7; box-shadow: 0 0 6px rgba(91,158,247,0.5); }
    @media (max-width: 760px) {
      #canvasWindow { inset: 4px; min-width: 0; }
      .canvas-toolbar .export-group { margin-left: 0; justify-content: flex-start; }
      #canvasChatBar { grid-template-columns: 1fr; }
      #canvasPreview { padding: 22px 18px; }
      #canvasEditor { padding: 16px 14px; }
    }
    .workspace { height: 100%; overflow: hidden; }
    .role {
      display: inline-flex;
      align-items: center;
      min-height: 22px;
      padding: 2px 7px;
      border: 1px solid var(--border-soft);
      border-radius: 999px;
      background: var(--panel-3);
      color: var(--muted);
      margin-bottom: 7px;
      font-size: 12px;
      font-weight: 650;
    }
    footer {
      display: grid;
      grid-template-columns: minmax(0, 1fr) auto auto auto auto auto auto;
      gap: 8px;
      padding: 10px 12px;
      background: var(--panel);
      border-top: 1px solid var(--border);
      box-shadow: 0 -1px 0 var(--border-soft);
    }
    #micBtn { background: none; border-color: var(--border); color: var(--muted); }
    #micBtn:hover { color: var(--accent); }
    #micBtn.unavailable { border-color: var(--border); color: var(--muted); opacity: .72; }
    #micBtn.active { border-color: var(--accent); color: var(--accent); }
    #micBtn.recording {
      border-color: var(--error); color: var(--error);
      animation: mic-pulse 1s ease-in-out infinite;
    }
    .voice-auto {
      display: inline-flex;
      align-items: center;
      gap: 5px;
      min-height: 38px;
      padding: 0 8px;
      border: 1px solid var(--border);
      border-radius: 10px;
      color: var(--muted);
      white-space: nowrap;
      font-size: 12px;
    }
    .voice-auto input { margin: 0; }
    .voice-auto.active {
      border-color: var(--accent);
      color: var(--accent);
      background: color-mix(in srgb, var(--accent) 9%, transparent);
    }
    @keyframes mic-pulse {
      0%, 100% { box-shadow: 0 0 0 0 rgba(255,141,141,0.4); }
      50%       { box-shadow: 0 0 0 6px rgba(255,141,141,0); }
    }
    #stopBtn { background: none; border-color: var(--error); color: var(--error); display: none; }
    #stopBtn.active { display: inline-flex; }
    #ctxBox { display: inline-flex; align-items: center; gap: 6px; }
    .ctx-bar {
      display: inline-block; width: 70px; height: 8px;
      background: var(--panel-2); border: 1px solid var(--border); border-radius: 4px;
      overflow: hidden; vertical-align: middle;
    }
    .ctx-fill {
      display: block; height: 100%; width: 0%;
      background: var(--accent); transition: width 0.25s ease, background 0.25s ease;
    }
    .ctx-fill.warn { background: #d6b34d; }
    .ctx-fill.danger { background: var(--error); }
    #sessionSelect { max-width: 220px; }
    #sessionDeleteBtn { background: none; border-color: var(--border); color: var(--muted); padding: 4px 8px; }
    #sessionDeleteBtn:hover { color: var(--error); border-color: var(--error); }
    .spinner {
      display: inline-block; width: 14px; height: 14px;
      border: 2px solid var(--border); border-top-color: var(--accent);
      border-radius: 50%; animation: spin 0.8s linear infinite;
      vertical-align: -2px; margin-right: 6px;
    }
    @keyframes spin { to { transform: rotate(360deg); } }
    #genIndicator {
      display: none; align-items: center; gap: 6px;
      color: var(--muted); font-size: 12px; padding: 4px 10px;
    }
    #genIndicator.active { display: inline-flex; }
    .thinking-block {
      margin: 6px 0; padding: 6px 10px;
      background: var(--panel-2); border-left: 3px solid var(--muted);
      color: var(--muted); font-size: 12px; font-style: italic;
      white-space: pre-wrap; border-radius: 4px;
    }
    .thinking-block summary {
      cursor: pointer; user-select: none; font-style: normal;
      color: var(--muted); font-size: 11px; text-transform: uppercase;
      letter-spacing: 0.5px; margin-bottom: 4px;
    }
    .thinking-block[open] summary { margin-bottom: 6px; }
    #attachBtn { background: none; border-color: var(--border); color: var(--muted); }
    #attachBtn:hover { color: var(--accent); }
    .attach-preview {
      display: flex; align-items: center; gap: 6px;
      background: var(--panel-2); border: 1px solid var(--border);
      border-radius: 6px; padding: 4px 8px; margin-bottom: 4px;
      font-size: 12px; color: var(--muted);
    }
    .attach-preview button {
      background: none; border: none; color: var(--error); cursor: pointer;
      padding: 0 2px; min-height: unset; font-size: 13px;
    }
    #input {
      resize: none;
      min-height: 44px;
      max-height: 140px;
      background: var(--panel-2);
      color: var(--text);
      border: 1px solid var(--border);
      border-radius: 10px;
      padding: 11px 12px;
      font: inherit;
    }
    .hint {
      color: var(--muted);
      margin-top: 6px;
      white-space: nowrap;
      overflow: hidden;
      text-overflow: ellipsis;
    }
    .mobile-actions {
      display: none;
    }
    .mobile-tabs {
      display: none;
    }
    @media (max-width: 760px) {
      body { height: 100dvh; }
      .app {
        height: 100dvh;
        grid-template-rows: auto 1fr auto auto;
      }
      header {
        padding: 8px;
        max-height: 52dvh;
        overflow: auto;
      }
      .topbar {
        grid-template-columns: 1fr;
        gap: 8px;
        margin-bottom: 8px;
      }
      .brand-mark { width: 30px; height: 30px; border-radius: 7px; }
      .brand-title { font-size: 15px; }
      .brand-subtitle { display: none; }
      .runtime-banner { margin-top: 6px; }
      .rt-pill { padding: 3px 7px; font-size: 11px; }
      .status {
        justify-content: flex-start;
        gap: 5px;
        max-height: 68px;
        overflow: hidden;
      }
      .status > span {
        min-height: 25px;
        padding: 3px 7px;
        font-size: 11px;
      }
      #ctxBox { display: none; }
      .mobile-actions {
        display: grid;
        grid-template-columns: 1fr auto auto;
        gap: 8px;
        margin-bottom: 8px;
      }
      .mobile-actions button {
        min-height: 40px;
      }
      .controls,
      .tts-controls,
      .hint {
        display: none;
      }
      .app.mobile-settings-open .controls {
        display: grid;
        grid-template-columns: 1fr;
        margin-top: 8px;
      }
      .app.mobile-settings-open .tts-controls {
        display: grid;
        grid-template-columns: 1fr;
      }
      .controls select,
      .controls button,
      .tts-controls select,
      .tts-controls button,
      .tts-controls input {
        width: 100%;
        min-height: 42px;
      }
      #sessionSelect { max-width: none; }
      .workspace {
        grid-template-columns: 1fr;
        grid-template-rows: minmax(0, 1fr);
      }
      main {
        padding: 10px 8px;
      }
      .msg {
        max-width: 100%;
        border-radius: 9px;
        padding: 9px 10px;
        margin-bottom: 9px;
      }
      #verbosePane {
        border-left: 0;
        border-top: 0;
      }
      .app:not(.verbose-on) #verbosePane { display: none; }
      .app.mobile-tab-chat #verbosePane { display: none; }
      .app.mobile-tab-chat main { display: block; }
      .app.mobile-tab-activity main { display: none; }
      .app.mobile-tab-activity #verbosePane { display: grid; }
      .app.mobile-tab-canvas main,
      .app.mobile-tab-canvas #verbosePane { display: none; }
      #verboseLog {
        padding: 8px;
        font-size: 12px;
      }
      footer {
        grid-template-columns: 1fr auto auto auto auto;
        gap: 6px;
        padding: 8px;
      }
      footer > div:first-child {
        min-width: 0;
      }
      #input {
        min-height: 46px;
        max-height: 112px;
        padding: 11px;
      }
      #attachBtn,
      #micBtn,
      #stopBtn,
      #send {
        min-width: 44px;
        min-height: 44px;
        padding: 7px 9px;
      }
      .voice-auto {
        min-height: 44px;
        padding: 0 9px;
      }
      #readLastBtn { display: none; }
      #genIndicator {
        grid-column: 1 / -1;
        justify-content: center;
        min-height: 28px;
        padding: 2px 8px;
      }
      .mobile-tabs {
        display: grid;
        grid-template-columns: repeat(3, 1fr);
        gap: 6px;
        padding: 6px 8px 8px;
        background: var(--panel);
        border-top: 1px solid var(--border);
      }
      .mobile-tabs button {
        min-height: 40px;
        border-radius: 10px;
        font-size: 12px;
      }
      .mobile-tabs button.active {
        border-color: var(--accent);
        color: var(--accent);
        background: color-mix(in srgb, var(--accent) 14%, var(--panel-2));
      }
      #canvasWindow {
        inset: 0;
        min-width: 0;
        min-height: 0;
        border-radius: 0;
        resize: none;
        z-index: 800;
      }
      #canvasWindow.open {
        display: flex;
      }
      .canvas-titlebar {
        cursor: default;
        min-height: 44px;
      }
      .canvas-toolbar {
        max-height: 30dvh;
        overflow: auto;
      }
      #canvasBody {
        min-height: 0;
      }
      #canvasChat {
        max-height: 22dvh;
      }
      #canvasChatBar {
        grid-template-columns: 1fr;
      }
    }
  </style>
</head>
<body>
  <div class="app">
    <header>
      <div class="topbar">
        <div>
          <div class="brand-wrap">
          <button class="brand" id="assistantNameButton" type="button" title="Change assistant name">
            <div class="brand-mark">L</div>
            <div>
              <div class="brand-title" id="assistantNameTitle">Lilith</div>
              <div class="brand-subtitle">Ollama Command Deck control surface</div>
            </div>
          </button>
          <div class="name-popover" id="assistantNamePopover">
            <label for="assistantNameInput">Assistant name</label>
            <div class="name-popover-row">
              <input id="assistantNameInput" type="text" maxlength="40" title="Assistant name" placeholder="Lilith">
              <button id="assistantNameSaveBtn" type="button">Save</button>
            </div>
          </div>
          </div>
          <div id="runtimeBanner" class="runtime-banner" title="Runtime configuration — click to open setup">
            <span id="rtDanger" class="rt-pill rt-danger hide">DANGEROUS MODE</span>
            <span class="rt-pill" id="rtOllama">Ollama: …</span>
            <span class="rt-pill" id="rtTools">Tools: …</span>
            <a href="/setup" class="rt-pill rt-link">Edit config</a>
          </div>
        </div>
        <div class="status">
          <span>Ollama <strong id="ollamaStatus" title="API status">checking</strong></span>
          <span>Model <strong id="modelLabel">loading</strong></span>
          <span>Mode <strong id="modeLabel">Chat</strong></span>
          <span>Profile <strong id="agentProfileLabel">General</strong></span>
          <span>Verbose <strong id="verboseLabel">OFF</strong></span>
          <span>Voice <strong id="personalityLabel">default</strong></span>
          <span id="ctxBox" title="Last prompt tokens / runtime context · advertised model maximum">
            Ctx <strong id="ctxLabel">—</strong>
            <span class="ctx-bar"><span class="ctx-fill" id="ctxFill"></span></span>
          </span>
        </div>
      </div>
      <div class="mobile-actions" aria-label="Mobile quick actions">
        <button id="mobileSettingsBtn" type="button">Settings</button>
        <button id="mobileActivityBtn" type="button">Activity</button>
        <button id="mobileCanvasBtn" type="button">Canvas</button>
      </div>
      <div class="controls">
        <select id="model"></select>
        <select id="modeSelect">
          <optgroup label="── Interaction ──">
            <option value="chat">💬 Chat / Voice</option>
            <option value="live">🎙 Live Chat</option>
            <option value="agent" selected>🤖 Agent</option>
          </optgroup>
          <optgroup label="── Chat Modes ──">
            <option value="coding">💻 Coding</option>
            <option value="creative">🎨 Creative</option>
            <option value="concise">⚡ Concise</option>
            <option value="teaching">📚 Teaching</option>
          </optgroup>
        </select>
        <select id="personality" title="Personality">
          <option value="default">🎭 default</option>
          <option value="friendly">😊 friendly</option>
          <option value="snarky">😏 snarky</option>
          <option value="rude">🙄 rude</option>
          <option value="formal">🎩 formal</option>
          <option value="pirate">🏴‍☠️ pirate</option>
          <option value="philosopher">🧐 philosopher</option>
          <option value="chef">👨‍🍳 chef</option>
        </select>
        <select id="agentProfile" title="Agent profile">
          <option value="general">Agent: General</option>
          <option value="ops">Agent: Ops</option>
          <option value="home">Agent: Home</option>
          <option value="code">Agent: Code</option>
          <option value="builder">Agent: Builder</option>
          <option value="research">Agent: Research</option>
          <option value="deep_research">Agent: Deep Research</option>
          <option value="writing">Agent: Writing</option>
          <option value="brief">Agent: Brief</option>
          <option value="debug">Agent: Debug</option>
          <option value="frontend">Agent: Frontend</option>
          <option value="skill_creator">Agent: Skill Creator</option>
        </select>
        <label class="check-control" title="Let Command Deck temporarily choose a better profile for each prompt. Your selected profile stays saved.">
          <input id="autoProfile" type="checkbox">
          <span>Auto profile</span>
        </label>
        <select id="themeSelect" title="Theme">
          <option value="system">Theme: System</option>
          <option value="dark">Theme: Dark</option>
          <option value="light">Theme: Light</option>
          <option value="dim">Theme: Dim</option>
          <option value="warm">Theme: Warm Light</option>
          <option value="contrast">Theme: High Contrast</option>
        </select>
        <select id="gpuSelect" title="GPU / Ollama endpoint"></select>
        <label class="check-control" title="Keep the selected Ollama model loaded for 30 minutes after each reply">
          <input id="keepAlive" type="checkbox">
          <span>Keep alive</span>
        </label>
        <select id="sessionSelect" title="Chat sessions"><option value="__new__">＋ New session</option></select>
        <button id="sessionDeleteBtn" type="button" title="Delete current session" style="display:none">🗑</button>
        <button id="verboseBtn" type="button">Verbose</button>
        <button id="canvasToggleBtn" type="button" title="Open document canvas">📄 Canvas</button>
      </div>
      <div class="tts-controls">
        <span style="color:var(--muted);white-space:nowrap;">Voice</span>
        <select id="voiceSelect"><option value="">Loading voices…</option></select>
        <span class="tts-speed-label">Speed: <strong id="ttsRateLabel">Normal</strong></span>
        <input id="ttsRate" type="range" min="-75" max="100" step="25" value="0" title="Speech speed">
        <span class="tts-speed-label">Pitch: <strong id="ttsPitchLabel">0</strong></span>
        <input id="ttsPitch" type="range" min="-12" max="12" step="1" value="0" title="Voice pitch">
        <span class="tts-speed-label">Tone</span>
        <select id="ttsTone" title="Voice tone">
          <option value="natural">Natural</option>
          <option value="dark">Dark</option>
          <option value="bright">Bright</option>
          <option value="radio">Radio</option>
          <option value="robotic">Robotic</option>
        </select>
        <span class="tts-speed-label">Vol: <strong id="ttsVolumeLabel">0%</strong></span>
        <input id="ttsVolume" type="range" min="-50" max="50" step="5" value="0" title="Voice volume">
        <span class="tts-speed-label">Text: <strong id="textZoomLabel">100%</strong></span>
        <input id="textZoom" type="range" min="85" max="140" step="5" value="100" title="Text zoom">
        <span class="tts-speed-label">Pause: <strong id="voicePauseLabel">1.8s</strong></span>
        <input id="voicePauseDelay" type="range" min="800" max="4000" step="100" value="1800" title="Voice conversation pause before auto-send">
        <button id="ttsStopBtn" type="button">Stop</button>
      </div>
      <div class="hint">Commands: /help /model /hosts /ssh /local /search /agent on|off /chat /verbose on|off /clear /quit</div>
    </header>
    <div class="workspace">
      <main id="chat"></main>
      <aside id="verbosePane" aria-label="Verbose event log">
        <div class="verbose-title">
          <span>Activity & Context</span>
          <button id="verboseClearBtn" type="button" title="Clear verbose log">Clear</button>
        </div>
        <div id="verboseLog"></div>
      </aside>
    </div>

    <!-- Floating canvas popup -->
    <div id="canvasWindow">
      <div class="canvas-titlebar" id="canvasDragBar">
        <span>📄 Document Canvas</span>
        <button type="button" id="canvasMaxBtn" title="Maximize canvas">□</button>
        <button type="button" id="canvasMinBtn" title="Minimize">─</button>
        <button type="button" id="canvasCloseBtn" title="Close">✕</button>
      </div>
      <div class="canvas-toolbar">
        <select id="canvasFileSelect" title="Saved canvas files"><option value="__current__">📝 Untitled</option></select>
        <button type="button" id="canvasNewBtn" title="New empty canvas">＋ New</button>
        <button type="button" id="canvasSaveBtn" title="Save canvas to file">💾 Save</button>
        <button type="button" id="canvasFileDeleteBtn" title="Delete saved canvas file" style="display:none">🗑</button>
        <button type="button" id="canvasAutoBtn">Auto: OFF</button>
        <button type="button" id="canvasViewBtn" title="Toggle rendered preview">👁 Preview</button>
        <button type="button" id="canvasSendAiBtn" title="Paste canvas content into chat so AI can read/edit it">📋 Send to AI</button>
        <button type="button" id="canvasClearBtn">Clear</button>
        <button type="button" id="canvasCopyBtn">Copy</button>
        <span id="canvasStatus">Unsaved canvas</span>
        <div class="export-group">
          <button type="button" class="export-btn" data-fmt="md">⬇ .md</button>
          <button type="button" class="export-btn" data-fmt="txt">⬇ .txt</button>
          <button type="button" class="export-btn" data-fmt="rtf" title="Opens in Word &amp; LibreOffice">⬇ .rtf</button>
          <button type="button" class="export-btn" data-fmt="pdf">⬇ PDF</button>
          <button type="button" class="export-btn" data-fmt="csv">⬇ .csv</button>
        </div>
      </div>
      <div id="canvasBody">
        <textarea id="canvasEditor" placeholder="Document canvas — paste, type, or send AI responses here.&#10;&#10;Markdown is supported: headings, lists, tables, code blocks, links."></textarea>
        <div id="canvasPreview" style="display:none"></div>
      </div>
      <div id="canvasChat"></div>
      <div id="canvasChatBar">
        <textarea id="canvasChatInput" placeholder="Ask about this canvas only. Example: check the section I added, reformat this document, or summarize the architecture section."></textarea>
        <button type="button" id="canvasChatSendBtn">Ask Canvas</button>
      </div>
    </div>
    <footer>
      <div style="display:flex;flex-direction:column;gap:4px;">
        <div id="attachPreview" style="display:none" class="attach-preview">
          <span id="attachName"></span>
          <button type="button" id="attachClearBtn" title="Remove attachment">✕</button>
        </div>
        <textarea id="input" placeholder="Type a message or slash command. Enter sends, Shift+Enter inserts a newline."></textarea>
      </div>
      <button id="attachBtn" type="button" title="Attach image or document">📎</button>
      <input type="file" id="fileInput" accept="image/*,.pdf,.docx,.doc,.csv,.txt,.md,.json" style="display:none">
      <button id="micBtn" type="button" title="Click to start/stop voice input">🎙</button>
      <label class="voice-auto" id="voiceAutoWrap" title="Send transcribed speech immediately after recording stops">
        <input id="voiceAutoSend" type="checkbox">
        <span>Auto-send</span>
      </label>
      <button id="readLastBtn" type="button" disabled title="Read last reply aloud">🔊 Read</button>
      <span id="genIndicator"><span class="spinner"></span><span id="genLabel">Thinking…</span></span>
      <button id="stopBtn" type="button" title="Stop generation">⏹ Stop</button>
      <button id="send" type="button">Send</button>
    </footer>
    <nav class="mobile-tabs" aria-label="Mobile navigation">
      <button type="button" class="active" data-mobile-tab="chat">Chat</button>
      <button type="button" data-mobile-tab="activity">Activity</button>
      <button type="button" data-mobile-tab="canvas">Canvas</button>
    </nav>
  </div>
  <script>
    const SETTINGS_KEY = "ollamaWebSettings";
    function loadSavedSettings() {
      try { return JSON.parse(localStorage.getItem(SETTINGS_KEY) || "{}") || {}; }
      catch (_) { return {}; }
    }
    const savedSettings = loadSavedSettings();
    const state = {
      model: "",
      messages: [],
      agentMode: true,
      verbose: true,
      chatMode: "default",
      voiceAutoSend: false,
      personality: "default",
      agentProfile: "general",
      activeAgentProfile: "general",
      autoProfile: false,
      assistantName: "Lilith",
      textZoom: 100,
      voicePauseDelay: 1800,
      keepAlive: false,
      gpu: "0",
      ...savedSettings,
    };
    const chat = document.getElementById("chat");
    const appRoot = document.querySelector(".app");
    const verboseLog = document.getElementById("verboseLog");
    const verboseClearBtn = document.getElementById("verboseClearBtn");
    const input = document.getElementById("input");
    const modelSelect = document.getElementById("model");
    const modeSelect = document.getElementById("modeSelect");
    const personality = document.getElementById("personality");
    const agentProfile = document.getElementById("agentProfile");
    const autoProfile = document.getElementById("autoProfile");
    const keepAlive = document.getElementById("keepAlive");
    const assistantNameInput = document.getElementById("assistantNameInput");
    const assistantNameTitle = document.getElementById("assistantNameTitle");
    const assistantNameButton = document.getElementById("assistantNameButton");
    const assistantNamePopover = document.getElementById("assistantNamePopover");
    const assistantNameSaveBtn = document.getElementById("assistantNameSaveBtn");
    const themeSelect = document.getElementById("themeSelect");
    const verboseBtn = document.getElementById("verboseBtn");
    const mobileSettingsBtn = document.getElementById("mobileSettingsBtn");
    const mobileActivityBtn = document.getElementById("mobileActivityBtn");
    const mobileCanvasBtn = document.getElementById("mobileCanvasBtn");
    const mobileTabBtns = [...document.querySelectorAll("[data-mobile-tab]")];

    function cleanAssistantName(value) {
      return String(value || "Lilith").replace(/[^A-Za-z0-9 _.-]/g, "").trim().slice(0, 40) || "Lilith";
    }

    function saveSettings() {
      localStorage.setItem(SETTINGS_KEY, JSON.stringify({
        model: state.model || "",
        agentMode: Boolean(state.agentMode),
        verbose: Boolean(state.verbose),
        chatMode: state.chatMode || "default",
        voiceAutoSend: Boolean(state.voiceAutoSend),
        personality: state.personality || "default",
        agentProfile: state.agentProfile || "general",
        autoProfile: Boolean(state.autoProfile),
        assistantName: cleanAssistantName(state.assistantName),
        textZoom: Number(state.textZoom) || 100,
        voicePauseDelay: Math.max(800, Math.min(4000, Number(state.voicePauseDelay) || 1800)),
        keepAlive: Boolean(state.keepAlive),
        gpu: state.gpu || "0",
      }));
    }

    function resolveTheme(theme) {
      if (theme === "system") {
        return window.matchMedia && window.matchMedia("(prefers-color-scheme: light)").matches ? "light" : "dark";
      }
      return theme || "dark";
    }

    function applyTheme(theme) {
      const selected = theme || "system";
      themeSelect.value = selected;
      document.documentElement.dataset.theme = resolveTheme(selected);
      localStorage.setItem("ollamaWebTheme", selected);
    }

    applyTheme(localStorage.getItem("ollamaWebTheme") || "system");
    if (window.matchMedia) {
      window.matchMedia("(prefers-color-scheme: light)").addEventListener("change", () => {
        if ((localStorage.getItem("ollamaWebTheme") || "system") === "system") applyTheme("system");
      });
    }
    themeSelect.addEventListener("change", () => applyTheme(themeSelect.value));
    personality.value = state.personality;
    agentProfile.value = state.agentProfile;
    autoProfile.checked = Boolean(state.autoProfile);
    state.activeAgentProfile = state.agentProfile || "general";
    assistantNameInput.value = cleanAssistantName(state.assistantName);

    function openAssistantNamePopover() {
      assistantNameInput.value = cleanAssistantName(state.assistantName);
      assistantNamePopover.classList.add("open");
      assistantNameInput.focus();
      assistantNameInput.select();
    }

    function closeAssistantNamePopover() {
      assistantNamePopover.classList.remove("open");
    }

    function saveAssistantNameFromPopover() {
      state.assistantName = cleanAssistantName(assistantNameInput.value);
      updateStatus();
      closeAssistantNamePopover();
    }

    // Map modeSelect value → state fields
    function applyModeSelect(value) {
      if (value === "agent") {
        state.agentMode = true;
        state.chatMode = "default";
      } else {
        state.agentMode = false;
        state.chatMode = value === "chat" ? "conversation" : value;
      }
    }

    modeSelect.addEventListener("change", () => {
      applyModeSelect(modeSelect.value);
      updateStatus();
    });

    // Sync modeSelect → current state
    function syncModeSelect() {
      if (state.agentMode) {
        modeSelect.value = "agent";
      } else if (state.chatMode === "default" || state.chatMode === "conversation") {
        modeSelect.value = "chat";
      } else if (state.chatMode === "live") {
        modeSelect.value = "live";
      } else {
        modeSelect.value = state.chatMode;
      }
    }
    const gpuSelect = document.getElementById("gpuSelect");
    const micBtn = document.getElementById("micBtn");
    const voiceAutoSend = document.getElementById("voiceAutoSend");
    const voiceAutoWrap = document.getElementById("voiceAutoWrap");
    const voiceSelect = document.getElementById("voiceSelect");
    const ttsStopBtn = document.getElementById("ttsStopBtn");
    const ttsRate = document.getElementById("ttsRate");
    const ttsRateLabel = document.getElementById("ttsRateLabel");
    const ttsPitch = document.getElementById("ttsPitch");
    const ttsPitchLabel = document.getElementById("ttsPitchLabel");
    const ttsTone = document.getElementById("ttsTone");
    const ttsVolume = document.getElementById("ttsVolume");
    const ttsVolumeLabel = document.getElementById("ttsVolumeLabel");
    const textZoom = document.getElementById("textZoom");
    const textZoomLabel = document.getElementById("textZoomLabel");
    const voicePauseDelay = document.getElementById("voicePauseDelay");
    const voicePauseLabel = document.getElementById("voicePauseLabel");
    const readLastBtn = document.getElementById("readLastBtn");
    const stopBtn = document.getElementById("stopBtn");
    const attachBtn = document.getElementById("attachBtn");
    const fileInput = document.getElementById("fileInput");
    const attachPreview = document.getElementById("attachPreview");
    const attachName = document.getElementById("attachName");
    const attachClearBtn = document.getElementById("attachClearBtn");

    let pendingAttachment = null;  // {type, data, filename} or null
    let currentController = null;

    // ── File attachment ───────────────────────────────────────────────────────
    attachBtn.addEventListener("click", () => fileInput.click());
    attachClearBtn.addEventListener("click", () => clearAttachment());

    function clearAttachment() {
      pendingAttachment = null;
      attachPreview.style.display = "none";
      attachName.textContent = "";
      fileInput.value = "";
    }

    fileInput.addEventListener("change", () => {
      const file = fileInput.files[0];
      if (!file) return;
      const isImage = file.type.startsWith("image/");
      const reader = new FileReader();
      reader.onload = (e) => {
        const result = e.target.result;
        if (isImage) {
          // result is a DataURL like "data:image/png;base64,..."
          const b64 = result.split(",")[1];
          pendingAttachment = {type: "image", data: b64, filename: file.name, mime: file.type};
        } else {
          // Text/PDF — send raw bytes as base64, server decodes
          const b64 = result.split(",")[1];
          pendingAttachment = {type: "document", data: b64, filename: file.name};
        }
        attachName.textContent = file.name;
        attachPreview.style.display = "flex";
      };
      reader.readAsDataURL(file);
    });
    // ─────────────────────────────────────────────────────────────────────────

    const RATE_LABELS = {"-75": "Slow", "-50": "Slower", "-25": "Slow-ish", "0": "Normal", "25": "Fast-ish", "50": "Faster", "75": "Fast", "100": "Fastest"};
    function updateRateLabel() {
      ttsRateLabel.textContent = RATE_LABELS[ttsRate.value] || (ttsRate.value > 0 ? "+" + ttsRate.value + "%" : ttsRate.value + "%");
    }
    function updateMixerLabels() {
      const pitch = parseInt(ttsPitch.value || "0", 10);
      const volume = parseInt(ttsVolume.value || "0", 10);
      ttsPitchLabel.textContent = (pitch > 0 ? "+" : "") + pitch;
      ttsVolumeLabel.textContent = (volume > 0 ? "+" : "") + volume + "%";
    }
    function applyTextZoom(value) {
      const zoom = Math.max(85, Math.min(140, parseInt(value || "100", 10) || 100));
      state.textZoom = zoom;
      textZoom.value = String(zoom);
      textZoomLabel.textContent = zoom + "%";
      document.documentElement.style.setProperty("--ui-font-size", (14 * zoom / 100).toFixed(1) + "px");
      saveSettings();
    }
    function applyVoicePauseDelay(value) {
      const delay = Math.max(800, Math.min(4000, parseInt(value || "1800", 10) || 1800));
      state.voicePauseDelay = delay;
      voicePauseDelay.value = String(delay);
      voicePauseLabel.textContent = (delay / 1000).toFixed(1) + "s";
      saveSettings();
    }
    function getVoicePauseDelay() {
      return Math.max(800, Math.min(4000, parseInt(state.voicePauseDelay || voicePauseDelay.value || "1800", 10) || 1800));
    }
    ttsRate.value = localStorage.getItem("ollamaWebTtsRate") || "0";
    ttsPitch.value = localStorage.getItem("ollamaWebTtsPitch") || "0";
    ttsTone.value = localStorage.getItem("ollamaWebTtsTone") || "natural";
    ttsVolume.value = localStorage.getItem("ollamaWebTtsVolume") || "0";
    textZoom.value = String(state.textZoom || 100);
    voicePauseDelay.value = String(state.voicePauseDelay || 1800);
    updateRateLabel();
    updateMixerLabels();
    applyTextZoom(textZoom.value);
    applyVoicePauseDelay(voicePauseDelay.value);
    ttsRate.addEventListener("input", () => { updateRateLabel(); localStorage.setItem("ollamaWebTtsRate", ttsRate.value); });
    ttsPitch.addEventListener("input", () => { updateMixerLabels(); localStorage.setItem("ollamaWebTtsPitch", ttsPitch.value); });
    ttsTone.addEventListener("change", () => { localStorage.setItem("ollamaWebTtsTone", ttsTone.value); });
    ttsVolume.addEventListener("input", () => { updateMixerLabels(); localStorage.setItem("ollamaWebTtsVolume", ttsVolume.value); });
    textZoom.addEventListener("input", () => applyTextZoom(textZoom.value));
    voicePauseDelay.addEventListener("input", () => applyVoicePauseDelay(voicePauseDelay.value));
    const inputHistory = JSON.parse(localStorage.getItem("ollamaWebInputHistory") || "[]");

    // ── TTS (server-side edge-tts neural voices) ─────────────────────────────
    let activeTtsBtn = null;
    let activeAudio = null;
    let activeAudioDone = null;
    let ttsQueue = Promise.resolve();
    let ttsQueueActive = false;
    let ttsQueueSerial = 0;
    let ttsQueueCancelToken = 0;
    let lastAnswerText = "";

    // ── GPU selector ─────────────────────────────────────────────────────────
    async function loadGpus() {
      try {
        const res = await fetch("/api/gpu-list");
        const data = await res.json();
        if (!data.ok || data.gpus.length < 2) return; // hide if only one GPU
        gpuSelect.innerHTML = "";
        for (const g of data.gpus) {
          const opt = document.createElement("option");
          opt.value = g.id;
          opt.textContent = g.label;
          gpuSelect.append(opt);
        }
        const saved = localStorage.getItem("ollamaWebGpu") || "0";
        gpuSelect.value = data.gpus.some(g => g.id === saved) ? saved : data.gpus[0].id;
        state.gpu = gpuSelect.value;
        gpuSelect.style.display = "";
      } catch {}
    }
    loadGpus();
    gpuSelect.addEventListener("change", () => {
      state.gpu = gpuSelect.value;
      localStorage.setItem("ollamaWebGpu", state.gpu);
      loadModels(); // reload models for the selected GPU
    });
    // ── Mic / Whisper STT ─────────────────────────────────────────────────────
    let mediaRecorder = null;
    let audioChunks = [];
    let sttReady = false;
    let voiceConversationActive = false;
    let voiceMonitorTimer = null;
    let voiceAudioContext = null;
    let voiceStream = null;
    function hasBrowserMicApi() {
      return Boolean(window.isSecureContext && navigator.mediaDevices && navigator.mediaDevices.getUserMedia);
    }
    function voiceIsRecording() {
      return Boolean(mediaRecorder && mediaRecorder.state === "recording");
    }
    function stopVoiceMonitor() {
      if (voiceMonitorTimer) {
        clearInterval(voiceMonitorTimer);
        voiceMonitorTimer = null;
      }
      if (voiceAudioContext) {
        voiceAudioContext.close().catch(() => {});
        voiceAudioContext = null;
      }
    }
    function setVoiceConversationActive(on) {
      voiceConversationActive = Boolean(on);
      micBtn.classList.toggle("recording", voiceConversationActive || voiceIsRecording());
      micBtn.textContent = voiceConversationActive
        ? (voiceIsRecording() ? "⏹" : "🎙")
        : (voiceIsRecording() ? "⏹" : "🎙");
      updateMicUi();
    }
    function updateMicUi() {
      const browserMicReady = hasBrowserMicApi();
      micBtn.title = !browserMicReady
        ? "Voice input requires HTTPS or http://localhost in the browser."
        : sttReady
        ? "Voice input: click to start/stop recording. With Auto-send enabled, click once to start/stop hands-free conversation mode."
        : "Voice input unavailable. Click for setup details.";
      voiceAutoSend.checked = Boolean(state.voiceAutoSend);
      voiceAutoWrap.classList.toggle("active", Boolean(state.voiceAutoSend));
      micBtn.classList.toggle("active", Boolean(state.voiceAutoSend) || voiceConversationActive);
      micBtn.classList.toggle("unavailable", !sttReady || !browserMicReady);
    }
    async function checkSttAvailable() {
      try {
        const res = await fetch("/api/stt-status");
        const data = await res.json();
        sttReady = Boolean(data.ok && data.enabled);
      } catch {
        sttReady = false;
      }
      updateMicUi();
    }
    checkSttAvailable();

    voiceAutoSend.addEventListener("change", () => {
      state.voiceAutoSend = voiceAutoSend.checked;
      if (state.voiceAutoSend) {
        state.agentMode = false;
        state.chatMode = "live";
      } else {
        setVoiceConversationActive(false);
        cancelTtsQueue();
        if (voiceIsRecording()) mediaRecorder.stop();
      }
      saveSettings();
      updateMicUi();
      addMessage("Tool", state.voiceAutoSend ? `Voice conversation mode enabled. Click the mic once, speak, and I will auto-send after about ${(getVoicePauseDelay() / 1000).toFixed(1)}s of silence.` : "Voice auto-send disabled.", "tool");
    });

    function scheduleVoiceConversationRestart(delay = 850) {
      if (!voiceConversationActive || !state.voiceAutoSend || currentController || voiceIsRecording() || activeAudio || ttsQueueActive) return;
      setTimeout(() => {
        if (voiceConversationActive && state.voiceAutoSend && !currentController && !voiceIsRecording() && !activeAudio && !ttsQueueActive) {
          startVoiceRecording(true);
        }
      }, delay);
    }

    function setupVoiceAutoStop(stream, startedAt) {
      stopVoiceMonitor();
      const AudioCtx = window.AudioContext || window.webkitAudioContext;
      if (!AudioCtx) {
        setTimeout(() => {
          if (voiceIsRecording()) mediaRecorder.stop();
        }, 6000);
        return;
      }
      voiceAudioContext = new AudioCtx();
      const source = voiceAudioContext.createMediaStreamSource(stream);
      const analyser = voiceAudioContext.createAnalyser();
      analyser.fftSize = 1024;
      source.connect(analyser);
      const data = new Uint8Array(analyser.fftSize);
      let speechStarted = false;
      let silenceSince = null;
      const speechThreshold = 18;
      const silenceThreshold = 10;
      const silenceStopDelay = getVoicePauseDelay();
      voiceMonitorTimer = setInterval(() => {
        if (!voiceIsRecording()) {
          stopVoiceMonitor();
          return;
        }
        analyser.getByteTimeDomainData(data);
        let sum = 0;
        for (const value of data) {
          const centered = value - 128;
          sum += centered * centered;
        }
        const rms = Math.sqrt(sum / data.length);
        const elapsed = Date.now() - startedAt;
        if (rms >= speechThreshold) {
          speechStarted = true;
          silenceSince = null;
        } else if (speechStarted && rms <= silenceThreshold) {
          if (!silenceSince) silenceSince = Date.now();
          if (Date.now() - silenceSince > silenceStopDelay && elapsed > Math.max(1200, silenceStopDelay)) {
            mediaRecorder.stop();
          }
        }
        if (!speechStarted && elapsed > 9000) {
          mediaRecorder.stop();
        }
      }, 120);
    }

    async function startVoiceRecording(autoStop = false) {
      if (voiceIsRecording()) return;
      try {
        voiceStream = await navigator.mediaDevices.getUserMedia({audio: true});
        audioChunks = [];
        mediaRecorder = new MediaRecorder(voiceStream);
        const startedAt = Date.now();
        mediaRecorder.ondataavailable = e => { if (e.data.size > 0) audioChunks.push(e.data); };
        mediaRecorder.onstop = async () => {
          stopVoiceMonitor();
          micBtn.classList.remove("recording");
          micBtn.textContent = voiceConversationActive ? "🎙" : "🎙";
          voiceStream.getTracks().forEach(t => t.stop());
          voiceStream = null;
          const blob = new Blob(audioChunks, {type: mediaRecorder.mimeType || "audio/webm"});
          micBtn.disabled = true;
          micBtn.textContent = "⏳";
          let transcribedText = "";
          try {
            const res = await fetch("/api/transcribe", {
              method: "POST",
              headers: {"Content-Type": blob.type},
              body: blob,
            });
            const data = await res.json();
            if (data.ok && data.text) {
              transcribedText = data.text.trim();
              input.value = (input.value ? input.value + " " : "") + transcribedText;
              input.focus();
              if (state.voiceAutoSend && !currentController) {
                state.agentMode = false;
                state.chatMode = "live";
                updateStatus();
                await sendMessage();
              }
            } else if (data.error) {
              addMessage("Error", "Transcription failed: " + data.error, "error");
            }
          } catch (err) {
            addMessage("Error", "Transcription error: " + err, "error");
          } finally {
            micBtn.disabled = false;
            micBtn.textContent = "🎙";
            updateMicUi();
            if (autoStop && (!currentController || !input.value.trim())) {
              scheduleVoiceConversationRestart(transcribedText ? 850 : 350);
            }
          }
        };
        mediaRecorder.start();
        micBtn.classList.add("recording");
        micBtn.textContent = "⏹";
        updateMicUi();
        if (autoStop) setupVoiceAutoStop(voiceStream, startedAt);
      } catch (err) {
        setVoiceConversationActive(false);
        addMessage("Error", "Microphone access denied: " + err, "error");
      }
    }

    micBtn.addEventListener("click", async () => {
      if (state.agentMode) {
        state.agentMode = false;
        state.chatMode = "live";
        updateStatus();
      }
      if (!hasBrowserMicApi()) {
        addMessage("Error", "Browser microphone access requires a secure context. Open Command Deck at https://... or use http://localhost:8765 on the same machine. Plain http://LAN-IP:8765 will not expose navigator.mediaDevices.", "error");
        return;
      }
      if (!sttReady) {
        addMessage("Error", "Voice input is unavailable. Install faster-whisper or the GPU STT dependencies from requirements-gpu.txt, make sure [stt] enabled = true, then restart.", "error");
        checkSttAvailable();
        return;
      }
      if (state.voiceAutoSend) {
        if (voiceConversationActive) {
          setVoiceConversationActive(false);
          cancelTtsQueue();
          if (voiceIsRecording()) mediaRecorder.stop();
          addMessage("Tool", "Voice conversation mode stopped.", "tool");
        } else {
          state.agentMode = false;
          state.chatMode = "live";
          updateStatus();
          setVoiceConversationActive(true);
          addMessage("Tool", `Voice conversation mode listening. Speak naturally; I will send after about ${(getVoicePauseDelay() / 1000).toFixed(1)}s of silence and read replies aloud.`, "tool");
          await startVoiceRecording(true);
        }
        return;
      }
      if (mediaRecorder && mediaRecorder.state === "recording") {
        mediaRecorder.stop();
        return;
      }
      await startVoiceRecording(false);
    });
    // ─────────────────────────────────────────────────────────────────────────

    async function loadTtsVoices() {
      try {
        const res = await fetch("/api/tts-voices");
        const data = await res.json();
        if (!data.ok || !data.voices?.length) {
          voiceSelect.innerHTML = "<option value=''>No TTS backend available</option>";
          return;
        }
        const saved = localStorage.getItem("ollamaWebTtsVoice");
        voiceSelect.innerHTML = "";
        // Split into local/offline and Edge (online) groups
        const isLocalVoice = v => v.name === "preset:lilith_dark" || v.name.startsWith("af_") || v.name.startsWith("bf_") || (v.label || "").includes("Piper offline");
        const kokoro = data.voices.filter(isLocalVoice);
        const edge   = data.voices.filter(v => !isLocalVoice(v));
        if (kokoro.length) {
          const grp = document.createElement("optgroup");
          grp.label = "── Kokoro (self-hosted / offline) ──";
          for (const v of kokoro) {
            const opt = document.createElement("option");
            opt.value = v.name; opt.textContent = v.label;
            grp.append(opt);
          }
          voiceSelect.append(grp);
        }
        if (edge.length) {
          const grp = document.createElement("optgroup");
          grp.label = "── Edge TTS (internet required) ──";
          for (const v of edge) {
            const opt = document.createElement("option");
            opt.value = v.name; opt.textContent = v.label;
            grp.append(opt);
          }
          voiceSelect.append(grp);
        }
        if (saved && data.voices.some(v => v.name === saved)) {
          voiceSelect.value = saved;
        } else if (data.voices.some(v => v.name === "preset:lilith_dark")) {
          voiceSelect.value = "preset:lilith_dark";
        }
      } catch (e) {
        voiceSelect.innerHTML = "<option value=''>Voice load failed</option>";
      }
    }
    loadTtsVoices();

    voiceSelect.addEventListener("change", () => {
      localStorage.setItem("ollamaWebTtsVoice", voiceSelect.value);
    });

    function stopTts() {
      if (activeAudio) { activeAudio.pause(); activeAudio = null; }
      if (activeAudioDone) { activeAudioDone(); activeAudioDone = null; }
      if (activeTtsBtn) { activeTtsBtn.classList.remove("speaking"); activeTtsBtn.textContent = "🔊"; }
      activeTtsBtn = null;
    }

    function cancelTtsQueue() {
      ttsQueueCancelToken++;
      ttsQueueActive = false;
      ttsQueue = Promise.resolve();
      stopTts();
    }

    // ── Markdown renderer ─────────────────────────────────────────────────────
    function escHtml(s) {
      return s.replace(/&/g,"&amp;").replace(/</g,"&lt;").replace(/>/g,"&gt;");
    }

    function isDiagramText(text) {
      return /[┌┐└┘├┤┬┴┼─│▶→←▲▼]/.test(text) || /\+[-+]{8,}\+/.test(text);
    }

    function normalizeDiagram(text) {
      if (!isDiagramText(text)) return text;
      const lines = text.replace(/\t/g, "    ").split("\n").map(line => line.replace(/\s+$/, ""));
      const nonEmpty = lines.filter(line => line.trim());
      if (!nonEmpty.length) return text;
      const minIndent = Math.min(...nonEmpty.map(line => (line.match(/^\s*/) || [""])[0].length));
      return lines.map(line => line.slice(minIndent)).join("\n");
    }

    function renderMarkdown(text) {
      const lines = text.split("\n");
      let html = "";
      let i = 0;
      while (i < lines.length) {
        const line = lines[i];
        // Fenced code block
        const fenceMatch = line.match(/^```(\w*)/);
        if (fenceMatch) {
          const lang = fenceMatch[1] || "";
          const codeLines = [];
          i++;
          while (i < lines.length && !lines[i].startsWith("```")) {
            codeLines.push(lines[i]);
            i++;
          }
          const codeText = normalizeDiagram(codeLines.join("\n"));
          const diagramClass = isDiagramText(codeText) ? " diagram" : "";
          html += `<div class="code-header"><span>${escHtml(lang)}</span><button type="button" class="copy-btn">Copy</button></div>`;
          html += `<pre class="${diagramClass.trim()}"><code>${escHtml(codeText)}</code></pre>`;
          i++;
          continue;
        }
        // Horizontal rule
        if (/^(---+|\*\*\*+|___+)\s*$/.test(line)) { html += "<hr>"; i++; continue; }
        // Headings
        const h6 = line.match(/^###### (.+)/); if (h6) { html += `<h6>${inlineMd(h6[1])}</h6>`; i++; continue; }
        const h5 = line.match(/^##### (.+)/);  if (h5) { html += `<h5>${inlineMd(h5[1])}</h5>`; i++; continue; }
        const h4 = line.match(/^#### (.+)/);   if (h4) { html += `<h4>${inlineMd(h4[1])}</h4>`; i++; continue; }
        const h3 = line.match(/^### (.+)/);    if (h3) { html += `<h3>${inlineMd(h3[1])}</h3>`; i++; continue; }
        const h2 = line.match(/^## (.+)/);     if (h2) { html += `<h2>${inlineMd(h2[1])}</h2>`; i++; continue; }
        const h1 = line.match(/^# (.+)/);      if (h1) { html += `<h1>${inlineMd(h1[1])}</h1>`; i++; continue; }
        // GFM-style table — header row, separator row, then body rows
        if (line.includes("|") && i + 1 < lines.length && /^\s*\|?\s*[:\-]+(\s*\|\s*[:\-]+)*\s*\|?\s*$/.test(lines[i + 1])) {
          const splitRow = (row) => row.replace(/^\s*\|/, "").replace(/\|\s*$/, "").split("|").map(c => c.trim());
          const aligns = splitRow(lines[i + 1]).map(c => {
            if (/^:-+:$/.test(c)) return "center";
            if (/-+:$/.test(c)) return "right";
            return "left";
          });
          const headers = splitRow(line);
          html += `<div class="table-wrap"><table><thead><tr>`;
          headers.forEach((h, idx) => {
            html += `<th style="text-align:${aligns[idx] || "left"}">${inlineMd(h)}</th>`;
          });
          html += "</tr></thead><tbody>";
          i += 2;
          while (i < lines.length && lines[i].includes("|") && lines[i].trim() !== "") {
            const cells = splitRow(lines[i]);
            html += "<tr>";
            cells.forEach((c, idx) => {
              html += `<td style="text-align:${aligns[idx] || "left"}">${inlineMd(c)}</td>`;
            });
            html += "</tr>";
            i++;
          }
          html += "</tbody></table></div>";
          continue;
        }
        // Blockquote
        if (line.startsWith("> ")) { html += `<blockquote>${inlineMd(line.slice(2))}</blockquote>`; i++; continue; }
        // Unordered list
        if (line.match(/^[-*] /)) {
          html += "<ul>";
          while (i < lines.length && lines[i].match(/^[-*] /)) {
            html += `<li>${inlineMd(lines[i].slice(2))}</li>`;
            i++;
          }
          html += "</ul>";
          continue;
        }
        // Ordered list
        if (line.match(/^\d+\. /)) {
          html += "<ol>";
          while (i < lines.length && lines[i].match(/^\d+\. /)) {
            html += `<li>${inlineMd(lines[i].replace(/^\d+\. /, ""))}</li>`;
            i++;
          }
          html += "</ol>";
          continue;
        }
        // Blank line
        if (line.trim() === "") { html += "<p>"; i++; continue; }
        // Paragraph
        html += `<p>${inlineMd(line)}</p>`;
        i++;
      }
      return html;
    }

    function inlineMd(s) {
      return escHtml(s)
        // Links — must run before bold/italic so the text inside isn't mangled
        .replace(/\[([^\]]+)\]\((https?:\/\/[^)\s]+)\)/g,
          '<a href="$2" target="_blank" rel="noopener">$1</a>')
        // Auto-link bare URLs
        .replace(/(^|[\s(])(https?:\/\/[^\s<)]+)/g,
          '$1<a href="$2" target="_blank" rel="noopener">$2</a>')
        .replace(/\*\*(.+?)\*\*/g, "<strong>$1</strong>")
        .replace(/\*(.+?)\*/g, "<em>$1</em>")
        .replace(/`([^`]+)`/g, "<code>$1</code>");
    }

    async function copyText(text) {
      if (navigator.clipboard && window.isSecureContext) {
        try { await navigator.clipboard.writeText(text); return true; } catch (_) {}
      }
      // Fallback for non-secure contexts and older browsers
      const ta = document.createElement("textarea");
      ta.value = text;
      ta.style.position = "fixed";
      ta.style.opacity = "0";
      document.body.appendChild(ta);
      ta.select();
      let ok = false;
      try { ok = document.execCommand("copy"); } catch (_) {}
      document.body.removeChild(ta);
      return ok;
    }

    // Delegated copy-button handler — works for buttons rendered by streamed markdown
    document.addEventListener("click", async (ev) => {
      const btn = ev.target.closest(".copy-btn");
      if (!btn) return;
      const header = btn.closest(".code-header");
      const pre = header && header.nextElementSibling;
      const code = pre && pre.querySelector("code");
      if (!code) return;
      const ok = await copyText(code.textContent || "");
      btn.textContent = ok ? "Copied!" : "Copy failed";
      setTimeout(() => { btn.textContent = "Copy"; }, 1500);
    });

    // Live render during streaming — safe with incomplete text
    function liveRender(body, text) {
      // If we're inside an open code fence, show plain text to avoid flicker.
      // Preserve the answer-text class when present (streaming reasoning models).
      const isAnswerHolder = body.classList && body.classList.contains("answer-text");
      const fences = (text.match(/^```/gm) || []).length;
      if (fences % 2 !== 0) {
        body.textContent = text;
        if (!isAnswerHolder) body.className = "";
      } else {
        if (!isAnswerHolder) body.className = "msg-body";
        body.innerHTML = renderMarkdown(text);
      }
    }

    // Finalize a streaming body div: render markdown (copy buttons handled by delegation)
    // Preserves any prepended thinking block — only the answer-text region is re-rendered.
    function finalizeBody(body, rawText) {
      body.className = "msg-body";
      const thinking = body.querySelector(".thinking-block");
      let answerHolder = body.querySelector(".answer-text");
      if (!answerHolder) {
        answerHolder = document.createElement("div");
        answerHolder.className = "answer-text";
      }
      answerHolder.innerHTML = renderMarkdown(rawText);
      body.innerHTML = "";
      if (thinking) {
        thinking.open = false;  // collapse on completion
        body.append(thinking);
      }
      body.append(answerHolder);
    }
    // ─────────────────────────────────────────────────────────────────────────

    // ── Canvas popup ──────────────────────────────────────────────────────────
    const canvasWindow = document.getElementById("canvasWindow");
    const canvasBody = document.getElementById("canvasBody");
    const canvasEditor = document.getElementById("canvasEditor");
    const canvasToggleBtn = document.getElementById("canvasToggleBtn");
    const canvasStatus = document.getElementById("canvasStatus");
    const canvasChat = document.getElementById("canvasChat");
    const canvasChatInput = document.getElementById("canvasChatInput");
    const canvasChatSendBtn = document.getElementById("canvasChatSendBtn");
    let canvasMessages = [];
    let canvasController = null;
    let canvasMinimized = false;
    let canvasMaximized = localStorage.getItem("ollamaCanvasMaximized") !== "0";
    let canvasRestoreRect = null;

    function setCanvasStatus(text) {
      if (canvasStatus) canvasStatus.textContent = text;
    }

    function addCanvasChatMessage(role, text, cls) {
      const node = document.createElement("div");
      node.className = "canvas-chat-msg " + (cls || role.toLowerCase());
      const label = document.createElement("span");
      label.className = "canvas-chat-role";
      label.textContent = role + ":";
      const body = document.createElement("div");
      body.className = cls === "ai" ? "msg-body" : "";
      if (cls === "ai") body.innerHTML = renderMarkdown(text || "");
      else body.textContent = text || "";
      node.append(label, body);
      canvasChat.append(node);
      canvasChat.scrollTop = canvasChat.scrollHeight;
      return body;
    }

    function clearCanvasChat() {
      canvasMessages = [];
      canvasChat.innerHTML = "";
      canvasChatInput.value = "";
    }

    function canvasEditRequested(text) {
      return /\b(update|replace|rewrite|revise|reprocess|fix|format|clean|polish|improve|add|insert|integrate|organize)\b/i.test(text);
    }

    function setCanvasOpen(open) {
      canvasWindow.classList.toggle("open", open);
      canvasToggleBtn.classList.toggle("active", open);
      localStorage.setItem("ollamaCanvasOpen", open ? "1" : "0");
      if (open) setCanvasMaximized(canvasMaximized);
      if (!open && appRoot.classList.contains("mobile-tab-canvas")) setMobileTab("chat");
      updateStatus();
    }

    function setMobileTab(tab) {
      const selected = ["chat", "activity", "canvas"].includes(tab) ? tab : "chat";
      appRoot.classList.toggle("mobile-tab-chat", selected === "chat");
      appRoot.classList.toggle("mobile-tab-activity", selected === "activity");
      appRoot.classList.toggle("mobile-tab-canvas", selected === "canvas");
      for (const btn of mobileTabBtns) {
        btn.classList.toggle("active", btn.dataset.mobileTab === selected);
      }
      if (selected === "activity") {
        state.verbose = true;
        updateStatus();
      }
      if (selected === "canvas") {
        setCanvasOpen(true);
        canvasToggleBtn.classList.remove("updated");
      }
      if (selected === "chat") {
        chat.scrollTop = chat.scrollHeight;
      }
    }

    function setCanvasMaximized(on) {
      canvasMaximized = on;
      const btn = document.getElementById("canvasMaxBtn");
      canvasWindow.classList.toggle("maximized", on);
      if (btn) {
        btn.textContent = on ? "❐" : "□";
        btn.title = on ? "Restore canvas window" : "Maximize canvas";
      }
      if (on) {
        canvasWindow.style.left = "";
        canvasWindow.style.top = "";
        canvasWindow.style.right = "";
        canvasWindow.style.bottom = "";
        canvasWindow.style.width = "";
        canvasWindow.style.height = "";
        canvasWindow.style.transform = "";
      } else if (canvasRestoreRect) {
        canvasWindow.style.left = canvasRestoreRect.left + "px";
        canvasWindow.style.top = canvasRestoreRect.top + "px";
        canvasWindow.style.width = canvasRestoreRect.width + "px";
        canvasWindow.style.height = canvasRestoreRect.height + "px";
        canvasWindow.style.right = "auto";
        canvasWindow.style.bottom = "auto";
        canvasWindow.style.transform = "none";
      } else {
        const width = Math.min(980, Math.max(380, window.innerWidth - 48));
        const height = Math.min(760, Math.max(360, window.innerHeight - 48));
        canvasWindow.style.left = Math.max(12, Math.round((window.innerWidth - width) / 2)) + "px";
        canvasWindow.style.top = "24px";
        canvasWindow.style.width = width + "px";
        canvasWindow.style.height = height + "px";
        canvasWindow.style.right = "auto";
        canvasWindow.style.bottom = "auto";
        canvasWindow.style.transform = "none";
      }
      localStorage.setItem("ollamaCanvasMaximized", on ? "1" : "0");
    }

    // Drag-to-move
    (function() {
      const bar = document.getElementById("canvasDragBar");
      let ox = 0, oy = 0, startL = 0, startT = 0;
      bar.addEventListener("mousedown", e => {
        if (e.target.tagName === "BUTTON") return;
        if (canvasMaximized) setCanvasMaximized(false);
        const rect = canvasWindow.getBoundingClientRect();
        ox = e.clientX; oy = e.clientY;
        startL = rect.left; startT = rect.top;
        canvasWindow.style.transform = "none";
        canvasWindow.style.left = startL + "px";
        canvasWindow.style.top  = startT + "px";
        function onMove(e) {
          const maxLeft = Math.max(0, window.innerWidth - canvasWindow.offsetWidth);
          const maxTop = Math.max(0, window.innerHeight - 48);
          const nextLeft = Math.min(maxLeft, Math.max(0, startL + e.clientX - ox));
          const nextTop = Math.min(maxTop, Math.max(0, startT + e.clientY - oy));
          canvasWindow.style.left = nextLeft + "px";
          canvasWindow.style.top  = nextTop + "px";
        }
        function onUp() {
          document.removeEventListener("mousemove", onMove);
          document.removeEventListener("mouseup", onUp);
        }
        document.addEventListener("mousemove", onMove);
        document.addEventListener("mouseup", onUp);
      });
    })();

    document.getElementById("canvasMaxBtn").addEventListener("click", () => {
      if (!canvasMaximized) {
        const rect = canvasWindow.getBoundingClientRect();
        canvasRestoreRect = {left: rect.left, top: rect.top, width: rect.width, height: rect.height};
      }
      setCanvasMaximized(!canvasMaximized);
    });

    document.getElementById("canvasMinBtn").addEventListener("click", () => {
      canvasMinimized = !canvasMinimized;
      canvasBody.style.display = canvasMinimized ? "none" : "flex";
      document.querySelector(".canvas-toolbar").style.display = canvasMinimized ? "none" : "";
      document.getElementById("canvasMinBtn").textContent = canvasMinimized ? "□" : "─";
    });

    document.getElementById("canvasCloseBtn").addEventListener("click", () => setCanvasOpen(false));
    document.getElementById("canvasClearBtn").addEventListener("click", () => {
      canvasEditor.value = "";
      const preview = document.getElementById("canvasPreview");
      if (preview && preview.style.display !== "none") preview.innerHTML = "";
      setCanvasStatus("Canvas cleared");
    });

    // ── Canvas file management ────────────────────────────────────────────────
    const canvasFileSelect = document.getElementById("canvasFileSelect");
    const canvasFileDeleteBtn = document.getElementById("canvasFileDeleteBtn");
    let currentCanvasFile = null;

    function clearCanvasForSession(statusText) {
      canvasEditor.value = "";
      currentCanvasFile = null;
      if (canvasFileSelect) canvasFileSelect.value = "__current__";
      if (canvasFileDeleteBtn) canvasFileDeleteBtn.style.display = "none";
      const preview = document.getElementById("canvasPreview");
      if (preview) preview.innerHTML = "";
      clearCanvasChat();
      refreshCanvasPreview();
      setCanvasStatus(statusText || "Canvas cleared for session");
    }

    function refreshCanvasPreview() {
      if (!canvasPreviewing) return;
      const preview = document.getElementById("canvasPreview");
      if (preview) preview.innerHTML = renderMarkdown(canvasEditor.value);
    }
    async function refreshCanvasFileList(selectName) {
      try {
        const res = await fetch("/api/canvas-files");
        const data = await res.json();
        const files = (data && data.files) || [];
        canvasFileSelect.innerHTML = '<option value="__current__">📝 Untitled</option>';
        for (const f of files) {
          const opt = document.createElement("option");
          opt.value = f.name;
          const date = (f.updated_at || "").replace("T", " ").slice(0, 16);
          opt.textContent = (date ? date + " · " : "") + f.name;
          opt.title = f.name;
          canvasFileSelect.append(opt);
        }
        if (selectName && files.some(f => f.name === selectName)) {
          canvasFileSelect.value = selectName;
          currentCanvasFile = selectName;
        } else if (!currentCanvasFile) {
          canvasFileSelect.value = "__current__";
        }
        canvasFileDeleteBtn.style.display = currentCanvasFile ? "" : "none";
      } catch (_) {}
    }
    canvasFileSelect.addEventListener("change", async () => {
      const name = canvasFileSelect.value;
      if (name === "__current__") {
        currentCanvasFile = null;
        canvasFileDeleteBtn.style.display = "none";
        return;
      }
      try {
        const res = await fetch("/api/canvas-files?name=" + encodeURIComponent(name));
        const data = await res.json();
        if (data.ok) {
          canvasEditor.value = data.content || "";
          currentCanvasFile = data.name;
          canvasFileDeleteBtn.style.display = "";
          clearCanvasChat();
          refreshCanvasPreview();
          canvasEditor.scrollTop = 0;
          setCanvasStatus(`Loaded ${data.name}.md`);
        }
      } catch (_) {}
    });
    document.getElementById("canvasNewBtn").addEventListener("click", () => {
      canvasEditor.value = "";
      currentCanvasFile = null;
      canvasFileSelect.value = "__current__";
      canvasFileDeleteBtn.style.display = "none";
      clearCanvasChat();
      refreshCanvasPreview();
      setCanvasStatus("New unsaved canvas");
    });
    document.getElementById("canvasSaveBtn").addEventListener("click", async () => {
      const defaultName = currentCanvasFile || "";
      const name = (prompt("Save canvas as (letters, numbers, dot, dash, underscore):", defaultName) || "").trim();
      if (!name) return;
      try {
        const res = await fetch("/api/canvas-files/save", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({name, content: canvasEditor.value}),
        });
        const data = await res.json();
        if (!data.ok) {
          alert("Save failed: " + (data.error || "unknown"));
          return;
        }
        currentCanvasFile = data.file.name;
        await refreshCanvasFileList(currentCanvasFile);
        setCanvasStatus(`Saved ${data.file.name}.md`);
      } catch (e) { alert("Save failed: " + e); }
    });
    async function saveCurrentCanvasIfNamed(statusPrefix) {
      if (!currentCanvasFile) return;
      try {
        step(`Saving canvas file ${currentCanvasFile}.md`);
        const res = await fetch("/api/canvas-files/save", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({name: currentCanvasFile, content: canvasEditor.value}),
        });
        const data = await res.json();
        if (data.ok) {
          await refreshCanvasFileList(data.file.name);
          setCanvasStatus(`${statusPrefix || "Saved"} ${data.file.name}.md`);
          step(`Saved canvas file ${data.file.name}.md`);
        } else {
          setCanvasStatus("Canvas updated, save failed");
          step("Canvas file save failed");
        }
      } catch (_) {
        setCanvasStatus("Canvas updated, save failed");
        step("Canvas file save failed");
      }
    }
    canvasFileDeleteBtn.addEventListener("click", async () => {
      if (!currentCanvasFile) return;
      if (!confirm("Delete saved canvas '" + currentCanvasFile + "'?")) return;
      try {
        await fetch("/api/canvas-files/delete", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({name: currentCanvasFile}),
        });
        currentCanvasFile = null;
        canvasEditor.value = "";
        canvasFileDeleteBtn.style.display = "none";
        clearCanvasChat();
        refreshCanvasPreview();
        await refreshCanvasFileList();
        setCanvasStatus("Deleted saved canvas");
      } catch (_) {}
    });
    refreshCanvasFileList();
    canvasEditor.addEventListener("input", () => {
      refreshCanvasPreview();
      setCanvasStatus(currentCanvasFile ? `Editing ${currentCanvasFile}.md` : "Unsaved canvas");
    });

    // Edit/Preview toggle
    let canvasPreviewing = false;
    const canvasViewBtn = document.getElementById("canvasViewBtn");
    function setCanvasPreviewing(on) {
      canvasPreviewing = on;
      const preview = document.getElementById("canvasPreview");
      if (canvasPreviewing) {
        preview.innerHTML = renderMarkdown(canvasEditor.value);
        preview.style.display = "";
        canvasEditor.style.display = "none";
        canvasViewBtn.textContent = "✏ Edit";
        canvasViewBtn.classList.add("active");
        setCanvasStatus(currentCanvasFile ? `Previewing ${currentCanvasFile}.md` : "Previewing unsaved canvas");
      } else {
        preview.style.display = "none";
        canvasEditor.style.display = "";
        canvasViewBtn.textContent = "👁 Preview";
        canvasViewBtn.classList.remove("active");
        setCanvasStatus(currentCanvasFile ? `Editing ${currentCanvasFile}.md` : "Unsaved canvas");
      }
    }
    canvasViewBtn.addEventListener("click", () => setCanvasPreviewing(!canvasPreviewing));
    document.getElementById("canvasCopyBtn").addEventListener("click", async () => {
      const btn = document.getElementById("canvasCopyBtn");
      const ok = await copyText(canvasEditor.value);
      btn.textContent = ok ? "Copied!" : "Copy failed";
      setTimeout(() => { btn.textContent = "Copy"; }, 1500);
    });

    // Send canvas content to AI as the working document.
    document.getElementById("canvasSendAiBtn").addEventListener("click", () => {
      const content = canvasEditor.value.trim();
      if (!content) return;
      if (canvasController || currentController) return;
      const extraInstruction = input.value.trim();
      const instruction = extraInstruction
        ? `Reprocess the current canvas and replace the canvas with the complete updated Markdown document. Preserve the existing document, integrate any new section I typed into it, and apply this instruction:\n\n${extraInstruction}`
        : "Reprocess the current canvas and replace the canvas with the complete updated Markdown document. Preserve the existing document, integrate any new section I typed into it, clean up structure, and improve formatting.";
      input.value = "";
      setCanvasStatus("Sending canvas to AI...");
      sendCanvasMessage(instruction, true);
    });

    // Export functions
    function downloadFile(content, filename, mime) {
      const blob = new Blob([content], {type: mime});
      const a = document.createElement("a");
      a.href = URL.createObjectURL(blob);
      a.download = filename;
      a.click(); URL.revokeObjectURL(a.href);
    }

    function toRtf(text) {
      const esc = text
        .replace(/\\/g, "\\\\").replace(/\{/g, "\\{").replace(/\}/g, "\\}")
        .replace(/\r?\n/g, "\\par\n");
      return "{\\rtf1\\ansi\\deff0{\\fonttbl{\\f0\\froman\\fcharset0 Times New Roman;}}" +
             "{\\colortbl ;\\red232\\green234\\blue237;}\\f0\\fs24\\cf1 " + esc + "}";
    }

    function toCsv(text) {
      return text.split("\n").map(line => '"' + line.replace(/"/g, '""') + '"').join("\n");
    }

    function exportPdf(text) {
      const w = window.open("", "_blank");
      const rendered = renderMarkdown(text);
      w.document.write(`<!doctype html><html><head><meta charset="utf-8">
        <title>Canvas Export</title>
        <style>
          body{font:12pt/1.6 Georgia,"Times New Roman",serif;max-width:7.5in;margin:0.6in auto;padding:0 0.4in;color:#111;}
          h1,h2,h3,h4,h5,h6{font-family:"Helvetica Neue",Arial,sans-serif;color:#111;line-height:1.25;margin:1.4em 0 0.5em;font-weight:600;}
          h1{font-size:24pt;border-bottom:2px solid #ddd;padding-bottom:6px;margin-top:0;}
          h2{font-size:16pt;border-bottom:1px solid #e5e5e5;padding-bottom:4px;}
          h3{font-size:13pt;}
          h4{font-size:11pt;text-transform:uppercase;letter-spacing:0.5px;color:#444;}
          p{margin:0.5em 0;}
          ul,ol{margin:0.4em 0 0.7em;padding-left:0.4in;}
          li{margin:0.15em 0;}
          a{color:#1a5fb4;text-decoration:none;}
          blockquote{margin:1em 0;padding:6px 14px;border-left:3px solid #c7c7c7;background:#f4f3ee;color:#444;font-style:italic;}
          hr{border:none;border-top:1px solid #d0d0d0;margin:1.4em 0;}
          code{font-family:"SFMono-Regular",Consolas,monospace;font-size:0.9em;background:#ececec;padding:1px 4px;border-radius:3px;}
          pre{background:#f1f1ef;border:1px solid #e0e0e0;border-radius:4px;padding:10px 12px;overflow-x:auto;line-height:1.45;page-break-inside:avoid;}
          pre code{background:none;padding:0;}
          .code-header{display:none;}
          table{border-collapse:collapse;margin:1em 0;width:100%;font-size:10.5pt;page-break-inside:avoid;}
          th,td{border:1px solid #c7c7c7;padding:5px 9px;vertical-align:top;}
          thead th{background:#efefe9;font-weight:600;}
          tbody tr:nth-child(even){background:#f7f6f1;}
          @page{margin:0.75in;}
          @media print{body{margin:0;max-width:none;}}
        </style></head><body>${rendered}</body></html>`);
      w.document.close();
      w.focus();
      setTimeout(() => { w.print(); }, 400);
    }

    document.querySelectorAll(".export-btn").forEach(btn => {
      btn.addEventListener("click", () => {
        const fmt = btn.dataset.fmt;
        const text = canvasEditor.value;
        const base = "document";
        if (fmt === "md")  downloadFile(text, base + ".md",  "text/markdown");
        if (fmt === "txt") downloadFile(text, base + ".txt", "text/plain");
        if (fmt === "rtf") downloadFile(toRtf(text), base + ".rtf", "application/rtf");
        if (fmt === "csv") downloadFile(toCsv(text), base + ".csv", "text/csv");
        if (fmt === "pdf") exportPdf(text);
      });
    });

    function writeCanvas(text, mode) {
      // mode: "replace" overwrites, "append" appends with separator
      const value = String(text || "");
      if (mode === "replace") {
        canvasEditor.value = value;
      } else {
        if (canvasEditor.value && !canvasEditor.value.endsWith("\n\n")) {
          canvasEditor.value += canvasEditor.value.endsWith("\n") ? "\n" : "\n\n";
        }
        canvasEditor.value += value;
      }
      setCanvasOpen(true);
      if (canvasMinimized) {
        canvasMinimized = false;
        canvasBody.style.display = "flex";
        document.querySelector(".canvas-toolbar").style.display = "";
        document.getElementById("canvasMinBtn").textContent = "─";
      }
      if (canvasPreviewing) {
        const preview = document.getElementById("canvasPreview");
        if (preview) {
          preview.innerHTML = renderMarkdown(canvasEditor.value);
          preview.scrollTop = mode === "replace" ? 0 : preview.scrollHeight;
        }
      } else {
        canvasEditor.style.display = "";
        canvasEditor.scrollTop = mode === "replace" ? 0 : canvasEditor.scrollHeight;
      }
      canvasEditor.focus();
      canvasToggleBtn.classList.add("updated");
      setCanvasStatus(mode === "replace" ? "Canvas replaced from chat" : "Appended from chat");
      saveCurrentCanvasIfNamed(mode === "replace" ? "Replaced and saved" : "Appended and saved");
    }
    function appendToCanvas(text) { writeCanvas(text, "append"); }

    async function sendCanvasMessage(instruction, forceReplace) {
      const canvasContent = canvasEditor.value.trim();
      const text = String(instruction || "").trim();
      if (!canvasContent || !text || canvasController) return;
      const shouldReplace = Boolean(forceReplace || canvasEditRequested(text));
      canvasMessages.push({role: "user", content: text});
      canvasMessages = canvasMessages.slice(-12);
      addCanvasChatMessage("You", text, "user");
      const aiBody = addCanvasChatMessage("Canvas AI", "", "ai");
      setCanvasStatus(shouldReplace ? "Canvas AI updating document..." : "Canvas AI reading document...");
      canvasController = new AbortController();
      canvasChatSendBtn.disabled = true;
      try {
        const payload = {
          ...state,
          messages: canvasMessages,
          text,
          attachment: null,
          canvasContent,
          canvasRequested: shouldReplace,
          agentMode: false,
          gpu: state.gpu,
        };
        const res = await fetch("/api/chat-stream", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify(payload),
          signal: canvasController.signal,
        });
        if (!res.ok || !res.body) throw new Error("HTTP " + res.status);
        const reader = res.body.getReader();
        const decoder = new TextDecoder();
        let buffer = "";
        let answerText = "";
        const handleCanvasEvent = (event) => {
          if (event.type === "error") {
            addCanvasChatMessage("Error", event.error || "Canvas request failed", "tool");
          } else if (event.type === "cmd") {
            if (state.verbose) addCanvasChatMessage("step", event.command || "", "tool");
          } else if (event.type === "context") {
            if (event.profile) {
              state.activeAgentProfile = event.profile;
              updateStatus();
            }
            if (state.verbose) addCanvasChatMessage("context", event.text || "", "tool");
          } else if (event.type === "thinking") {
            if (state.verbose) setCanvasStatus("Canvas AI thinking...");
          } else if (event.type === "chunk") {
            answerText += event.text || "";
            aiBody.innerHTML = renderMarkdown(answerText);
            canvasChat.scrollTop = canvasChat.scrollHeight;
          } else if (event.type === "final") {
            if (event.text) {
              answerText = event.text;
              aiBody.innerHTML = renderMarkdown(answerText);
            }
          }
        };
        while (true) {
          const {value, done} = await reader.read();
          if (done) break;
          buffer += decoder.decode(value, {stream: true});
          const parts = buffer.split("\n\n");
          buffer = parts.pop() || "";
          for (const part of parts) {
            const line = part.split("\n").find(item => item.startsWith("data: "));
            if (line) handleCanvasEvent(JSON.parse(line.slice(6)));
          }
        }
        canvasMessages.push({role: "assistant", content: answerText});
        canvasMessages = canvasMessages.slice(-12);
        if (shouldReplace && answerText.trim()) {
          writeCanvas(answerText.trim(), "replace");
          setCanvasStatus("Canvas updated from canvas chat");
        } else {
          setCanvasStatus("Canvas chat answered");
        }
      } catch (err) {
        if (err.name !== "AbortError") addCanvasChatMessage("Error", String(err), "tool");
        setCanvasStatus(err.name === "AbortError" ? "Canvas chat stopped" : "Canvas chat failed");
      } finally {
        canvasController = null;
        canvasChatSendBtn.disabled = false;
      }
    }

    canvasChatSendBtn.addEventListener("click", () => {
      const instruction = canvasChatInput.value.trim();
      if (!instruction) return;
      canvasChatInput.value = "";
      sendCanvasMessage(instruction, false);
    });
    canvasChatInput.addEventListener("keydown", (event) => {
      if (event.key === "Enter" && !event.shiftKey) {
        event.preventDefault();
        canvasChatSendBtn.click();
      }
    });

    let autoCanvas = false;
    const canvasAutoBtn = document.getElementById("canvasAutoBtn");
    function setAutoCanvas(on) {
      autoCanvas = false;
      canvasAutoBtn.textContent = "Canvas: Manual";
      canvasAutoBtn.classList.remove("active");
      canvasAutoBtn.disabled = true;
      canvasAutoBtn.title = "Canvas is written only when your prompt explicitly mentions canvas.";
      localStorage.removeItem("ollamaAutoCanvas");
    }
    setAutoCanvas(false);

    canvasToggleBtn.addEventListener("click", () => {
      setCanvasOpen(!canvasWindow.classList.contains("open"));
      canvasToggleBtn.classList.remove("updated");
    });

    // Only write to canvas when the user explicitly mentions "canvas".
    const CANVAS_KEYWORDS = /\bcanvas\b/i;
    const CANVAS_REPLACE_KEYWORDS = /\b(post|write|put|save|create|replace|set|make|update|reprocess|rewrite|revise|improve|clean|format)\b[^.]{0,80}\b(in|into|to|on|the|current)?\s*canvas\b|\bcanvas\b[^.]{0,50}\b(should|with|to|and)\b[^.]{0,50}\b(be|contain|show|have|replace|update)\b|\breplace\s+the\s+canvas\b/i;
    function shouldAutoCanvas(userText) {
      return CANVAS_KEYWORDS.test(userText);
    }
    function canvasModeFor(userText) {
      return CANVAS_KEYWORDS.test(userText) ? "replace" : (CANVAS_REPLACE_KEYWORDS.test(userText) ? "replace" : "append");
    }

    if (localStorage.getItem("ollamaCanvasOpen") === "1") setCanvasOpen(true);
    // ─────────────────────────────────────────────────────────────────────────

    function addSpeakButton(body, rawText) {
      const label = body.previousElementSibling;
      if (!label) return;
      const btn = document.createElement("button");
      btn.type = "button";
      btn.className = "tts-btn";
      btn.title = "Read aloud";
      btn.textContent = "🔊";
      btn.addEventListener("click", () => speakText(rawText || body.textContent, btn));
      label.append(btn);
      // Canvas append button
      const canvasBtn = document.createElement("button");
      canvasBtn.type = "button";
      canvasBtn.className = "tts-btn";
      canvasBtn.title = "Append to canvas";
      canvasBtn.textContent = "→ Canvas";
      canvasBtn.style.fontSize = "11px";
      canvasBtn.addEventListener("click", () => appendToCanvas(rawText || body.textContent));
      label.append(canvasBtn);
      const replaceCanvasBtn = document.createElement("button");
      replaceCanvasBtn.type = "button";
      replaceCanvasBtn.className = "tts-btn";
      replaceCanvasBtn.title = "Replace canvas with this response";
      replaceCanvasBtn.textContent = "↻ Canvas";
      replaceCanvasBtn.style.fontSize = "11px";
      replaceCanvasBtn.addEventListener("click", () => writeCanvas(rawText || body.textContent, "replace"));
      label.append(replaceCanvasBtn);
    }

    function cleanTextForTts(text) {
      const lines = String(text || "")
        .replace(/```[\s\S]*?```/g, " Code block omitted. ")
        .replace(/`([^`]+)`/g, "$1")
        .replace(/!\[([^\]]*)\]\([^)]+\)/g, "$1")
        .replace(/\[([^\]]+)\]\([^)]+\)/g, "$1")
        .split(/\r?\n/);
      const out = [];
      for (let line of lines) {
        line = line.trim();
        line = line.replace(/[✅❌⚠️✔✖✗✓]/g, "");
        if (!line) {
          out.push("");
          continue;
        }
        if (/^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$/.test(line)) continue;
        if (line.includes("|") && (line.match(/\|/g) || []).length >= 2) {
          const cells = line.split("|").map(c => c.trim()).filter(Boolean);
          if (cells.length) out.push(cells.join("; ") + ".");
          continue;
        }
        line = line
          .replace(/^#{1,6}\s*/, "")
          .replace(/^\s*[-*+]\s+/, "")
          .replace(/^\s*\d+[.)]\s+/, "")
          .replace(/^\s*>\s?/, "")
          .replace(/[*_~]{1,3}/g, "")
          .replace(/\s{2,}/g, " ")
          .trim();
        if (line) out.push(line);
      }
      return out.join("\n").replace(/\n{3,}/g, "\n\n").replace(/[ \t]+/g, " ").trim();
    }

    async function speakText(text, btn, options = {}) {
      if (options.stopExisting !== false) stopTts();
      text = options.alreadyClean ? String(text || "").trim() : cleanTextForTts(text);
      if (!text) return;
      const voice = voiceSelect.value || "preset:lilith_dark";
      const rate = parseInt(ttsRate.value, 10);
      const rateStr = (rate >= 0 ? "+" : "") + rate + "%";
      const pitch = parseInt(ttsPitch.value || "0", 10);
      const tone = ttsTone.value || "natural";
      const volume = parseInt(ttsVolume.value || "0", 10);
      activeTtsBtn = btn || null;
      if (activeTtsBtn) { activeTtsBtn.classList.add("speaking"); activeTtsBtn.textContent = "⏸"; }
      try {
        const res = await fetch("/api/tts", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({text, voice, rate: rateStr, pitch, tone, volume})
        });
        if (!res.ok) {
          let detail = "";
          try {
            const err = await res.json();
            detail = err.error ? ": " + err.error : "";
          } catch (_) {}
          throw new Error("TTS server error " + res.status + detail);
        }
        const blob = await res.blob();
        const url = URL.createObjectURL(blob);
        const audio = new Audio(url);
        activeAudio = audio;
        await new Promise((resolve, reject) => {
          let settled = false;
          const cleanup = () => {
            if (settled) return;
            settled = true;
            URL.revokeObjectURL(url);
            if (activeTtsBtn) { activeTtsBtn.classList.remove("speaking"); activeTtsBtn.textContent = "🔊"; }
            activeTtsBtn = null;
            activeAudio = null;
            activeAudioDone = null;
          };
          activeAudioDone = () => {
            cleanup();
            resolve();
          };
          audio.onended = () => {
            cleanup();
            resolve();
          };
          audio.onerror = () => {
            cleanup();
            reject(new Error("audio playback failed"));
          };
          audio.play().catch(err => {
            cleanup();
            reject(err);
          });
        });
      } catch (err) {
        if (activeTtsBtn) { activeTtsBtn.classList.remove("speaking"); activeTtsBtn.textContent = "🔊"; }
        activeAudio = null;
        activeAudioDone = null;
        activeTtsBtn = null;
        addMessage("Error", "TTS failed: " + err, "error");
      }
    }

    function queueSpeakText(text) {
      const cleaned = cleanTextForTts(text);
      if (!cleaned) return ttsQueue;
      ttsQueueActive = true;
      const serial = ++ttsQueueSerial;
      const cancelToken = ttsQueueCancelToken;
      stopBtn.classList.add("active");
      const next = ttsQueue
        .catch(() => {})
        .then(() => {
          if (cancelToken !== ttsQueueCancelToken || !voiceConversationActive || !state.voiceAutoSend) return;
          return speakText(cleaned, null, {stopExisting: false, alreadyClean: true});
        })
        .catch(err => addMessage("Error", "TTS failed: " + err, "error"));
      ttsQueue = next.finally(() => {
        if (ttsQueueSerial === serial) {
          ttsQueueActive = false;
          if (!currentController) stopBtn.classList.remove("active");
        }
      });
      return ttsQueue;
    }

    function takeSpeakableSegments(buffer, force = false) {
      const cleaned = cleanTextForTts(buffer);
      if (!cleaned) return {segments: [], rest: ""};
      if (force) return {segments: [cleaned], rest: ""};
      const segments = [];
      let rest = cleaned;
      const sentenceEnd = /[.!?](?:["')\]]+)?\s+/g;
      let match;
      let cut = 0;
      while ((match = sentenceEnd.exec(cleaned)) !== null) {
        const end = match.index + match[0].length;
        const segment = cleaned.slice(cut, end).trim();
        if (segment.length >= 24) {
          segments.push(segment);
          cut = end;
        }
      }
      if (cut > 0) rest = cleaned.slice(cut).trim();
      if (!segments.length && cleaned.length >= 240) {
        const comma = cleaned.lastIndexOf(",", 220);
        const space = cleaned.lastIndexOf(" ", 220);
        const splitAt = comma > 120 ? comma + 1 : (space > 120 ? space : 220);
        segments.push(cleaned.slice(0, splitAt).trim());
        rest = cleaned.slice(splitAt).trim();
      }
      return {segments, rest};
    }

    ttsStopBtn.addEventListener("click", stopTts);

    readLastBtn.addEventListener("click", () => {
      if (lastAnswerText) speakText(lastAnswerText, readLastBtn);
    });
    // ─────────────────────────────────────────────────────────────────────────
    let historyIndex = -1;
    let draftBeforeHistory = "";

    function updateStatus() {
      state.assistantName = cleanAssistantName(state.assistantName);
      document.getElementById("modelLabel").textContent = state.model || "none";
      document.getElementById("verboseLabel").textContent = state.verbose ? "ON" : "OFF";
      appRoot.classList.toggle("verbose-on", state.verbose);
      const modeNames = {chat:"Chat/Voice", live:"Live Chat", agent:"Agent", coding:"Coding", creative:"Creative", concise:"Concise", teaching:"Teaching"};
      const modeKey = state.agentMode ? "agent" : (state.chatMode === "default" || state.chatMode === "conversation" ? "chat" : state.chatMode);
      document.getElementById("modeLabel").textContent = modeNames[modeKey] || modeKey;
      document.getElementById("personalityLabel").textContent = state.personality;
      const activeProfile = state.activeAgentProfile || state.agentProfile;
      document.getElementById("agentProfileLabel").textContent = (
        state.autoProfile && activeProfile !== state.agentProfile
          ? `${state.agentProfile} -> ${activeProfile}`
          : state.agentProfile
      );
      agentProfile.value = state.agentProfile;
      autoProfile.checked = Boolean(state.autoProfile);
      keepAlive.checked = Boolean(state.keepAlive);
      assistantNameInput.value = state.assistantName;
      assistantNameTitle.textContent = state.assistantName;
      verboseBtn.classList.toggle("active", state.verbose);
      if (mobileActivityBtn) mobileActivityBtn.classList.toggle("active", state.verbose);
      if (mobileCanvasBtn) mobileCanvasBtn.classList.toggle("active", canvasWindow?.classList.contains("open"));
      syncModeSelect();
      saveSettings();
    }

    const RENDER_CLASSES = new Set(["ollama", "agent"]);
    const VERBOSE_ROLES = new Set(["tool", "cmd", "step", "stats", "context"]);

    function addVerboseMessage(role, text, cls) {
      const node = document.createElement("section");
      const kind = cls || String(role || "tool").toLowerCase();
      node.className = "vmsg " + kind;
      const label = document.createElement("span");
      label.className = "role";
      label.textContent = role + ":";
      const body = document.createElement("div");
      body.textContent = text || "";
      node.append(label, body);
      verboseLog.append(node);
      verboseLog.scrollTop = verboseLog.scrollHeight;
      return node;
    }

    if (verboseClearBtn) {
      verboseClearBtn.addEventListener("click", () => {
        verboseLog.innerHTML = "";
      });
    }

    function addMessage(role, text, cls) {
      const node = document.createElement("section");
      const resolvedCls = cls || role.toLowerCase();
      if (VERBOSE_ROLES.has(resolvedCls)) {
        const verboseNode = addVerboseMessage(role, text, String(role || resolvedCls).toLowerCase());
        return verboseNode.querySelector("div");
      }
      node.className = "msg " + resolvedCls;
      const label = document.createElement("span");
      label.className = "role";
      label.textContent = role + ":";
      const body = document.createElement("div");
      if (RENDER_CLASSES.has(resolvedCls) && text) {
        finalizeBody(body, text);
        node.append(label, body);
        chat.append(node);
        chat.scrollTop = chat.scrollHeight;
        addSpeakButton(body, text);
      } else {
        body.textContent = text || "";
        node.append(label, body);
        chat.append(node);
        chat.scrollTop = chat.scrollHeight;
      }
      return body;
    }

    function createMessage(role, cls) {
      const node = document.createElement("section");
      node.className = "msg " + (cls || role.toLowerCase());
      const label = document.createElement("span");
      label.className = "role";
      label.textContent = role + ":";
      const body = document.createElement("div");
      node.append(label, body);
      chat.append(node);
      chat.scrollTop = chat.scrollHeight;
      return body;
    }

    function applyResponseState(data) {
      if (data.mode === "agent") { state.agentMode = true; syncModeSelect(); }
      if (data.mode === "chat") { state.agentMode = false; syncModeSelect(); }
      if (data.chatMode) { state.chatMode = data.chatMode; syncModeSelect(); }
      if (typeof data.verbose === "boolean") state.verbose = data.verbose;
      if (data.models) {
        const current = data.model || state.model;
        modelSelect.innerHTML = "";
        for (const model of data.models) {
          const opt = document.createElement("option");
          opt.value = model;
          opt.textContent = model;
          modelSelect.append(opt);
        }
        state.model = current;
        modelSelect.value = current;
      }
      if (data.model) {
        state.model = data.model;
        if (![...modelSelect.options].some(opt => opt.value === data.model)) {
          const opt = document.createElement("option");
          opt.value = data.model;
          opt.textContent = data.model;
          modelSelect.append(opt);
        }
        modelSelect.value = data.model;
      }
      if (data.chatMode) {
        state.chatMode = data.chatMode;
        syncModeSelect();
      }
      if (data.personality) {
        state.personality = data.personality;
        personality.value = data.personality;
      }
      if (data.agentProfile) {
        state.agentProfile = data.agentProfile;
        state.activeAgentProfile = data.agentProfile;
        agentProfile.value = data.agentProfile;
      }
      updateStatus();
    }

    const CAP_ICONS = {chat:"💬", agent:"🤖", vision:"🖼", code:"💻", think:"🧠", document:"📄"};

    async function loadModels() {
      const gpu = state.gpu || "0";
      const res = await fetch("/api/models?gpu=" + encodeURIComponent(gpu));
      const data = await res.json();
      modelSelect.innerHTML = "";
      if (!data.ok) {
        addMessage("Error", data.error, "error");
        return;
      }
      for (const model of data.models) {
        const opt = document.createElement("option");
        opt.value = model.name;
        const caps = (model.capabilities || []);
        // Show capability icons after the model name
        const tagStr = caps.map(c => CAP_ICONS[c] || c).join(" ");
        opt.textContent = tagStr ? model.name + "  " + tagStr : model.name;
        opt.title = caps.length ? "Supports: " + caps.join(", ") : "";
        modelSelect.append(opt);
      }
      const preferred = ["qwen3.5:latest", "qwen3:latest", "qwen3.5", "qwen3"];
      const savedModel = state.model && data.models.some(m => m.name === state.model) ? state.model : "";
      const found = preferred.find(p => data.models.some(m => m.name === p));
      state.model = savedModel || found || data.models[0]?.name || "";
      modelSelect.value = state.model;
      updateStatus();
      refreshContextWindow();
    }

    // ── Context window indicator ──────────────────────────────────────────────
    let contextLength = null;  // runtime/configured context tokens, null until unknown
    let advertisedContextLength = null;  // model architecture max context tokens
    let lastPromptTokens = 0;
    async function refreshContextWindow() {
      contextLength = null;
      advertisedContextLength = null;
      renderContextBar();
      if (!state.model) return;
      try {
        const res = await fetch("/api/model-info?model=" + encodeURIComponent(state.model)
          + "&gpu=" + encodeURIComponent(state.gpu || "0"));
        const data = await res.json();
        if (data.ok) {
          contextLength = data.runtime_context_length || null;
          advertisedContextLength = data.advertised_context_length || data.context_length || null;
          renderContextBar();
        }
      } catch (_) {}
    }
    function renderContextBar() {
      const label = document.getElementById("ctxLabel");
      const fill = document.getElementById("ctxFill");
      const fmt = (n) => n >= 1000 ? (n / 1000).toFixed(n >= 10000 ? 0 : 1) + "k" : String(n);
      const maxForBar = contextLength || advertisedContextLength;
      if (!maxForBar) {
        label.textContent = lastPromptTokens ? fmt(lastPromptTokens) + " / runtime ?" : "—";
        fill.style.width = "0%";
        fill.classList.remove("warn", "danger");
        document.getElementById("ctxBox").title = "Last prompt tokens / runtime context unknown";
        return;
      }
      const pct = Math.min(100, (lastPromptTokens / maxForBar) * 100);
      const runtimeText = contextLength ? fmt(contextLength) + " runtime" : "runtime ?";
      const advertisedText = advertisedContextLength ? " · " + fmt(advertisedContextLength) + " advertised" : "";
      label.textContent = fmt(lastPromptTokens) + " / " + runtimeText + advertisedText
        + " (" + Math.round(pct) + "%)";
      document.getElementById("ctxBox").title = "Prompt tokens / runtime context"
        + (advertisedContextLength ? " · advertised model max " + advertisedContextLength.toLocaleString() : "");
      fill.style.width = pct + "%";
      fill.classList.toggle("warn", pct >= 70 && pct < 90);
      fill.classList.toggle("danger", pct >= 90);
    }

    // ── Sessions ──────────────────────────────────────────────────────────────
    const sessionSelect = document.getElementById("sessionSelect");
    const sessionDeleteBtn = document.getElementById("sessionDeleteBtn");
    let currentSessionId = null;
    let saveTimer = null;

    function fmtSessionLabel(s) {
      const date = (s.updated_at || s.created_at || "").replace("T", " ").replace("Z", "");
      const title = (s.title || "(untitled)").slice(0, 40);
      return (date ? date.slice(0, 16) + " · " : "") + title;
    }
    async function refreshSessionList(selectId) {
      try {
        const res = await fetch("/api/sessions");
        const data = await res.json();
        const sessions = (data && data.sessions) || [];
        sessionSelect.innerHTML = '<option value="__new__">＋ New session</option>';
        for (const s of sessions) {
          const opt = document.createElement("option");
          opt.value = s.id;
          opt.textContent = fmtSessionLabel(s);
          opt.title = s.title || "";
          sessionSelect.append(opt);
        }
        if (selectId && sessions.some(s => s.id === selectId)) {
          sessionSelect.value = selectId;
          currentSessionId = selectId;
        }
        sessionDeleteBtn.style.display = currentSessionId ? "" : "none";
      } catch (_) {}
    }
    function rerenderChatFromMessages(messages) {
      chat.innerHTML = "";
      verboseLog.innerHTML = "";
      addMessage("Tool", "Web TUI ready. Select a model and start chatting.", "tool");
      for (const m of messages) {
        if (m.role === "user") addMessage("You", m.content, "user");
        else if (m.role === "assistant") addMessage(cleanAssistantName(state.assistantName), m.content, "ollama");
      }
    }
    async function loadSessionById(id) {
      try {
        const res = await fetch("/api/sessions?id=" + encodeURIComponent(id));
        const data = await res.json();
        if (!data.ok) return;
        const s = data.session;
        // Clear canvas when switching to a different session
        if (id !== currentSessionId) {
          clearCanvasForSession("Canvas cleared for session");
        }
        state.messages = (s.messages || []).filter(m => m.role === "user" || m.role === "assistant");
        if (s.model) { state.model = s.model; modelSelect.value = s.model; }
        if (s.chatMode) state.chatMode = s.chatMode;
        if (s.personality) { state.personality = s.personality; document.getElementById("personality").value = s.personality; }
        if (s.agentProfile) { state.agentProfile = s.agentProfile; state.activeAgentProfile = s.agentProfile; agentProfile.value = s.agentProfile; }
        if (typeof s.autoProfile === "boolean") state.autoProfile = s.autoProfile;
        if (s.assistantName) { state.assistantName = cleanAssistantName(s.assistantName); }
        if (typeof s.agentMode === "boolean") state.agentMode = s.agentMode;
        currentSessionId = s.id;
        sessionDeleteBtn.style.display = "";
        rerenderChatFromMessages(state.messages);
        updateStatus();
        refreshContextWindow();
      } catch (e) {
        addMessage("Error", String(e), "error");
      }
    }
    function startNewSession() {
      currentSessionId = null;
      state.messages = [];
      sessionSelect.value = "__new__";
      sessionDeleteBtn.style.display = "none";
      lastPromptTokens = 0;
      renderContextBar();
      chat.innerHTML = "";
      verboseLog.innerHTML = "";
      clearCanvasForSession("New session canvas");
      addMessage("Tool", "New session started.", "tool");
    }
    sessionSelect.addEventListener("change", () => {
      const v = sessionSelect.value;
      if (v === "__new__") { startNewSession(); return; }
      loadSessionById(v);
    });
    sessionDeleteBtn.addEventListener("click", async () => {
      if (!currentSessionId) return;
      if (!confirm("Delete this session?")) return;
      try {
        await fetch("/api/sessions/delete", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({id: currentSessionId}),
        });
        const deletedId = currentSessionId;
        currentSessionId = null;
        startNewSession();
        await refreshSessionList();
      } catch (_) {}
    });
    function scheduleSessionSave() {
      if (saveTimer) clearTimeout(saveTimer);
      saveTimer = setTimeout(async () => {
        if (!state.messages.length) return;
        try {
          const res = await fetch("/api/sessions/save", {
            method: "POST",
            headers: {"Content-Type": "application/json"},
            body: JSON.stringify({
              id: currentSessionId,
              model: state.model,
              messages: state.messages,
              agentMode: state.agentMode,
              chatMode: state.chatMode,
              personality: state.personality,
              agentProfile: state.agentProfile,
              autoProfile: Boolean(state.autoProfile),
              assistantName: state.assistantName,
            }),
          });
          const data = await res.json();
          if (data.ok && data.session) {
            const newId = data.session.id;
            const wasNew = !currentSessionId;
            currentSessionId = newId;
            await refreshSessionList(newId);
            if (wasNew) sessionDeleteBtn.style.display = "";
          }
        } catch (_) {}
      }, 250);
    }
    refreshSessionList();

    function setGenerating(on, label) {
      stopBtn.classList.toggle("active", on || ttsQueueActive || Boolean(activeAudio));
      document.getElementById("send").disabled = on;
      const ind = document.getElementById("genIndicator");
      ind.classList.toggle("active", on);
      document.getElementById("genLabel").textContent = label || "Thinking…";
    }

    function step(message) {
      if (!state.verbose) return;
      addMessage("step", message, "tool");
    }

    function stopConversationActivity() {
      setVoiceConversationActive(false);
      if (voiceIsRecording()) {
        try { mediaRecorder.stop(); } catch (_) {}
      }
      cancelTtsQueue();
      if (currentController) {
        currentController.abort();
        currentController = null;
      }
      setGenerating(false);
      updateMicUi();
    }

    stopBtn.addEventListener("click", () => {
      stopConversationActivity();
    });

    async function sendMessage() {
      const text = input.value.trim();
      if (!text && !pendingAttachment) return;
      input.value = "";
      if (!inputHistory.length || inputHistory[inputHistory.length - 1] !== text) {
        inputHistory.push(text);
        if (inputHistory.length > 500) inputHistory.shift();
        localStorage.setItem("ollamaWebInputHistory", JSON.stringify(inputHistory));
      }
      historyIndex = -1;
      draftBeforeHistory = "";
      if (text === "/clear") {
        chat.innerHTML = "";
        state.messages = [];
        addMessage("Tool", "Chat cleared.", "tool");
        return;
      }
      if (text === "/verbose on") { state.verbose = true; updateStatus(); }
      if (text === "/verbose off") { state.verbose = false; updateStatus(); }
      if (text === "/agent on") { state.agentMode = true; updateStatus(); }
      if (text === "/agent off") { state.agentMode = false; updateStatus(); }
      if (text.startsWith("/chat_mode ")) { state.chatMode = text.slice(11).trim() || state.chatMode; state.agentMode = false; updateStatus(); }
      if (text.startsWith("/chat_personality ")) state.personality = text.slice(18).trim() || state.personality;
      if (text.startsWith("/agent_profile ")) {
        state.agentProfile = text.slice(15).trim() || state.agentProfile;
        state.activeAgentProfile = state.agentProfile;
        updateStatus();
      }
      const displayText = text || (pendingAttachment ? `[${pendingAttachment.filename}]` : "");
      addMessage("You", displayText, "user");
      step("Preparing request");
      const deepResearchMode = state.agentProfile === "deep_research";
      if (deepResearchMode) {
        state.agentMode = true;
        updateStatus();
        setCanvasOpen(true);
        step("Deep Research profile active; final report will replace canvas");
      }
      if (pendingAttachment?.type === "image") {
        const img = document.createElement("img");
        img.src = "data:" + pendingAttachment.mime + ";base64," + pendingAttachment.data;
        img.style.cssText = "max-width:200px;max-height:150px;border-radius:4px;margin-top:4px;display:block";
        chat.lastElementChild.querySelector("div").append(img);
        chat.scrollTop = chat.scrollHeight;
        step(`Attached image: ${pendingAttachment.filename || "image"}`);
      }
      const canvasRequested = deepResearchMode || shouldAutoCanvas(text);
      const attachment = pendingAttachment;
      // Pass canvas content to server when user asks to check/edit/update it
      const CANVAS_AI_KEYWORDS = /\b(canvas|check|update|edit|rewrite|revise|improve|summarize|extend|read|document|fix|clean|format|translate|shorten|expand|what.s in|what is)\b/i;
      const canvasContent = (CANVAS_AI_KEYWORDS.test(text) && canvasEditor.value.trim())
        ? canvasEditor.value.trim() : "";
      if (attachment?.type === "document") step(`Attached document: ${attachment.filename || "document"}`);
      if (canvasContent) step(`Reading canvas context (${canvasContent.length} chars)`);
      if (canvasRequested) step("Canvas output requested; final answer will replace canvas");
      clearAttachment();
      const pending = addVerboseMessage("step", "Working...", "step");
      currentController = new AbortController();
      setGenerating(true);
      try {
        step(`Sending request to ${state.model || "selected model"} on GPU ${state.gpu || "0"}`);
        const res = await fetch("/api/chat-stream", {
          method: "POST",
          headers: {"Content-Type": "application/json"},
          body: JSON.stringify({...state, text, attachment, canvasContent, canvasRequested, gpu: state.gpu}),
          signal: currentController.signal,
        });
        pending.remove();
        if (!res.ok || !res.body) throw new Error("HTTP " + res.status);
        const reader = res.body.getReader();
        const decoder = new TextDecoder();
        let buffer = "";
        let answerBody = null;
        let answerText = "";
        let thinkingBlock = null;
        let thinkingText = "";
        let forceCanvas = false;
        let sawFirstChunk = false;
        const streamTtsEnabled = voiceConversationActive && state.voiceAutoSend;
        let streamTtsBuffer = "";
        let streamTtsStarted = false;
        const queueStreamingSpeech = (textPart, force = false) => {
          if (!streamTtsEnabled || (!textPart && !force)) return;
          streamTtsBuffer += textPart;
          const result = takeSpeakableSegments(streamTtsBuffer, force);
          streamTtsBuffer = result.rest;
          for (const segment of result.segments) {
            if (!streamTtsStarted) {
              step("Reading assistant response aloud as it streams");
              streamTtsStarted = true;
            }
            queueSpeakText(segment);
          }
        };
        const ensureThinkingBlock = () => {
          if (thinkingBlock) return thinkingBlock;
          if (!answerBody) answerBody = createMessage(cleanAssistantName(state.assistantName), "ollama");
          thinkingBlock = document.createElement("details");
          thinkingBlock.className = "thinking-block";
          thinkingBlock.open = true;
          const summary = document.createElement("summary");
          summary.textContent = "Thinking";
          const body = document.createElement("div");
          body.className = "thinking-body";
          thinkingBlock.append(summary, body);
          answerBody.prepend(thinkingBlock);
          return thinkingBlock;
        };
        const handleEvent = (event) => {
          if (event.type === "error") {
            addMessage("Error", event.error, "error");
          } else if (event.type === "cmd") {
            if (/Rendering \d+ page/.test(event.command)) forceCanvas = true;
            addMessage("cmd", event.command, "tool");
          } else if (event.type === "context") {
            if (event.profile) {
              state.activeAgentProfile = event.profile;
              updateStatus();
            }
            addMessage("context", event.text || "", "context");
          } else if (event.type === "tool") {
            addMessage(event.role || "Tool", event.text || "", "tool");
            step("Tool output received; adding it to model context");
          } else if (event.type === "thinking") {
            if (!thinkingText) step("Model is producing reasoning trace");
            thinkingText += event.text || "";
            const block = ensureThinkingBlock();
            block.querySelector(".thinking-body").textContent = thinkingText;
            setGenerating(true, "Thinking…");
            chat.scrollTop = chat.scrollHeight;
          } else if (event.type === "chunk") {
            if (!answerBody) answerBody = createMessage(cleanAssistantName(state.assistantName), "ollama");
            if (!sawFirstChunk) {
              step("Receiving assistant output");
              sawFirstChunk = true;
            }
            answerText += event.text || "";
            queueStreamingSpeech(event.text || "");
            // Re-render the answer text (without disturbing thinking block)
            let answerHolder = answerBody.querySelector(".answer-text");
            if (!answerHolder) {
              answerHolder = document.createElement("div");
              answerHolder.className = "answer-text";
              answerBody.append(answerHolder);
            }
            liveRender(answerHolder, answerText);
            setGenerating(true, "Generating…");
            chat.scrollTop = chat.scrollHeight;
          } else if (event.type === "stats" && event.stats) {
            if (state.verbose && event.stats.text) {
              addMessage("stats", event.stats.text, "tool");
            }
            if (typeof event.stats.prompt_tokens === "number") {
              lastPromptTokens = event.stats.prompt_tokens;
              renderContextBar();
            }
          } else if (event.type === "final") {
            applyResponseState(event);
            step("Received final response");
            if (event.commands?.length) {
              for (const command of event.commands) addMessage("cmd", command, "tool");
            }
            if (event.tool?.command) addMessage("cmd", event.tool.command, "tool");
            if (event.tool?.text) addMessage(event.tool.role || "Tool", event.tool.text, "tool");
            if (event.text) {
              answerText = event.text;
              if (!sawFirstChunk) queueStreamingSpeech(event.text, true);
              if (!answerBody) answerBody = createMessage(event.role || cleanAssistantName(state.assistantName), "ollama");
              let answerHolder = answerBody.querySelector(".answer-text");
              if (!answerHolder) {
                answerHolder = document.createElement("div");
                answerHolder.className = "answer-text";
                answerBody.append(answerHolder);
              }
              liveRender(answerHolder, answerText);
            }
            if (event.stats) {
              if (state.verbose && event.stats.text) addMessage("stats", event.stats.text, "tool");
              if (typeof event.stats.prompt_tokens === "number") {
                lastPromptTokens = event.stats.prompt_tokens;
                renderContextBar();
              }
            }
          }
        };
        while (true) {
          const {value, done} = await reader.read();
          if (done) break;
          buffer += decoder.decode(value, {stream: true});
          const parts = buffer.split("\n\n");
          buffer = parts.pop() || "";
          for (const part of parts) {
            const line = part.split("\n").find(item => item.startsWith("data: "));
            if (line) handleEvent(JSON.parse(line.slice(6)));
          }
        }
        applyResponseState({});
        if (streamTtsEnabled) {
          queueStreamingSpeech("", true);
          await ttsQueue;
        }
        if (answerBody) {
          finalizeBody(answerBody, answerText);
          addSpeakButton(answerBody, answerText);
          step("Rendered final markdown response");
        }
        if (answerText) {
          lastAnswerText = answerText;
          readLastBtn.disabled = false;
          if (canvasRequested) {
            const mode = canvasModeFor(text);
            const canvasText = String(answerText || "").trim();
            console.log("[canvas]", mode, "len=", canvasText.length, "preview=", canvasText.slice(0, 120));
            step(`${mode === "replace" ? "Replacing" : "Appending to"} canvas (${canvasText.length} chars)`);
            writeCanvas(canvasText, mode);
            addMessage("Tool", `→ Canvas (${mode}, ${canvasText.length} chars)`, "tool");
          }
        }
        state.messages.push({role: "user", content: displayText});
        state.messages.push({role: "assistant", content: answerText});
        state.messages = state.messages.slice(-30);
        step("Saving chat session");
        scheduleSessionSave();
      } catch (err) {
        if (pending.isConnected) pending.remove();
        if (err.name !== "AbortError") {
          addMessage("Error", String(err), "error");
          step("Request failed");
        } else {
          addMessage("Tool", "Generation stopped.", "tool");
          step("Request stopped by user");
        }
      } finally {
        currentController = null;
        setGenerating(false);
        scheduleVoiceConversationRestart();
      }
    }

    modelSelect.addEventListener("change", () => { state.model = modelSelect.value; updateStatus(); refreshContextWindow(); });
    personality.addEventListener("change", () => { state.personality = personality.value; updateStatus(); });
    agentProfile.addEventListener("change", () => {
      state.agentProfile = agentProfile.value;
      state.activeAgentProfile = agentProfile.value;
      updateStatus();
    });
    autoProfile.addEventListener("change", () => {
      state.autoProfile = autoProfile.checked;
      state.activeAgentProfile = state.agentProfile;
      updateStatus();
    });
    keepAlive.addEventListener("change", () => { state.keepAlive = keepAlive.checked; updateStatus(); });
    assistantNameButton.addEventListener("click", (event) => {
      event.stopPropagation();
      if (assistantNamePopover.classList.contains("open")) closeAssistantNamePopover();
      else openAssistantNamePopover();
    });
    assistantNameSaveBtn.addEventListener("click", saveAssistantNameFromPopover);
    assistantNameInput.addEventListener("keydown", event => {
      if (event.key === "Enter") {
        event.preventDefault();
        saveAssistantNameFromPopover();
        input.focus();
      } else if (event.key === "Escape") {
        event.preventDefault();
        closeAssistantNamePopover();
      }
    });
    document.addEventListener("click", event => {
      if (!assistantNamePopover.contains(event.target) && !assistantNameButton.contains(event.target)) {
        closeAssistantNamePopover();
      }
    });
    verboseBtn.addEventListener("click", () => { state.verbose = !state.verbose; updateStatus(); });
    if (mobileSettingsBtn) {
      mobileSettingsBtn.addEventListener("click", () => {
        appRoot.classList.toggle("mobile-settings-open");
        mobileSettingsBtn.classList.toggle("active", appRoot.classList.contains("mobile-settings-open"));
      });
    }
    if (mobileActivityBtn) mobileActivityBtn.addEventListener("click", () => setMobileTab("activity"));
    if (mobileCanvasBtn) mobileCanvasBtn.addEventListener("click", () => setMobileTab("canvas"));
    for (const btn of mobileTabBtns) {
      btn.addEventListener("click", () => setMobileTab(btn.dataset.mobileTab || "chat"));
    }
    setMobileTab("chat");
    input.addEventListener("input", () => {
      if (historyIndex !== -1) {
        historyIndex = -1;
        draftBeforeHistory = "";
      }
    });
    document.getElementById("send").addEventListener("click", sendMessage);
    input.addEventListener("keydown", event => {
      if (event.key === "Enter" && !event.shiftKey) {
        event.preventDefault();
        if (!currentController) sendMessage();
        return;
      }
      if (event.key === "ArrowUp" && !event.shiftKey && !event.altKey && !event.ctrlKey && !event.metaKey) {
        const atFirstLine = input.selectionStart === input.selectionEnd &&
          input.value.slice(0, input.selectionStart).indexOf("\n") === -1;
        if (atFirstLine && inputHistory.length) {
          event.preventDefault();
          if (historyIndex === -1) {
            draftBeforeHistory = input.value;
            historyIndex = inputHistory.length - 1;
          } else {
            historyIndex = Math.max(0, historyIndex - 1);
          }
          input.value = inputHistory[historyIndex];
          input.setSelectionRange(input.value.length, input.value.length);
        }
        return;
      }
      if (event.key === "ArrowDown" && !event.shiftKey && !event.altKey && !event.ctrlKey && !event.metaKey) {
        const afterCursor = input.value.slice(input.selectionStart);
        const atLastLine = input.selectionStart === input.selectionEnd && afterCursor.indexOf("\n") === -1;
        if (atLastLine && historyIndex !== -1) {
          event.preventDefault();
          if (historyIndex < inputHistory.length - 1) {
            historyIndex += 1;
            input.value = inputHistory[historyIndex];
          } else {
            historyIndex = -1;
            input.value = draftBeforeHistory;
            draftBeforeHistory = "";
          }
          input.setSelectionRange(input.value.length, input.value.length);
        }
      }
    });
    // ── Ollama API status indicator ───────────────────────────────────────────
    const ollamaStatus = document.getElementById("ollamaStatus");
    async function checkOllamaStatus() {
      try {
        const res = await fetch("/api/models", {cache: "no-store"});
        const data = await res.json();
        if (data.ok) {
          ollamaStatus.textContent = "⬤ online";
          ollamaStatus.className = "online";
          ollamaStatus.title = "Ollama API reachable";
        } else {
          ollamaStatus.textContent = "⬤ error";
          ollamaStatus.className = "offline";
          ollamaStatus.title = data.error || "API returned an error";
        }
      } catch {
        ollamaStatus.textContent = "⬤ offline";
        ollamaStatus.className = "offline";
        ollamaStatus.title = "Cannot reach Ollama API";
      }
    }
    checkOllamaStatus();
    setInterval(checkOllamaStatus, 15000);
    // ─────────────────────────────────────────────────────────────────────────

    async function refreshRuntimeBanner(){
      try {
        const r = await fetch("/api/runtime-status");
        if(!r.ok) return;
        const d = await r.json();
        const danger = document.getElementById("rtDanger");
        if(d.dangerous_mode){ danger.classList.remove("hide"); }
        else { danger.classList.add("hide"); }
        document.getElementById("rtOllama").textContent = "Ollama: " + (d.ollama_url || "?");
        const t = d.tools || {};
        const enabled = Object.keys(t).filter(k => t[k]);
        const labels = {local_command:"local", ssh_command:"ssh", internet_search:"search", current_datetime:"time", ollama_models:"models", thunderbird_readonly:"thunderbird"};
        const pretty = enabled.map(k => labels[k] || k);
        document.getElementById("rtTools").textContent = "Tools: " + (pretty.length ? pretty.join(", ") : "none");
      } catch(_){}
    }
    refreshRuntimeBanner();
    setInterval(refreshRuntimeBanner, 30000);

    addMessage("Tool", "Web TUI ready. Select a model and start chatting.", "tool");
    loadModels();
  </script>
</body>
</html>
"""

FAVICON_SVG = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 32 32">
<rect width="32" height="32" rx="6" fill="#181b1f"/>
<path d="M7 10h18v12H7z" fill="#7cc7ff"/>
<path d="M10 14h12v2H10zm0 4h8v2h-8z" fill="#101214"/>
</svg>
"""

_SETUP_DONE_RUNTIME = False

# Short-lived tokens issued by /api/setup-auth so /api/setup can verify
# the operator re-confirmed credentials within the same browsing session.
# Maps token -> (username, expires_epoch).
_SETUP_TOKENS: dict[str, tuple[str, float]] = {}
_SETUP_TOKEN_TTL = 600  # 10 minutes


def _issue_setup_token(username: str) -> str:
    import secrets as _secrets
    import time as _time
    token = _secrets.token_urlsafe(32)
    _SETUP_TOKENS[token] = (username, _time.time() + _SETUP_TOKEN_TTL)
    # Garbage-collect expired tokens.
    now = _time.time()
    for k in [k for k, (_u, exp) in _SETUP_TOKENS.items() if exp < now]:
        _SETUP_TOKENS.pop(k, None)
    return token


def _consume_setup_token(token: str) -> str | None:
    import time as _time
    entry = _SETUP_TOKENS.pop(token, None)
    if not entry:
        return None
    username, expires = entry
    if expires < _time.time():
        return None
    return username


def _current_setup_config() -> dict:
    """Snapshot the editable subset of config.toml for pre-populating the setup form."""
    import urllib.parse
    ssh_cfg = ""
    ssh_kh = ""
    try:
        from ollama_tools import config as _cfg_mod
        ssh_section = _cfg_mod._cfg.get("ssh", {})
        ssh_cfg = str(ssh_section.get("config_path") or "/data/.ssh/config")
        ssh_kh = str(ssh_section.get("known_hosts_path") or "/data/.ssh/known_hosts")
    except Exception:
        ssh_cfg = "/data/.ssh/config"
        ssh_kh = "/data/.ssh/known_hosts"
    return {
        "ollama_url": get_ollama_base_url(),
        "work_dir": str(get_work_dir()),
        "enable_local": get_tool_enabled("local_command"),
        "enable_ssh": get_tool_enabled("ssh_command"),
        "enable_search": get_tool_enabled("internet_search"),
        "enable_mcp": get_tool_enabled("monitoring") or get_tool_enabled("mcp"),
        "enable_thunderbird": get_thunderbird_enabled(),
        "dangerous_mode": get_dangerous_mode(),
        "auth_enabled": get_auth_enabled() if get_setup_completed() else True,
        "skip_login": (get_auth_skip_allowed() or not get_auth_enabled()) if get_setup_completed() else False,
        "piper_url": get_piper_url(),
        "searxng_url": get_searxng_url(),
        # Brave key intentionally not echoed back to the browser.
        "ssh_config_path": ssh_cfg,
        "ssh_known_hosts_path": ssh_kh,
    }


def _runtime_status_snapshot() -> dict:
    return {
        "dangerous_mode": get_dangerous_mode(),
        "ollama_url": get_ollama_base_url(),
        "work_dir": str(get_work_dir()),
        "tools": {
            "local_command": get_tool_enabled("local_command"),
            "ssh_command": get_tool_enabled("ssh_command"),
            "internet_search": get_tool_enabled("internet_search"),
            "current_datetime": get_tool_enabled("current_datetime"),
            "ollama_models": get_tool_enabled("ollama_models"),
            "thunderbird_readonly": get_thunderbird_enabled(),
        },
        "setup_completed": get_setup_completed(),
        "has_user": has_users(),
        "auth_enabled": get_auth_enabled(),
    }

LOGIN_HTML = """<!doctype html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Ollama Command Deck Login</title>
<style>body{margin:0;background:#0f141b;color:#e6edf3;font:15px system-ui;display:grid;place-items:center;min-height:100vh}main{width:min(420px,92vw);border:1px solid #30363d;background:#161b22;padding:24px;border-radius:8px}label{display:block;margin:12px 0 5px;color:#9fb0c3}input,button{box-sizing:border-box;width:100%;padding:10px;border-radius:6px;border:1px solid #30363d;background:#0d1117;color:#e6edf3}button{margin-top:18px;background:#1f6feb;border-color:#1f6feb;font-weight:700;cursor:pointer}.err{color:#ff7b72}</style>
</head><body><main><h1>Ollama Command Deck</h1><p>Sign in to continue.</p><p id="err" class="err"></p>
<label>Username</label><input id="u" autocomplete="username" autofocus>
<label>Password</label><input id="p" type="password" autocomplete="current-password">
<button id="b">Sign In</button>
<script>
async function login(){err.textContent="";const res=await fetch("/api/auth/login",{method:"POST",headers:{"Content-Type":"application/json"},body:JSON.stringify({username:u.value,password:p.value})});if(res.ok){location.href="/"}else{err.textContent="Invalid username or password."}}
b.onclick=login;p.addEventListener("keydown",e=>{if(e.key==="Enter")login()});
</script></main></body></html>"""

SETUP_HTML = r"""<!doctype html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Ollama Command Deck Setup</title>
<style>
body{margin:0;background:#0f141b;color:#e6edf3;font:15px system-ui;display:grid;place-items:start center;min-height:100vh;padding:24px}
main{width:min(820px,96vw);border:1px solid #30363d;background:#161b22;padding:24px;border-radius:8px}
h1{margin:0 0 4px}
h2{margin:24px 0 8px;font-size:17px;color:#cdd9e5;border-bottom:1px solid #21262d;padding-bottom:6px}
label{display:block;margin:12px 0 5px;color:#9fb0c3;font-size:13px}
input{box-sizing:border-box;width:100%;padding:9px;border-radius:6px;border:1px solid #30363d;background:#0d1117;color:#e6edf3;font:14px system-ui}
.row{display:grid;grid-template-columns:1fr 1fr;gap:12px}
.check{display:flex;gap:8px;align-items:flex-start;margin:8px 0;font-size:14px;line-height:1.4}
.check input{width:auto;margin-top:3px}
.check span code{background:#0d1117;padding:1px 4px;border-radius:3px}
button{padding:10px 14px;border-radius:6px;border:1px solid #1f6feb;background:#1f6feb;color:#fff;font-weight:700;cursor:pointer;font-size:14px}
button.danger{border-color:#da3633;background:#da3633}
button.ghost{border-color:#30363d;background:transparent;color:#9fb0c3;font-weight:500}
.warn{border-left:4px solid #f2cc60;background:#2d250f;padding:12px;font-size:14px;line-height:1.5}
.danger-banner{border-left:4px solid #da3633;background:#2d0f0f;padding:12px;margin-top:12px}
.err{color:#ff7b72}
.ok{color:#7ee787}
.hide{display:none}
.muted{color:#7d8590;font-size:13px}
.actions{margin-top:18px;display:flex;gap:10px;flex-wrap:wrap}
hr{border:none;border-top:1px solid #21262d;margin:18px 0}
</style>
</head><body><main>
<h1 id="title">Loading…</h1>
<p id="subtitle" class="muted"></p>
<p class="warn">__WARNING__</p>
<p id="msg"></p>

<!-- Section: existing-user re-auth gate (shown only when a user already exists) -->
<section id="authGate" class="hide">
  <h2>Confirm identity</h2>
  <p class="muted">A user already exists on this deployment. Enter the current credentials to edit the configuration.</p>
  <div class="row">
    <div><label>Current username</label><input id="auth_user"></div>
    <div><label>Current password</label><input id="auth_pass" type="password" autocomplete="current-password"></div>
  </div>
  <div class="actions">
    <button id="authBtn">Unlock editor</button>
    <button id="forgotBtn" class="ghost" type="button">Forgot password — reset everything</button>
  </div>

  <div id="forgotPanel" class="hide danger-banner">
    <p><b>Reset everything.</b> This deletes the existing user, all chat history, all canvas files, all overrides, and the rendered config. After reset you'll see the fresh-install form.</p>
    <div class="actions">
      <button id="confirmReset" class="danger" type="button">Yes, wipe and start over</button>
      <button id="cancelReset" class="ghost" type="button">Cancel</button>
    </div>
  </div>
</section>

<!-- Section: editor (shown for fresh install or after re-auth) -->
<section id="editor" class="hide">
  <h2 id="editorTitle">Configuration</h2>

  <div class="check"><input id="auth_enabled" type="checkbox" checked><span>Require sign-in for the Web UI. Turn this off only on a trusted LAN or behind another auth layer.</span></div>

  <div id="userBlock">
    <div class="row">
      <div><label>Username</label><input id="username" value="admin"></div>
      <div><label id="passwordLabel">Password</label><input id="password" type="password" placeholder=""></div>
    </div>
  </div>

  <h2>Ollama</h2>
  <label>Ollama API URL</label>
  <input id="ollama_url" value="http://localhost:11434">
  <label>Ollama API key / token (optional)</label>
  <input id="ollama_api_key" type="password" placeholder="Leave blank if Ollama does not require one">
  <p id="apiKeyHint" class="muted hide">Leave blank to keep the existing key. Type a new value to replace it.</p>

  <h2>Working directory</h2>
  <label>Container working directory (commands run here unless dangerous mode is on)</label>
  <input id="work_dir" value="/workspace">

  <h2>Tools</h2>
  <div class="check"><input id="enable_local" type="checkbox"><span>Enable local command tools. High risk.</span></div>
  <div class="check"><input id="enable_ssh" type="checkbox"><span>Enable SSH tools. The agent reads <code>/data/.ssh/config</code> for host aliases. High risk.</span></div>
  <div class="check"><input id="enable_search" type="checkbox"><span>Enable internet search tools. Medium risk.</span></div>
  <div class="check"><input id="enable_mcp" type="checkbox"><span>Enable MCP server/tooling.</span></div>
  <div class="check"><input id="enable_thunderbird" type="checkbox"><span>Enable read-only Thunderbird bridge. Thunderbird can send selected/search-result message snippets to Command Deck for analysis. No send, delete, move, or compose permissions.</span></div>
  <label>Thunderbird bridge token</label>
  <input id="thunderbird_token" readonly placeholder="Generated after enabling and saving">
  <p class="muted">Install the example add-on from <code>examples/thunderbird-readonly</code> and paste this token into its settings.</p>

  <h2>Search backends (optional)</h2>
  <label>SearXNG URL</label>
  <input id="searxng_url" placeholder="https://searx.example/">
  <label>Brave Search API key</label>
  <input id="brave_api_key" type="password" placeholder="Leave blank to disable Brave">

  <h2>Text-to-speech (optional)</h2>
  <label>Piper / Kokoro URL</label>
  <input id="piper_url" placeholder="http://piper:8880">

  <h2>SSH paths</h2>
  <div class="row">
    <div><label>SSH config path</label><input id="ssh_config_path" value="/data/.ssh/config"></div>
    <div><label>known_hosts path</label><input id="ssh_known_hosts_path" value="/data/.ssh/known_hosts"></div>
  </div>

  <hr>
  <div class="check"><input id="dangerous_mode" type="checkbox"><span><b>Dangerous mode.</b> Lets the agent run any non-privilege-escalation command anywhere in the container — bypasses work-directory confinement and the destructive-command blocklist. Privilege escalation (sudo, su, doas, pkexec, passwd, visudo) and a fork-bomb pattern remain blocked. Docker isolation becomes the only meaningful boundary. <b>Extreme risk.</b></span></div>
  <div class="check"><input id="ack" type="checkbox"><span>I understand these tools can harm this computer or configured devices if misused.</span></div>

  <div class="actions">
    <button id="saveBtn">Save and apply</button>
    <a href="/" class="ghost" style="text-decoration:none"><button class="ghost" type="button">Cancel</button></a>
  </div>
  <p class="muted" id="applyHint">Saving rewrites <code>/config/config.toml</code>, reloads the runtime, and rotates the session secret — all open browser sessions will need to log in again.</p>
</section>

<script>
(async function(){
  const state = await (await fetch("/api/setup-state")).json();
  const msg = document.getElementById("msg");
  const authGate = document.getElementById("authGate");
  const editor = document.getElementById("editor");
    const editorTitle = document.getElementById("editorTitle");
    const passwordLabel = document.getElementById("passwordLabel");
    const passwordInput = document.getElementById("password");
    const userBlock = document.getElementById("userBlock");
  const apiKeyHint = document.getElementById("apiKeyHint");
  const apiKeyInput = document.getElementById("ollama_api_key");

  let authToken = "";

  function populate(cfg){
    document.getElementById("username").value = cfg.username || "admin";
    document.getElementById("auth_enabled").checked = cfg.auth_enabled !== false && cfg.skip_login !== true;
    document.getElementById("ollama_url").value = cfg.ollama_url || "http://localhost:11434";
    document.getElementById("work_dir").value = cfg.work_dir || "/workspace";
    document.getElementById("enable_local").checked = !!cfg.enable_local;
    document.getElementById("enable_ssh").checked = !!cfg.enable_ssh;
    document.getElementById("enable_search").checked = !!cfg.enable_search;
    document.getElementById("enable_mcp").checked = !!cfg.enable_mcp;
    document.getElementById("enable_thunderbird").checked = !!cfg.enable_thunderbird;
    document.getElementById("thunderbird_token").value = cfg.enable_thunderbird ? "(save to reveal token)" : "";
    document.getElementById("dangerous_mode").checked = !!cfg.dangerous_mode;
    document.getElementById("searxng_url").value = cfg.searxng_url || "";
    document.getElementById("brave_api_key").value = "";
    document.getElementById("piper_url").value = cfg.piper_url || "";
    document.getElementById("ssh_config_path").value = cfg.ssh_config_path || "/data/.ssh/config";
    document.getElementById("ssh_known_hosts_path").value = cfg.ssh_known_hosts_path || "/data/.ssh/known_hosts";
  }

  function showEditor(mode){
    editor.classList.remove("hide");
    if(mode === "edit"){
      document.getElementById("title").textContent = "Edit configuration";
      document.getElementById("subtitle").textContent = "Existing deployment — leave password blank to keep it.";
      editorTitle.textContent = "User";
      passwordLabel.textContent = "New password (leave blank to keep current)";
      passwordInput.placeholder = "(unchanged)";
      apiKeyHint.classList.remove("hide");
      apiKeyInput.placeholder = "(unchanged — leave blank to keep)";
    } else {
      document.getElementById("title").textContent = "First-run setup";
      document.getElementById("subtitle").textContent = "Configure this Ollama Command Deck deployment.";
      editorTitle.textContent = "Create admin user";
    }
  }

  function syncAuthFields(){
    const authEnabled = document.getElementById("auth_enabled").checked;
    userBlock.classList.toggle("hide", !authEnabled);
    document.getElementById("password").required = authEnabled && !state.has_user;
  }

  if(state.has_user && state.auth_enabled !== false){
    authGate.classList.remove("hide");
    document.getElementById("auth_user").value = state.username || "";
    document.getElementById("authBtn").onclick = async function(){
      msg.className = ""; msg.textContent = "";
      const u = document.getElementById("auth_user").value;
      const p = document.getElementById("auth_pass").value;
      const r = await fetch("/api/setup-auth", {method:"POST", headers:{"Content-Type":"application/json"}, body: JSON.stringify({username:u, password:p})});
      const d = await r.json().catch(()=>({}));
      if(!r.ok || !d.ok){ msg.className="err"; msg.textContent = d.error || "Invalid credentials."; return; }
      authToken = d.token;
      authGate.classList.add("hide");
      populate(state.config);
      document.getElementById("username").value = u;
      showEditor("edit");
      syncAuthFields();
    };
    document.getElementById("forgotBtn").onclick = function(){
      document.getElementById("forgotPanel").classList.remove("hide");
    };
    document.getElementById("cancelReset").onclick = function(){
      document.getElementById("forgotPanel").classList.add("hide");
    };
    document.getElementById("confirmReset").onclick = async function(){
      msg.className = ""; msg.textContent = "Resetting…";
      const r = await fetch("/api/setup-reset", {method:"POST"});
      const d = await r.json().catch(()=>({}));
      if(!r.ok || !d.ok){ msg.className="err"; msg.textContent = d.error || "Reset failed."; return; }
      location.reload();
    };
  } else if (state.has_user) {
    populate(state.config);
    showEditor("edit");
    syncAuthFields();
  } else {
    populate(state.config);
    showEditor("new");
    syncAuthFields();
  }

  document.getElementById("auth_enabled").addEventListener("change", syncAuthFields);

  document.getElementById("saveBtn").onclick = async function(){
    msg.className = ""; msg.textContent = "";
    if(!document.getElementById("ack").checked){ msg.className="err"; msg.textContent="You must acknowledge the warning."; return; }
    const body = {
      auth_token: authToken,
      username: document.getElementById("username").value,
      password: document.getElementById("password").value,
      auth_enabled: document.getElementById("auth_enabled").checked,
      ollama_url: document.getElementById("ollama_url").value,
      ollama_api_key: document.getElementById("ollama_api_key").value,
      keep_existing_api_key: state.has_user && document.getElementById("ollama_api_key").value === "",
      work_dir: document.getElementById("work_dir").value,
      enable_local: document.getElementById("enable_local").checked,
      enable_ssh: document.getElementById("enable_ssh").checked,
      enable_search: document.getElementById("enable_search").checked,
      enable_mcp: document.getElementById("enable_mcp").checked,
      enable_thunderbird: document.getElementById("enable_thunderbird").checked,
      dangerous_mode: document.getElementById("dangerous_mode").checked,
      piper_url: document.getElementById("piper_url").value,
      searxng_url: document.getElementById("searxng_url").value,
      brave_api_key: document.getElementById("brave_api_key").value,
      ssh_config_path: document.getElementById("ssh_config_path").value,
      ssh_known_hosts_path: document.getElementById("ssh_known_hosts_path").value,
    };
    const r = await fetch("/api/setup", {method:"POST", headers:{"Content-Type":"application/json"}, body: JSON.stringify(body)});
    const d = await r.json().catch(()=>({}));
    if(!r.ok || !d.ok){ msg.className="err"; msg.textContent = d.error || "Save failed."; return; }
    if(d.thunderbird_token){
      document.getElementById("thunderbird_token").value = d.thunderbird_token;
    }
    msg.className = "ok";
    msg.textContent = d.thunderbird_token ? "Saved. Copy the Thunderbird token, then continue." : (body.auth_enabled ? "Saved. Redirecting to login…" : "Saved. Redirecting to app…");
    if(!d.thunderbird_token){
      setTimeout(function(){ location.href = body.auth_enabled ? "/login" : "/"; }, 1000);
    }
  };
})();
</script>
</main></body></html>""".replace("__WARNING__", WARNING_TEXT)


class Handler(BaseHTTPRequestHandler):
    server_version = "OllamaWebTUI/0.1"

    def _setup_required(self) -> bool:
        return is_deploy_mode() and not (
            _SETUP_DONE_RUNTIME
            or (get_setup_completed() and (not get_auth_enabled() or has_users()))
        )

    def _cookie(self, name: str) -> str | None:
        raw = self.headers.get("Cookie", "")
        for part in raw.split(";"):
            if "=" not in part:
                continue
            key, value = part.strip().split("=", 1)
            if key == name:
                return value
        return None

    def _current_user(self):
        if not get_auth_enabled():
            return {"username": "local", "role": "admin"}
        user = verify_session(self._cookie("ollama_hooks_session"))
        return {"username": user.username, "role": user.role} if user else None

    def _public_path(self) -> bool:
        clean = self.path.split("?", 1)[0]
        return clean in {
            "/favicon.ico", "/healthz", "/setup", "/login",
            "/api/setup", "/api/setup-state", "/api/setup-auth", "/api/setup-reset",
            "/api/auth/login", "/api/thunderbird/analyze",
        }

    def _require_access(self) -> bool:
        clean = self.path.split("?", 1)[0]
        if self._setup_required() and clean not in {"/setup", "/api/setup", "/api/setup-state", "/api/setup-auth", "/api/setup-reset", "/favicon.ico", "/healthz"}:
            if clean.startswith("/api/"):
                self.respond_json({"ok": False, "error": "setup required"}, status=428)
            else:
                self.redirect("/setup")
            return False
        if self._public_path():
            return True
        if self._current_user():
            return True
        if clean.startswith("/api/"):
            self.respond_json({"ok": False, "error": "authentication required"}, status=401)
        else:
            self.redirect("/login")
        return False

    def do_OPTIONS(self) -> None:
        if self.path == "/api/thunderbird/analyze":
            self.send_response(204)
            self.send_header("Access-Control-Allow-Origin", "*")
            self.send_header("Access-Control-Allow-Headers", "Content-Type")
            self.send_header("Access-Control-Allow-Methods", "POST, OPTIONS")
            self.end_headers()
            return
        self.send_error(404)

    def do_GET(self) -> None:
        if not self._require_access():
            return
        if self.path == "/healthz":
            self.respond_json({"ok": True, "setup_required": self._setup_required()})
            return
        if self.path == "/setup":
            self.respond_text(SETUP_HTML, "text/html; charset=utf-8")
            return
        if self.path == "/login":
            self.respond_text(LOGIN_HTML, "text/html; charset=utf-8")
            return
        if self.path == "/" or self.path == "/index.html":
            self.respond_text(INDEX_HTML, "text/html; charset=utf-8")
            return
        if self.path == "/favicon.ico":
            self.respond_text(FAVICON_SVG, "image/svg+xml")
            return
        if self.path.startswith("/api/models"):
            from urllib.parse import urlparse, parse_qs
            qs = parse_qs(urlparse(self.path).query)
            gpu_id = (qs.get("gpu") or ["0"])[0]
            base_url = _gpu_base_url(gpu_id)
            try:
                models_out = []
                for m in list_models(base_url=base_url):
                    d = asdict(m)
                    d["capabilities"] = model_capabilities(m.name)
                    d["gpu"] = gpu_id
                    models_out.append(d)
                data = {"ok": True, "models": models_out, "gpu": gpu_id}
            except Exception as exc:
                data = {"ok": False, "error": str(exc)}
            self.respond_json(data)
            return
        if self.path == "/api/gpu-list":
            self.respond_json({"ok": True, "gpus": _gpu_rows()})
            return
        if self.path == "/api/stt-status":
            self.respond_json({"ok": True, "enabled": is_whisper_available()})
            return
        if self.path.startswith("/api/model-info"):
            from urllib.parse import urlparse, parse_qs
            qs = parse_qs(urlparse(self.path).query)
            model = (qs.get("model") or [""])[0]
            gpu_id = (qs.get("gpu") or ["0"])[0]
            ctx = fetch_model_context_info(model, _gpu_base_url(gpu_id))
            self.respond_json({
                "ok": True,
                "context_length": ctx.get("advertised_context_length") or ctx.get("runtime_context_length"),
                **ctx,
            })
            return
        if self.path.startswith("/api/canvas-files"):
            from urllib.parse import urlparse, parse_qs
            qs = parse_qs(urlparse(self.path).query)
            name = (qs.get("name") or [""])[0]
            if name:
                content = load_canvas_file(name)
                if content is None:
                    self.respond_json({"ok": False, "error": "not found"}, status=404)
                else:
                    self.respond_json({"ok": True, "name": name, "content": content})
            else:
                self.respond_json({"ok": True, "files": list_canvas_files()})
            return
        if self.path.startswith("/api/sessions"):
            from urllib.parse import urlparse, parse_qs
            qs = parse_qs(urlparse(self.path).query)
            sid = (qs.get("id") or [""])[0]
            if sid:
                data = load_session(sid)
                if data is None:
                    self.respond_json({"ok": False, "error": "not found"}, status=404)
                else:
                    self.respond_json({"ok": True, "session": data})
            else:
                self.respond_json({"ok": True, "sessions": list_sessions()})
            return
        if self.path == "/api/tts-voices":
            global _PIPER_DYNAMIC_VOICE_NAMES, _PIPER_LAST_VOICE_NAMES
            voices: list[dict] = []
            if KOKORO_URL:
                dynamic_names = kokoro_voice_names()
                _PIPER_DYNAMIC_VOICE_NAMES = set(dynamic_names)
                _PIPER_LAST_VOICE_NAMES = set(dynamic_names)
                if dynamic_names:
                    voices.extend(
                        item for item in KOKORO_FEMALE_VOICES
                        if item["name"] == "preset:lilith_dark" or item["name"] in _PIPER_DYNAMIC_VOICE_NAMES
                    )
                else:
                    voices.extend(KOKORO_FEMALE_VOICES)
                known = {item["name"] for item in voices}
                for name in dynamic_names:
                    if name not in known:
                        label = name.replace("_", " ").replace("-", " ").title()
                        voices.append({"name": name, "label": f"{label} — Piper offline"})
            voices.extend(EDGE_TTS_FEMALE_VOICES)
            ok = bool(KOKORO_URL) or EDGE_TTS_AVAILABLE
            self.respond_json({"ok": ok, "voices": voices, "kokoro": bool(KOKORO_URL)})
            return
        if self.path == "/api/auth/me":
            self.respond_json({"ok": True, "user": self._current_user()})
            return
        if self.path == "/api/tools":
            self.respond_json({"ok": True, "tools": [asdict(item) for item in all_metadata()]})
            return
        if self.path == "/api/setup-state":
            self.respond_json({
                "ok": True,
                "has_user": has_users(),
                "username": get_first_username(),
                "auth_enabled": get_auth_enabled(),
                "config": _current_setup_config(),
            })
            return
        if self.path == "/api/runtime-status":
            self.respond_json({"ok": True, **_runtime_status_snapshot()})
            return
        self.send_error(404)

    def do_POST(self) -> None:
        global _SETUP_DONE_RUNTIME
        if not self._require_access():
            return
        if self.path == "/api/thunderbird/analyze":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                if length > 2_000_000:
                    self.respond_json({"ok": False, "error": "payload too large"}, status=413)
                    return
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                result = handle_thunderbird_analyze(payload)
                self.respond_json(result, status=200 if result.get("ok") else 403)
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=500)
            return
        if self.path == "/api/setup-auth":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                u = str(payload.get("username") or "")
                p = str(payload.get("password") or "")
                if authenticate(u, p) is None:
                    self.respond_json({"ok": False, "error": "Invalid credentials."}, status=401)
                    return
                self.respond_json({"ok": True, "token": _issue_setup_token(u)})
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/setup-reset":
            try:
                delete_all_users()
                reset_runtime_state()
                rotate_session_secret()
                reload_config()
                _SETUP_DONE_RUNTIME = False
                self.respond_json({"ok": True}, cookie="")
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=500)
            return
        if self.path == "/api/setup":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                ollama_url = normalize_ollama_url(str(payload.get("ollama_url") or ""))
                if not ollama_url:
                    self.respond_json({"ok": False, "error": "Ollama API URL is required"}, status=400)
                    return
                auth_enabled = bool(payload.get("auth_enabled", True))
                # If sign-in is currently enabled and a user exists, require a fresh setup-auth token.
                if get_auth_enabled() and has_users():
                    token = str(payload.get("auth_token") or "")
                    auth_user = _consume_setup_token(token) if token else None
                    if not auth_user:
                        self.respond_json({"ok": False, "error": "Re-authentication required."}, status=401)
                        return
                username = str(payload.get("username") or "admin").strip() or "admin"
                new_password = str(payload.get("password") or "")
                if auth_enabled and has_users():
                    if new_password:
                        if len(new_password) < 8:
                            self.respond_json({"ok": False, "error": "Password must be at least 8 characters."}, status=400)
                            return
                        # Username may have changed. Wipe old users.json so we
                        # don't leave stale entries, then create the single user.
                        delete_all_users()
                        create_admin_user(username, new_password)
                    elif username != get_first_username():
                        # Username changed but password not provided — re-hash
                        # is impossible without the current password. Refuse.
                        self.respond_json({"ok": False, "error": "Provide a password to rename the user."}, status=400)
                        return
                elif auth_enabled:
                    if not new_password:
                        self.respond_json({"ok": False, "error": "Password is required."}, status=400)
                        return
                    create_admin_user(username, new_password)

                # Decide what to do with the API key field.
                api_key = str(payload.get("ollama_api_key") or "")
                if bool(payload.get("keep_existing_api_key")) and not api_key:
                    # Re-use the existing key file by reading current value.
                    from ollama_tools.config import get_ollama_api_key as _get_key
                    api_key = _get_key()

                write_first_run_config(
                    ollama_url=ollama_url,
                    ollama_api_key=api_key,
                    work_dir=str(payload.get("work_dir") or "/workspace"),
                    enable_ssh=bool(payload.get("enable_ssh")),
                    enable_local=bool(payload.get("enable_local")),
                    enable_search=bool(payload.get("enable_search")),
                    enable_mcp=bool(payload.get("enable_mcp")),
                    enable_thunderbird=bool(payload.get("enable_thunderbird")),
                    dangerous_mode=bool(payload.get("dangerous_mode")),
                    piper_url=str(payload.get("piper_url") or ""),
                    searxng_url=str(payload.get("searxng_url") or ""),
                    brave_api_key=str(payload.get("brave_api_key") or ""),
                    ssh_config_path=str(payload.get("ssh_config_path") or "/data/.ssh/config"),
                    ssh_known_hosts_path=str(payload.get("ssh_known_hosts_path") or "/data/.ssh/known_hosts"),
                    auth_enabled=auth_enabled,
                )
                reload_config()
                # Force every existing browser session to log in again so the
                # new config is in effect from the next session forward.
                rotate_session_secret()
                _SETUP_TOKENS.clear()
                _SETUP_DONE_RUNTIME = True
                self.respond_json({
                    "ok": True,
                    "thunderbird_token": get_thunderbird_token() if get_thunderbird_enabled() else "",
                }, cookie="")
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/auth/login":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                user = authenticate(str(payload.get("username") or ""), str(payload.get("password") or ""))
                if not user:
                    self.respond_json({"ok": False, "error": "invalid login"}, status=401)
                    return
                self.respond_json({"ok": True, "user": {"username": user.username, "role": user.role}}, cookie=create_session(user))
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/auth/logout":
            self.respond_json({"ok": True}, cookie="")
            return
        if self.path == "/api/tools":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                kind = str(payload.get("kind") or "tool")
                name = str(payload.get("name") or "")
                enabled = bool(payload.get("enabled"))
                allowed = {item.name: item.kind for item in all_metadata()}
                if not name or allowed.get(name) != kind:
                    self.respond_json({"ok": False, "error": "unknown tool entry"}, status=400)
                    return
                if kind == "hook":
                    set_hook_override(name, enabled)
                elif kind == "skill":
                    set_skill_override(name, enabled)
                else:
                    set_tool_override(name, enabled)
                self.respond_json({"ok": True, "tools": [asdict(item) for item in all_metadata()]})
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/tts":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                if length > 32_000:
                    self.respond_json({"ok": False, "error": "payload too large"}, status=413)
                    return
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                text = clean_text_for_tts(str(payload.get("text", "")).strip())[:4000]
                # Whitelist voice to known safe names only
                local_voices = _PIPER_DYNAMIC_VOICE_NAMES or _PIPER_LAST_VOICE_NAMES
                valid_voices = _EDGE_VOICE_NAMES | local_voices | {"preset:lilith_dark"}
                raw_voice = str(payload.get("voice", "en-US-AriaNeural"))
                if raw_voice in valid_voices:
                    voice = raw_voice
                elif KOKORO_URL:
                    voice = "preset:lilith_dark"
                else:
                    voice = "en-US-AriaNeural"
                # Validate rate format: optional sign, digits, percent
                raw_rate = str(payload.get("rate", "+0%"))
                rate = raw_rate if re.fullmatch(r"[+-]?\d{1,3}%", raw_rate) else "+0%"
                pitch = _clamp_int(payload.get("pitch"), -12, 12, 0)
                tone = str(payload.get("tone") or "natural")
                if tone not in {"natural", "dark", "bright", "radio", "robotic"}:
                    tone = "natural"
                volume = _clamp_int(payload.get("volume"), -50, 50, 0)
                if not KOKORO_URL and not EDGE_TTS_AVAILABLE:
                    self.respond_json({"ok": False, "error": "No TTS backend available"}, status=503)
                    return
                if not text:
                    self.respond_json({"ok": False, "error": "no text"}, status=400)
                    return
                audio, audio_ct = tts_speak(text, voice, rate, pitch_steps=pitch, tone=tone, volume=volume)
                self.send_response(200)
                self.send_header("Content-Type", audio_ct)
                self.send_header("Content-Length", str(len(audio)))
                self.end_headers()
                self.wfile.write(audio)
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=500)
            return
        if self.path == "/api/chat-stream":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                # Allow up to 20 MB for image/document attachments
                if length > 20_000_000:
                    self.respond_json({"ok": False, "error": "payload too large"}, status=413)
                    return
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                self.respond_event_stream(payload)
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=500)
            return
        if self.path == "/api/canvas-files/save":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                if length > 5_000_000:
                    self.respond_json({"ok": False, "error": "payload too large"}, status=413)
                    return
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                name = str(payload.get("name") or "").strip()
                content = str(payload.get("content") or "")
                meta = save_canvas_file(name, content)
                if not meta:
                    self.respond_json({"ok": False, "error": "invalid name"}, status=400)
                else:
                    self.respond_json({"ok": True, "file": meta})
            except (ValueError, OSError) as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/canvas-files/delete":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                ok = delete_canvas_file(str(payload.get("name") or ""))
                self.respond_json({"ok": ok})
            except (ValueError, OSError) as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/sessions/save":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                if length > 5_000_000:
                    self.respond_json({"ok": False, "error": "payload too large"}, status=413)
                    return
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                summary = save_session(payload)
                self.respond_json({"ok": True, "session": summary})
            except (ValueError, OSError) as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/sessions/delete":
            try:
                length = int(self.headers.get("Content-Length", "0"))
                payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
                ok = delete_session(str(payload.get("id") or ""))
                self.respond_json({"ok": ok})
            except (ValueError, OSError) as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=400)
            return
        if self.path == "/api/transcribe":
            try:
                if not is_whisper_available():
                    self.respond_json({"ok": False, "error": "Whisper not available. Install faster-whisper."}, status=503)
                    return
                length = int(self.headers.get("Content-Length", "0"))
                if length > 50_000_000:
                    self.respond_json({"ok": False, "error": "Audio too large"}, status=413)
                    return
                audio_bytes = self.rfile.read(length)
                content_type = self.headers.get("Content-Type", "audio/webm")
                text = transcribe_audio(audio_bytes, content_type)
                self.respond_json({"ok": True, "text": text})
            except Exception as exc:
                self.respond_json({"ok": False, "error": str(exc)}, status=500)
            return
        if self.path != "/api/chat":
            self.send_error(404)
            return
        try:
            length = int(self.headers.get("Content-Length", "0"))
            payload = json.loads(self.rfile.read(length).decode("utf-8") or "{}")
            self.respond_json(handle_chat(payload))
        except Exception as exc:
            self.respond_json({"ok": False, "error": str(exc)}, status=500)

    def _send_security_headers(self) -> None:
        self.send_header("X-Content-Type-Options", "nosniff")
        self.send_header("X-Frame-Options", "DENY")
        self.send_header("X-XSS-Protection", "1; mode=block")
        self.send_header("Referrer-Policy", "no-referrer")
        self.send_header("Content-Security-Policy",
            "default-src 'self'; script-src 'unsafe-inline'; style-src 'unsafe-inline'; media-src blob:")

    def respond_text(self, text: str, content_type: str, status: int = 200) -> None:
        body = text.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Length", str(len(body)))
        if content_type.startswith("text/html"):
            self.send_header("Cache-Control", "no-store")
        self._send_security_headers()
        self.end_headers()
        self.wfile.write(body)

    def redirect(self, location: str) -> None:
        self.send_response(302)
        self.send_header("Location", location)
        self._send_security_headers()
        self.end_headers()

    def respond_json(self, data: dict[str, Any], status: int = 200, cookie: str | None = None) -> None:
        body = json.dumps(data).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        if self.path == "/api/thunderbird/analyze":
            self.send_header("Access-Control-Allow-Origin", "*")
        if cookie is not None:
            if cookie:
                secure = "; Secure" if get_cookie_secure() else ""
                same_site = get_cookie_same_site()
                self.send_header("Set-Cookie", f"ollama_hooks_session={cookie}; Path=/; HttpOnly; SameSite={same_site}{secure}")
            else:
                self.send_header("Set-Cookie", "ollama_hooks_session=; Path=/; Max-Age=0; HttpOnly; SameSite=Lax")
        self._send_security_headers()
        self.end_headers()
        self.wfile.write(body)

    def respond_event_stream(self, payload: dict[str, Any]) -> None:
        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Connection", "close")
        self.end_headers()
        gen = stream_chat_events(payload)
        try:
            for event in gen:
                data = "data: " + json.dumps(event) + "\n\n"
                try:
                    self.wfile.write(data.encode("utf-8"))
                    self.wfile.flush()
                except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError):
                    # Client disconnected (Stop button). Close the generator so the
                    # underlying Ollama HTTP connection is closed and the model
                    # stops generating tokens.
                    print("[web] client disconnected mid-stream — aborting generation", file=sys.stderr)
                    break
        finally:
            gen.close()

    def log_message(self, fmt: str, *args: Any) -> None:
        print(f"[web] {self.address_string()} - {fmt % args}", file=sys.stderr)


def _wrap_tls(server: ThreadingHTTPServer) -> ThreadingHTTPServer:
    import ssl
    ctx = ssl.SSLContext(ssl.PROTOCOL_TLS_SERVER)
    ctx.load_cert_chain(certfile=CERT_FILE, keyfile=KEY_FILE)
    server.socket = ctx.wrap_socket(server.socket, server_side=True)
    return server


def create_server() -> tuple[ThreadingHTTPServer, int]:
    last_error: OSError | None = None
    for port in range(PORT, PORT + max(1, PORT_SEARCH_LIMIT)):
        try:
            server = ThreadingHTTPServer((HOST, port), Handler)
            if TLS_ENABLED:
                server = _wrap_tls(server)
            return server, port
        except OSError as exc:
            last_error = exc
            if exc.errno != 98:
                raise
    raise OSError(
        f"Could not bind {HOST}:{PORT}-{PORT + max(1, PORT_SEARCH_LIMIT) - 1}: {last_error}"
    )


def display_host(host: str) -> str:
    return socket.gethostname() if host == "0.0.0.0" else host


def main() -> int:
    server, port = create_server()
    scheme = "https" if TLS_ENABLED else "http"
    print(f"Ollama Web TUI: {scheme}://{display_host(HOST)}:{port}")
    if port != PORT:
        print(f"Port {PORT} was in use, so using {port}.")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nStopping web server.")
    finally:
        server.server_close()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
